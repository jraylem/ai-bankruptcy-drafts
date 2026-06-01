"""DOCX template creation + fill service.

Two public operations drive the template lifecycle:

  create_template  — replace raw source values (template_property_marker) with
                     placeholders ([[variable_name]]) to produce the template.docx.
  fill_template    — replace placeholders back with resolved values to produce
                     the filled output.docx at draft/dry-run time.

Replacement is typography-tolerant (curly quotes, dashes, NBSP folded to ASCII)
and paragraph-aware (markers that straddle multiple runs still get replaced).
"""

import logging
import os
import re
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.core.agents.types.sources import FieldSource, GroupDropdownComposite
from src.core.agents.types.spec import TemplateField, TemplateVariable

logger = logging.getLogger(__name__)


class DocxTemplateService:
    """Create template docx files (markers → placeholders) and fill them (placeholders → resolved values) with typography-tolerant, run-straddle-aware replacement."""

    _GROUP_DROPDOWN_SOURCES = {
        FieldSource.GROUP_DROPDOWN_FROM_GMAIL,
        FieldSource.GROUP_DROPDOWN_FROM_COURT_DRIVE,
    }

    _TYPOGRAPHY_MAP = str.maketrans({
        "\u2018": "'",   # LEFT SINGLE QUOTATION MARK
        "\u2019": "'",   # RIGHT SINGLE QUOTATION MARK (Word-auto apostrophe)
        "\u201C": '"',   # LEFT DOUBLE QUOTATION MARK
        "\u201D": '"',   # RIGHT DOUBLE QUOTATION MARK
        "\u2013": "-",   # EN DASH
        "\u2014": "-",   # EM DASH
        "\u2026": "...", # HORIZONTAL ELLIPSIS
        "\u00A0": " ",   # NON-BREAKING SPACE
    })

    @staticmethod
    def _normalize_typography(text: str) -> str:
        """Fold curly quotes, dashes, ellipsis, and NBSP to their ASCII equivalents.

        Word auto-corrects straight quotes to curly quotes and single/double
        hyphens to dashes; LLM-extracted markers typically use ASCII.
        Without this fold, `'` (U+0027) in a marker will never match `’`
        (U+2019) in the DOCX, and substitution silently no-ops.
        """
        return text.translate(DocxTemplateService._TYPOGRAPHY_MAP)

    @staticmethod
    def _iter_all_paragraphs(doc):
        """Yield every paragraph in the document, body paragraphs first, then table-cell paragraphs.

        Flattens the table walk so callers don't need a 4-level nested loop.
        """
        yield from doc.paragraphs
        yield from (
            paragraph
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        )

    @staticmethod
    def extract_all_text(template_bytes: bytes) -> str:
        """Return the full visible text of a docx as a newline-joined string.

        Walks body paragraphs first, then every table-cell paragraph, in
        document order. Used by the bundling engine to feed a parent's
        produced draft text into the child's `extract_from_draft` slot
        resolution — the LLM needs the actual filed prose (real names,
        case numbers, the literal docket title) rather than the
        template's authoring markers.

        Empty paragraphs are dropped to keep the output dense.
        """
        doc = Document(BytesIO(template_bytes))
        return "\n".join(
            p.text for p in DocxTemplateService._iter_all_paragraphs(doc) if p.text
        )

    @staticmethod
    def flatten_word_fields(file_content: bytes) -> bytes:
        """Strip Word field codes from a .docx, freezing each field's
        cached text as plain content.

        Why: Word `{ DATE }` field codes (and `{ TIME }`, `{ FILENAME }`,
        etc.) re-evaluate every time the doc is opened, so a source
        document with `Dated: April 8, 2026` rendered via a DATE field
        would display TODAY's date in any viewer (Word, Syncfusion,
        OnlyOffice). The agent extracts the cached fallback ("April 8,
        2026") at compose time, but the field continues to live-update
        in the rendered template.docx — confusing paralegals and
        breaking the assumption that the marker the agent saw IS what
        renders.

        Flattening removes the field machinery (`<w:fldSimple>`,
        `<w:fldChar>`, `<w:instrText>`) and content controls (`<w:sdt>`)
        while preserving the cached result text. After this pass:
          - Every viewer renders the same frozen text.
          - The agent's marker matches what create_template will find
            and replace at compose time AND what paralegals see in any
            preview thereafter.
          - No information is lost beyond the field auto-update
            behavior, which is what we want to kill.

        Should be called on the uploaded bytes BEFORE
        `parse_document_v2` and BEFORE persisting to R2 — see
        `services/composer/generate.py` callsite.
        """
        from lxml import etree

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        import zipfile
        in_zip = zipfile.ZipFile(BytesIO(file_content), "r")
        out_buf = BytesIO()
        out_zip = zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED)

        for item in in_zip.infolist():
            raw = in_zip.read(item.filename)
            # Only touch the main document part + headers/footers; other
            # parts (styles, theme, etc.) are unaffected.
            if item.filename == "word/document.xml" or (
                item.filename.startswith("word/") and (
                    item.filename.startswith("word/header") or
                    item.filename.startswith("word/footer")
                ) and item.filename.endswith(".xml")
            ):
                try:
                    tree = etree.fromstring(raw)
                except etree.XMLSyntaxError:
                    # Defensive: malformed part, pass through unchanged.
                    out_zip.writestr(item, raw)
                    continue

                # 1. <w:fldSimple w:instr="..."> wraps cached result runs
                #    in a single element. Replace the element with its
                #    children, dropping the wrapper + instruction.
                for fld in tree.iter(W + "fldSimple"):
                    parent = fld.getparent()
                    if parent is None:
                        continue
                    idx = list(parent).index(fld)
                    # Splice children of fldSimple into parent at same
                    # position; preserve document order.
                    for child in list(fld):
                        parent.insert(idx, child)
                        idx += 1
                    parent.remove(fld)

                # 2. Complex field machinery:
                #    <w:r><w:fldChar fldCharType="begin"/></w:r>
                #    <w:r><w:instrText>{ DATE \@ "MMMM d, yyyy" }</w:instrText></w:r>
                #    <w:r><w:fldChar fldCharType="separate"/></w:r>
                #    <w:r><w:t>April 8, 2026</w:t></w:r>     ← keep this
                #    <w:r><w:fldChar fldCharType="end"/></w:r>
                #
                # Strategy: remove every <w:r> that CONTAINS a <w:fldChar>
                # or an <w:instrText> — they carry only machinery, no
                # user-visible text. The cached result runs (the plain
                # <w:r><w:t>...) between separate/end markers are left
                # untouched by this filter, so they survive as plain text.
                runs_to_remove = []
                for r in tree.iter(W + "r"):
                    if r.find(W + "fldChar") is not None or r.find(W + "instrText") is not None:
                        runs_to_remove.append(r)
                for r in runs_to_remove:
                    parent = r.getparent()
                    if parent is not None:
                        parent.remove(r)

                # 3. <w:sdt> content controls wrap content. Replace the
                #    <w:sdt> with the children of its <w:sdtContent> so
                #    the wrapped runs survive as plain content.
                for sdt in tree.iter(W + "sdt"):
                    parent = sdt.getparent()
                    if parent is None:
                        continue
                    idx = list(parent).index(sdt)
                    content_elem = sdt.find(W + "sdtContent")
                    if content_elem is not None:
                        for child in list(content_elem):
                            parent.insert(idx, child)
                            idx += 1
                    parent.remove(sdt)

                # Serialize back. Preserve the XML declaration the docx
                # spec requires (standalone="yes").
                new_xml = etree.tostring(
                    tree,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
                out_zip.writestr(item, new_xml)
            else:
                out_zip.writestr(item, raw)

        in_zip.close()
        out_zip.close()
        return out_buf.getvalue()

    @staticmethod
    def find_paragraph_containing(template_bytes: bytes, placeholder: str) -> str | None:
        """Return the text of the paragraph containing the given `[[placeholder]]` marker, or None if not found.

        Used by ChipFitAgent to scope heal context to a single paragraph so
        the LLM sees the exact surrounding sentence structure the filled
        value must integrate with.
        """
        if not placeholder:
            return None
        doc = Document(BytesIO(template_bytes))

        for paragraph in DocxTemplateService._iter_all_paragraphs(doc):
            if placeholder in paragraph.text:
                return paragraph.text
        return None

    @staticmethod
    def find_missing_placeholders(
        template_bytes: bytes,
        placeholders: list[str],
    ) -> set[str]:
        """Return the set of placeholders missing from the template document body and table cells.

        Used by the template-generation pipeline to drop orphan variables —
        those whose marker was hallucinated by the LLM, or whose marker was
        swallowed by a longer marker during create_template's longest-first
        replacement pass (e.g. 'her' extracted as its own variable when it
        lived inside another variable's marker — the parent's longer marker
        consumes the span first, so '[[her_var]]' never makes it into the
        rendered template.docx).
        """
        if not placeholders:
            return set()
        doc = Document(BytesIO(template_bytes))
        full_text = "\n".join(
            p.text for p in DocxTemplateService._iter_all_paragraphs(doc)
        )
        return {p for p in placeholders if p and p not in full_text}

    @staticmethod
    def extract_table_rows(table) -> list[str]:
        """Extract text content from a python-docx Table object.

        Iterates through each row and joins non-empty cell text with ' | ' so
        the row survives as a readable single line in the parsed-document
        string fed to the template agent. Empty rows are excluded.

        Example::

            | Name     | Value  |
            | Debtor   | John   |

            -> ["Name | Value", "Debtor | John"]
        """
        rows = []
        for row in table.rows:
            cell_texts = [c.text.strip() for c in row.cells if c.text.strip()]
            if cell_texts:
                rows.append(" | ".join(cell_texts))
        return rows

    @staticmethod
    def create_template(
        file_content: bytes,
        template_spec: list[TemplateVariable],
    ) -> bytes:
        """
        Replace template_property_marker (and any template_property_marker_aliases)
        with template_variable_string in a DOCX document.

        Handles text split across multiple runs due to formatting (bold, italic, superscript, etc).

        Each variable may carry a PRIMARY marker plus zero or more ALIAS markers
        (orthographic variants of the same value, e.g. 'Jane S Smith' vs
        'Jane S. Smith'). All markers across all variables are flattened into
        a single (marker, replacement) job list and processed longest-first
        GLOBALLY — so a short marker from variable A cannot consume a
        substring of a longer marker from variable B (e.g. "13" for chapter
        vs "April 13, 2026" for document_date), regardless of which variable
        owns each marker.

        Shared-marker disambiguation: when two or more variables share the
        same marker value (e.g. `civil_case_number` and `bankruptcy_case_number`
        both with marker "25-24573" because the source doc reused the value),
        each variable's `template_identifying_text_match` is used to locate
        the specific paragraph that should receive its placeholder. Unique-
        marker variables (the overwhelming common case) follow the existing
        global longest-first find-replace path unchanged.
        """
        doc = Document(BytesIO(file_content))

        # Two-pass scanner→resolver→mutator: collision-aware substitution
        # that defers short markers when a longer marker contains them
        # (handles the Fleisher CoS bug where cos_email_section_1's
        # cross-paragraph match fails and trustee_name steals its
        # region). Behind env flag for a controlled rollout — the legacy
        # immediate-mutation path stays the default until v2 has bake
        # time against the regression suite.
        if os.getenv("MARKER_SUBSTITUTION_V2") == "1":
            from .marker_substitution import substitute_markers
            substitute_markers(doc, template_spec)
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            template_bytes = output.read()
            DocxTemplateService._maybe_run_format_validator(
                file_content, template_bytes, template_spec,
            )
            return template_bytes

        # Group (replacement, identifying_text) entries by marker string.
        # Aliases of the same variable share its replacement, so they land
        # in the same group with identical replacements. Two DIFFERENT
        # variables that share the same marker land in the same group with
        # DIFFERENT replacements — that's the shared-marker case.
        groups: dict[str, list[tuple[str, str | None]]] = {}
        for variable in template_spec:
            replacement = variable.template_variable_string
            if not replacement:
                continue
            identifying_text = variable.template_identifying_text_match
            all_markers = [variable.template_property_marker] + list(variable.template_property_marker_aliases)
            for marker in all_markers:
                if marker:
                    groups.setdefault(marker, []).append((replacement, identifying_text))

        # Process markers longest-first GLOBALLY (preserves the existing
        # contract that "April 13, 2026" beats "13").
        ordered_markers = sorted(groups.keys(), key=len, reverse=True)
        consumed_paragraphs: set[int] = set()

        for marker in ordered_markers:
            entries = groups[marker]
            # Tabs are folded to spaces (docx captions often glue 'In re:' to
            # 'Case No:' via \t, but the LLM-emitted marker rarely carries tabs).
            # Newlines are PRESERVED in the marker so joint-debtor captions where
            # two names are separated by <w:br/> in a single paragraph can be
            # matched by stage 3 of _replace_in_paragraph.
            clean_marker = marker.replace("\t", " ").replace("\\n", "\n")
            distinct_replacements = {r for r, _ in entries}

            if len(distinct_replacements) == 1:
                # Common case: one variable owns this marker (with or without
                # aliases that share its placeholder). Global find-replace.
                clean_replacement = entries[0][0].replace("\t", " ").replace("\n", " ")
                DocxTemplateService._replace_in_document(doc, clean_marker, clean_replacement)
                continue

            # Shared-marker case: 2+ variables share this marker. Each lands
            # on the paragraph its identifying_text_match points to, with
            # consumed_paragraphs preventing collisions within the group.
            for replacement, identifying_text in entries:
                clean_replacement = replacement.replace("\t", " ").replace("\n", " ")
                DocxTemplateService._replace_first_in_context(
                    doc, clean_marker, clean_replacement, identifying_text, consumed_paragraphs,
                )

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        template_bytes = output.read()
        DocxTemplateService._maybe_run_format_validator(
            file_content, template_bytes, template_spec,
        )
        return template_bytes

    @staticmethod
    def _maybe_run_format_validator(
        original_bytes: bytes,
        template_bytes: bytes,
        template_spec: list[TemplateVariable],
    ) -> None:
        """Tier 1 format-drift gate: deterministic structural diff
        between the original `.docx` and the placeholder-marked
        output. Logs INFO when drift detected so the composer flow
        surfaces tab/break/paragraph drift before R2 upload.

        Behind `TEMPLATE_FORMAT_AUTOFIX_V2` env flag (default OFF).
        Tier 2 (LLM-assisted auto-fix) will land as a follow-up; this
        pass only reports — never modifies bytes, never blocks upload.
        Hard-fails are silent: any exception in the validator is
        swallowed with a WARNING so a validator bug never breaks the
        composer.
        """
        if os.getenv("TEMPLATE_FORMAT_AUTOFIX_V2") != "1":
            return
        try:
            from .template_format_validator import validate_template_format
            result = validate_template_format(
                original_bytes, template_bytes, template_spec,
            )
            if not result.ok:
                logger.info(
                    "TemplateFormatValidator: %s — drifted paragraphs (up to 5 shown): %s",
                    result.summary,
                    [
                        {
                            "idx": d.template_paragraph_index,
                            "tmpl": d.template_text[:120],
                            "reconstructed": d.reconstructed_text[:120],
                        }
                        for d in result.drifted_paragraphs[:5]
                    ],
                )
        except Exception as exc:  # noqa: BLE001 — opportunistic; never block composer
            logger.warning("TemplateFormatValidator raised %s; ignoring", exc)

    @staticmethod
    def _replace_first_in_context(
        doc,
        marker: str,
        replacement: str,
        identifying_text: str | None,
        consumed_paragraphs: set[int],
    ) -> None:
        """Place a single placeholder for a shared-marker variable on the
        paragraph indicated by `identifying_text`, replacing only the FIRST
        occurrence of `marker` there.

        Disambiguation order:
          1. If `identifying_text` is provided, prefer an unconsumed paragraph
             whose normalized text (lowercased, whitespace-collapsed) contains
             the normalized identifying_text AND the marker.
          2. Fall back to the first unconsumed paragraph that contains the
             marker, logging a warning so the spec author can tighten the
             identifying_text_match.
          3. If no unconsumed occurrence exists, log a warning. The placeholder
             will be missing from the rendered docx and `_drop_orphan_variables`
             will subsequently drop the variable.

        Tracks consumed paragraphs by INDEX in `_iter_all_paragraphs(doc)`
        order rather than by `id(p._element)` — lxml wraps the same underlying
        XML element with different proxy objects across iterations, so id()
        identity is not preserved.
        """
        paragraphs = list(DocxTemplateService._iter_all_paragraphs(doc))

        def _normalize(s: str) -> str:
            return " ".join(s.lower().split())

        normalized_marker = DocxTemplateService._normalize_typography(marker)

        target_index: int | None = None
        if identifying_text and identifying_text.strip():
            needle = _normalize(identifying_text)
            for idx, p in enumerate(paragraphs):
                if idx in consumed_paragraphs:
                    continue
                if not p.text:
                    continue
                normalized_p_text = DocxTemplateService._normalize_typography(p.text)
                if marker not in p.text and normalized_marker not in normalized_p_text:
                    continue
                hay = _normalize(p.text)
                if needle in hay or hay in needle:
                    target_index = idx
                    break

        if target_index is None:
            for idx, p in enumerate(paragraphs):
                if idx in consumed_paragraphs:
                    continue
                if not p.text:
                    continue
                normalized_p_text = DocxTemplateService._normalize_typography(p.text)
                if marker in p.text or normalized_marker in normalized_p_text:
                    target_index = idx
                    if identifying_text:
                        logger.warning(
                            "Shared-marker disambiguation fallback: identifying_text_match "
                            "%r did not match any unconsumed paragraph for marker %r; using "
                            "first unconsumed occurrence.",
                            identifying_text, marker,
                        )
                    break

        if target_index is None:
            logger.warning(
                "Shared-marker placement: marker %r has no unconsumed occurrence in the "
                "document; placeholder %r will be orphaned and dropped by "
                "_drop_orphan_variables.",
                marker, replacement,
            )
            return

        consumed_paragraphs.add(target_index)
        DocxTemplateService._replace_in_paragraph_first_only(paragraphs[target_index], marker, replacement)

    @staticmethod
    def _replace_in_paragraph_first_only(paragraph, marker: str, replacement: str) -> None:
        """Variant of `_replace_in_paragraph` that replaces only the FIRST
        occurrence of `marker` within the paragraph. Used by shared-marker
        disambiguation so each variable's placeholder lands on exactly one
        occurrence, even when the marker appears multiple times in one
        paragraph.

        Mirrors stages 1 and 2 of `_replace_in_paragraph` (per-run direct,
        paragraph-level merge). Stage 3 (span-with-breaks for joint-debtor
        `\\n`-markers) is NOT mirrored — shared-marker disambiguation
        operates on single-line markers; joint-debtor patterns use unique
        markers and continue to go through the fast path.
        """
        p_element = paragraph._element
        if not marker:
            return

        normalized_marker = DocxTemplateService._normalize_typography(marker)

        # Stage 1: per-run direct, only first hit.
        for t_elem in p_element.iter(qn('w:t')):
            if not t_elem.text:
                continue
            if marker in t_elem.text:
                t_elem.text = t_elem.text.replace(marker, replacement, 1)
                return
            norm_text = DocxTemplateService._normalize_typography(t_elem.text)
            if normalized_marker in norm_text:
                t_elem.text = norm_text.replace(normalized_marker, replacement, 1)
                return

        # Stage 2: paragraph-level collapse, only first hit.
        t_elems = list(p_element.iter(qn('w:t')))
        if not t_elems:
            return
        combined = "".join(t.text or "" for t in t_elems)
        if marker in combined:
            t_elems[0].text = combined.replace(marker, replacement, 1)
            for t in t_elems[1:]:
                t.text = ""
            return

        norm_combined = DocxTemplateService._normalize_typography(combined)
        if normalized_marker in norm_combined:
            t_elems[0].text = norm_combined.replace(normalized_marker, replacement, 1)
            for t in t_elems[1:]:
                t.text = ""

    @staticmethod
    def fill_template(
        template_bytes: bytes,
        template_fields: list[TemplateField],
        resolved_values: dict[str, str],
    ) -> tuple[bytes, list[str]]:
        """
        Replace [[variable_name]] placeholders in a template DOCX with resolved values.

        Inverse of DocxTemplateService.create_template. Walks the template_fields,
        looks up each field's resolved value by property_name, and substitutes
        the template_variable_string placeholder with that value across all
        paragraphs and table cells (preserving formatting via
        _replace_in_document).

        Returns:
            - filled_docx_bytes: the rendered DOCX as bytes
            - unresolved_placeholders: list of [[var]] strings that had no resolved value
        """
        doc = Document(BytesIO(template_bytes))
        unresolved: list[str] = []

        # Pre-compute per-paragraph caption-shape decisions on the
        # doc's ORIGINAL state. Substituting placeholder #1 into a
        # paragraph adds its resolved value's text into the runs, so
        # by the time we check shape for placeholder #2 the residue
        # would contain real alphanums from #1's value and report
        # inline-shape incorrectly. Caching the decision pre-mutation
        # makes per-placeholder substitution decisions consistent
        # across the whole paragraph.
        caption_shape_map = DocxTemplateService._compute_caption_shape_map(doc)

        for field in template_fields:
            if field.source in DocxTemplateService._GROUP_DROPDOWN_SOURCES and isinstance(field.source_params, GroupDropdownComposite):
                params = field.source_params
                for placeholder, variable_name in (
                    (params.left_template_variable_string, params.left_variable),
                    (params.right_template_variable_string, params.right_variable),
                ):
                    if not placeholder:
                        continue
                    value = resolved_values.get(variable_name)
                    if value:
                        DocxTemplateService._substitute_placeholder(
                            doc, placeholder, value, caption_shape_map=caption_shape_map,
                        )
                    else:
                        unresolved.append(placeholder)
                continue

            placeholder = field.template_variable_string
            if not placeholder:
                continue

            value = resolved_values.get(field.property_name)
            if value:
                DocxTemplateService._substitute_placeholder(
                    doc, placeholder, value, caption_shape_map=caption_shape_map,
                )
            else:
                unresolved.append(placeholder)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        filled_bytes = output.read()
        filled_bytes = DocxTemplateService._maybe_run_fill_validator(
            template_bytes, filled_bytes, resolved_values,
        )
        return filled_bytes, unresolved

    @staticmethod
    def _maybe_run_fill_validator(
        template_bytes: bytes,
        filled_bytes: bytes,
        resolved_values: dict[str, str],
    ) -> bytes:
        """Sync Tier 1 format-drift detector. Logs drift signals as
        INFO; does NOT call the LLM fixer (that path is async — see
        `maybe_autofix_fill_async`).

        Wired into `fill_template` (sync v1) and `_fill_template_v2`
        (sync v2) for opportunistic logging. The async caller
        (`finalize_run_v2`) does the Tier 2 LLM repair afterwards.
        """
        if os.getenv("TEMPLATE_FORMAT_AUTOFIX_V2") != "1":
            return filled_bytes
        try:
            from .template_format_validator import validate_fill_format
            result = validate_fill_format(
                template_bytes, filled_bytes, resolved_values,
            )
            if result.ok:
                return filled_bytes
            logger.info(
                "TemplateFillValidator: %s — drifted paragraphs (up to 5 shown): %s",
                result.summary,
                [
                    {
                        "idx": d.template_paragraph_index,
                        "tmpl": d.template_text[:120],
                        "filled": d.reconstructed_text[:120],
                    }
                    for d in result.drifted_paragraphs[:5]
                ],
            )
        except Exception as exc:  # noqa: BLE001 — opportunistic; never block fill
            logger.warning("TemplateFillValidator raised %s; ignoring", exc)
        return filled_bytes

    @staticmethod
    async def maybe_autofix_fill_async(
        template_bytes: bytes,
        filled_bytes: bytes,
        resolved_values: dict[str, str],
    ) -> bytes:
        """Async Tier 1 validator + Tier 2 LLM fixer.

        Runs the same `validate_fill_format` as the sync gate (so the
        async pipeline doesn't depend on the sync gate having fired
        first), then if drift detected calls Sonnet 4.6 to repair
        each drifted paragraph. Returns the (potentially fixed)
        bytes — falls back to the original `filled_bytes` if any
        step soft-fails.

        Use from async contexts (e.g. `finalize_run_v2`). The httpx
        connection pool inside LangChain's ChatAnthropic is tied to
        the calling event loop, so we MUST run the LLM call in the
        same loop the caller awaits us from — not in a thread-spawned
        fresh loop, which produces "Connection error" on the bridge.

        Behind the same `TEMPLATE_FORMAT_AUTOFIX_V2=1` env flag.
        """
        if os.getenv("TEMPLATE_FORMAT_AUTOFIX_V2") != "1":
            return filled_bytes
        try:
            from .template_format_validator import validate_fill_format
            result = validate_fill_format(
                template_bytes, filled_bytes, resolved_values,
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning("TemplateFillValidator (async) raised %s; skipping fix", exc)
            return filled_bytes
        if result.ok:
            return filled_bytes

        try:
            from .template_format_fixer import autofix_fill_drift
            fixed = await autofix_fill_drift(
                filled_bytes=filled_bytes,
                template_bytes=template_bytes,
                resolved_values=resolved_values,
                drifted_paragraphs=result.drifted_paragraphs,
            )
            if fixed != filled_bytes:
                logger.info(
                    "TemplateFormatFixer: applied LLM-driven repairs for %d drifted "
                    "paragraph(s); returning fixed bytes",
                    len(result.drifted_paragraphs),
                )
                return fixed
        except Exception as exc:  # noqa: BLE001
            logger.warning("TemplateFormatFixer raised %s; returning unfixed bytes", exc)
        return filled_bytes

    @staticmethod
    async def maybe_autofix_grammar_async(
        filled_bytes: bytes,
        resolved_values: dict[str, str],
    ) -> tuple[bytes, list]:
        """Async Tier 1 grammar validator + Tier 2 LLM grammar fixer.

        Runs `validate_fill_grammar` (cheap deterministic scan); if any
        suspect paragraphs surface, calls Sonnet 4.6 in a single
        document-level pass to repair the agreement mismatches.

        Returns `(post_fix_bytes, applied_repairs)`:
          - `post_fix_bytes` is the post-fix docx, or the original
            `filled_bytes` unchanged when nothing applied.
          - `applied_repairs` is a list of `GrammarRepairRecord`
            describing each swap that ACTUALLY landed — surfaced on
            `FinalizedRunV2.grammar_repairs` and shown in the FE
            Resolution Log so the paralegal can see exactly what
            changed.

        Behind `TEMPLATE_GRAMMAR_AUTOFIX_V2=1` env flag — independent
        of the format fixer's flag so each can be rolled separately.

        Use from async contexts (e.g. `finalize_run_v2`). Same
        event-loop discipline as `maybe_autofix_fill_async` — the
        httpx pool inside LangChain's ChatAnthropic is bound to the
        calling loop.
        """
        if os.getenv("TEMPLATE_GRAMMAR_AUTOFIX_V2") != "1":
            return filled_bytes, []
        try:
            from .template_grammar_validator import validate_fill_grammar
            result = validate_fill_grammar(filled_bytes, resolved_values)
        except Exception as exc:  # noqa: BLE001
            logger.warning(
                "TemplateGrammarValidator raised %s; skipping grammar fix", exc,
            )
            return filled_bytes, []
        if result.ok:
            return filled_bytes, []

        try:
            from .template_grammar_fixer import autofix_grammar_drift
            fixed, applied = await autofix_grammar_drift(
                filled_bytes=filled_bytes,
                resolved_values=resolved_values,
                suspect_paragraphs=result.suspect_paragraphs,
                cardinality_signals=result.cardinality_signals,
            )
            if applied:
                logger.info(
                    "TemplateGrammarFixer: applied %d agreement swap(s) "
                    "across %d paragraph(s) (out of %d suspect)",
                    len(applied),
                    len({r.paragraph_index for r in applied}),
                    len(result.suspect_paragraphs),
                )
                return fixed, applied
            logger.info(
                "TemplateGrammarFixer: no swaps applied across %d suspect "
                "paragraph(s) (all rejected by allowlist / case-style / "
                "missing-word checks)",
                len(result.suspect_paragraphs),
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning(
                "TemplateGrammarFixer raised %s; returning unfixed bytes", exc,
            )
        return filled_bytes, []

    @staticmethod
    def _substitute_placeholder(
        doc,
        placeholder: str,
        value: str,
        caption_shape_map: dict[int, bool] | None = None,
    ) -> None:
        """Dispatch placeholder substitution on whether the resolved value carries `\\n`.

        Values without a newline follow today's single-line path (with `\\t → space`
        normalization). Values containing `\\n` — joint-debtor names primarily —
        get a per-paragraph rendering decision because the same placeholder may
        appear in BOTH a caption-shape paragraph (placeholder is essentially the
        whole paragraph → render `\\n` as `<w:br/>` soft line break) AND inline
        body paragraphs ('The Debtor, [[name]], is employed' → join with ' and '
        so the sentence reads grammatically). The choice is made for each
        occurrence independently.

        `caption_shape_map`: optional pre-computed map (`id(paragraph._element) →
        is_caption_shape`) snapshot at fill_template kickoff, BEFORE any
        substitution mutates the runs. Avoids the trap where placeholder #1's
        resolved value adds alphanums to a paragraph and placeholder #2's live
        shape check then mis-classifies the paragraph as inline. Defaults to
        live `_is_caption_shape_paragraph` for back-compat with callers that
        haven't been threaded through fill_template's pre-snapshot path.
        """
        if "\n" not in value:
            cleaned = value.replace("\t", " ")
            DocxTemplateService._replace_in_document(doc, placeholder, cleaned)
            return

        lines = [line.replace("\t", " ") for line in value.split("\n")]
        inline_value = " and ".join(lines)
        for idx, paragraph in enumerate(DocxTemplateService._iter_all_paragraphs(doc)):
            if placeholder not in paragraph.text:
                continue
            if caption_shape_map is not None:
                is_caption = caption_shape_map.get(idx, False)
            else:
                is_caption = DocxTemplateService._is_caption_shape_paragraph(
                    paragraph, placeholder,
                )
            if is_caption:
                DocxTemplateService._render_lines_into_placeholder(
                    paragraph, placeholder, lines
                )
            else:
                DocxTemplateService._replace_in_paragraph(
                    paragraph, placeholder, inline_value
                )

    _PLACEHOLDER_TOKEN_RE = re.compile(r"\[\[[A-Za-z0-9_]+\]\]")

    @staticmethod
    def _compute_caption_shape_map(doc) -> dict[int, bool]:
        """Walk every paragraph in the doc's ORIGINAL state and return
        a map `paragraph_walk_index → is_caption_shape`. The cached
        decision survives subsequent substitution mutations so
        per-placeholder shape checks stay consistent across a
        multi-placeholder paragraph.

        Used by `fill_template` to pre-snapshot shape decisions BEFORE
        any `_substitute_placeholder` call runs. Without this, the
        first placeholder's resolved value gets written into the
        paragraph's runs, then the second placeholder's shape check
        sees those new alphanums and (incorrectly) routes to inline.

        Keyed by the paragraph's walk index because lxml returns
        fresh proxy objects for the same underlying `<w:p>` node on
        each access — `id(paragraph._element)` is NOT stable across
        walks. `_iter_all_paragraphs`'s body-then-tables order is
        deterministic, so walk index is a reliable cache key as long
        as paragraphs aren't added/removed between snapshot and use
        (which `_substitute_placeholder` never does — it only mutates
        run children, never `<w:p>` siblings).
        """
        shapes: dict[int, bool] = {}
        for idx, paragraph in enumerate(DocxTemplateService._iter_all_paragraphs(doc)):
            t_elems = list(paragraph._element.iter(qn("w:t")))
            combined = "".join(t.text or "" for t in t_elems)
            residue = DocxTemplateService._PLACEHOLDER_TOKEN_RE.sub("", combined)
            shapes[idx] = not any(ch.isalnum() for ch in residue)
        return shapes

    @staticmethod
    def _is_caption_shape_paragraph(paragraph, placeholder: str) -> bool:
        """Return True when the placeholder essentially IS the paragraph (caption-style).

        Caption-shape: residue after removing EVERY `[[…]]` placeholder
        token has no alphanumeric characters (just whitespace /
        punctuation / role markers like `, Debtors.`). Inline-shape:
        real prose surrounds the placeholder
        (`The Debtor, [[name]], is employed...`).

        Stripping ALL `[[…]]` tokens (not just the current one) matters
        for multi-placeholder paragraphs — e.g. a CoS recipients block
        like `[[cos_email_section_1]]\\n[[cos_email_section_2]]`.
        Removing only THIS placeholder would leave the SIBLING
        placeholder's characters in the residue, mis-classifying the
        paragraph as inline → multi-line resolved values get joined
        with " and " instead of rendering with `<w:br/>` (the Fleisher
        CoS bug). The smart strip restores caption-shape detection
        for those layouts.

        The check uses the same run-concatenated text the engine
        matches against, so a placeholder split across runs still
        resolves correctly. `placeholder` is accepted for backwards
        compatibility but no longer drives the residue computation —
        kept as a positional arg so call sites needn't change.
        """
        t_elems = list(paragraph._element.iter(qn("w:t")))
        combined = "".join(t.text or "" for t in t_elems)
        residue = DocxTemplateService._PLACEHOLDER_TOKEN_RE.sub("", combined)
        return not any(ch.isalnum() for ch in residue)

    @staticmethod
    def _render_lines_into_placeholder(paragraph, placeholder: str, lines: list[str]) -> None:
        """Replace `placeholder` in the paragraph with `lines` joined by soft line breaks.

        Collapses all `<w:t>` in the paragraph into the first run's text element
        (losing per-run formatting inside the rewritten span — acceptable for the
        short caption-line use case), then appends `<w:br w:type="line"/>` +
        `<w:t>line_n</w:t>` siblings inside the run's `<w:r>` element for each
        subsequent line. Trailing text after the placeholder is preserved.
        """
        p_element = paragraph._element
        t_elems = list(p_element.iter(qn("w:t")))
        if not t_elems:
            return
        combined = "".join(t.text or "" for t in t_elems)
        if placeholder not in combined:
            return

        before, _, after = combined.partition(placeholder)

        first_t = t_elems[0]
        run = first_t.getparent()
        for t in t_elems:
            t.text = ""

        first_t.text = before + lines[0]
        first_t.set(qn("xml:space"), "preserve")

        for line in lines[1:]:
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "line")
            run.append(br)
            new_t = OxmlElement("w:t")
            new_t.text = line
            new_t.set(qn("xml:space"), "preserve")
            run.append(new_t)

        if after:
            tail_t = OxmlElement("w:t")
            tail_t.text = after
            tail_t.set(qn("xml:space"), "preserve")
            run.append(tail_t)

    @staticmethod
    def _replace_in_document(doc: Document, marker: str, replacement: str) -> None:
        """Replace marker in all paragraphs and tables."""
        for paragraph in doc.paragraphs:
            DocxTemplateService._replace_in_paragraph(paragraph, marker, replacement)

        # For `\n`-bearing markers, also try spanning HARD paragraph breaks.
        # The per-paragraph + stage-3 soft-break passes above only span
        # `<w:br/>` within ONE paragraph. Word-authored recipient blocks
        # typically separate Name from Email via a hard return (separate
        # <w:p> elements), which the above passes can't bridge. This
        # bridges that gap so a `Name\nEmail` marker matches whether the
        # source used a soft or hard line break.
        if "\n" in marker:
            DocxTemplateService._replace_across_paragraphs(
                doc.element.body, marker, replacement,
            )

        for table in doc.tables:
            DocxTemplateService._replace_in_table(table, marker, replacement)

    @staticmethod
    def _replace_across_paragraphs(body_element, marker: str, replacement: str) -> bool:
        """Match a `\\n`-bearing `marker` against the body's direct-child <w:p>
        elements concatenated by `\\n`. On match spanning 2+ paragraphs,
        write `replacement` (plus any prefix from the first spanned
        paragraph) into the first spanned paragraph, remove intermediate
        spanned paragraphs entirely, and keep any suffix in the last
        spanned paragraph (or remove the last paragraph too if the
        suffix is empty).

        Returns True when a multi-paragraph replacement happened.

        Single-paragraph matches return False — the per-paragraph stage-1/2/3
        passes are responsible for those and have already run by the time
        this method is called.

        Skipped: typography-normalized fallback. The existing
        `_replace_span_with_breaks` skips it for the same reason — index
        drift from length-changing folds (e.g. `…` → `...`) would
        misalign the per-paragraph slicing below.
        """
        direct = body_element.findall(qn('w:p'))
        if len(direct) < 2:
            return False

        texts: list[str] = []
        for p in direct:
            t_elems = list(p.iter(qn('w:t')))
            texts.append("".join(t.text or "" for t in t_elems))

        # Stage 1: try the raw <w:t>-joined haystack first. Markers that
        # were authored against the raw stored form (including any trailing
        # whitespace Word kept inside <w:t>) match here.
        combined = "\n".join(texts)
        idx = combined.find(marker)

        if idx != -1:
            end = idx + len(marker)
            offsets: list[int] = []
            pos = 0
            for text in texts:
                offsets.append(pos)
                pos += len(text) + 1  # +1 for the `\n` separator (no trailing).
            end_offsets = [offsets[i] + len(texts[i]) for i in range(len(texts))]
            spanned: list[int] = []
            for i in range(len(direct)):
                if end_offsets[i] <= idx:
                    continue
                if offsets[i] >= end:
                    break
                spanned.append(i)
            if len(spanned) < 2:
                # Single-paragraph match — the per-paragraph pass already handled it.
                return False
            first_i = spanned[0]
            last_i = spanned[-1]
            prefix = texts[first_i][: max(0, idx - offsets[first_i])]
            suffix = texts[last_i][max(0, end - offsets[last_i]):]
        else:
            # Stage 2: fallback — per-paragraph rstrip to handle the common
            # case where Word stored trailing whitespace inside <w:t> but the
            # marker reflects parse_document_v2's `paragraph.text.strip()`
            # view (which trimmed that trailing whitespace). Without this,
            # joint-debtor captions where the first line happens to end with
            # a stored space silently fail to match.
            stripped_texts = [t.rstrip() for t in texts]
            stripped_combined = "\n".join(stripped_texts)
            stripped_idx = stripped_combined.find(marker)
            if stripped_idx == -1:
                return False
            stripped_end = stripped_idx + len(marker)
            stripped_offsets: list[int] = []
            pos = 0
            for st in stripped_texts:
                stripped_offsets.append(pos)
                pos += len(st) + 1
            stripped_end_offsets = [
                stripped_offsets[i] + len(stripped_texts[i])
                for i in range(len(stripped_texts))
            ]
            spanned = []
            for i in range(len(direct)):
                if stripped_end_offsets[i] <= stripped_idx:
                    continue
                if stripped_offsets[i] >= stripped_end:
                    break
                spanned.append(i)
            if len(spanned) < 2:
                return False
            first_i = spanned[0]
            last_i = spanned[-1]
            # rstrip only trims trailing whitespace, so [0:prefix_len] is
            # identical in stripped and raw forms.
            prefix_len = max(0, stripped_idx - stripped_offsets[first_i])
            prefix = texts[first_i][:prefix_len]
            # Suffix = chars in stripped last-paragraph AFTER the marker,
            # PLUS the whitespace rstrip removed from raw last-paragraph
            # (preserves the user's original trailing whitespace).
            suffix_start = max(0, stripped_end - stripped_offsets[last_i])
            stripped_suffix = stripped_texts[last_i][suffix_start:]
            rstripped_trailing = texts[last_i][len(stripped_texts[last_i]):]
            suffix = stripped_suffix + rstripped_trailing

        # First paragraph: collapse all <w:t> into the first run; write
        # prefix + replacement. Losing per-run formatting inside the
        # replaced span is acceptable (same trade-off as stage 2).
        first_p = direct[first_i]
        first_t_elems = list(first_p.iter(qn('w:t')))
        if first_t_elems:
            first_t_elems[0].text = prefix + replacement
            first_t_elems[0].set(qn("xml:space"), "preserve")
            for t in first_t_elems[1:]:
                t.text = ""

        # Last paragraph: keep suffix; remove the paragraph entirely if
        # suffix is empty (avoids leaving a visible blank line in the
        # rendered docx).
        last_p = direct[last_i]
        if last_i != first_i:
            last_t_elems = list(last_p.iter(qn('w:t')))
            if last_t_elems:
                last_t_elems[0].text = suffix
                last_t_elems[0].set(qn("xml:space"), "preserve")
                for t in last_t_elems[1:]:
                    t.text = ""
            if not suffix:
                parent = last_p.getparent()
                if parent is not None:
                    parent.remove(last_p)

        # Intermediate spanned paragraphs: remove entirely.
        for mid_i in spanned[1:-1]:
            mid_p = direct[mid_i]
            parent = mid_p.getparent()
            if parent is not None:
                parent.remove(mid_p)

        return True

    @staticmethod
    def _replace_in_table(table, marker: str, replacement: str) -> None:
        """Replace marker in all table cell paragraphs."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    DocxTemplateService._replace_in_paragraph(paragraph, marker, replacement)

    @staticmethod
    def _replace_in_paragraph(paragraph, marker: str, replacement: str) -> None:
        """Replace marker inside a paragraph, with typography tolerance.

        Stage 1 — per-run replacement (preserves fine-grained formatting):
        scan each `<w:t>` run; if the marker fits inside a single run, replace
        there. Word-boundary matching is used when the marker is bounded by
        alphanumerics on both ends (e.g. "13", "SMG") so short markers don't
        match substrings of longer tokens ("130", "1301"). If the direct
        match fails in a run that otherwise contains the marker with curly
        quotes / dashes, fall back to typography-normalized matching.

        Stage 2 — paragraph-level merge (handles markers that span runs):
        DOCX splits long paragraphs into multiple `<w:t>` runs at formatting
        boundaries, so markers covering many words often straddle runs and
        fail stage 1. Concatenate all runs, try direct match, then try a
        typography-normalized match (Word auto-corrects `'` → `’`, `-` → `—`,
        etc., while LLM-extracted markers use ASCII). If either match hits,
        collapse the replacement into the first run and clear the rest —
        losing per-run formatting inside the replaced span, which is
        acceptable for long narrative markers that become a single placeholder.
        """
        p_element = paragraph._element
        if not marker:
            return

        def _compile(m: str):
            use_boundary = m[0].isalnum() and m[-1].isalnum()
            return re.compile(r'\b' + re.escape(m) + r'\b') if use_boundary else None

        pattern = _compile(marker)
        normalized_marker = DocxTemplateService._normalize_typography(marker)
        pattern_norm = _compile(normalized_marker)

        def _apply_direct(text: str) -> str:
            if pattern is not None:
                return pattern.sub(replacement, text)
            return text.replace(marker, replacement)

        def _apply_normalized(text: str) -> str:
            """Normalize text to ASCII typography, then match on normalized form.
            The result is normalized too — surrounding curly quotes / dashes
            within the modified scope are folded to ASCII. That's acceptable
            because this path only fires when direct match failed, i.e. the
            DOCX has typography that the author never asked us to preserve."""
            norm = DocxTemplateService._normalize_typography(text)
            if pattern_norm is not None:
                return pattern_norm.sub(replacement, norm)
            return norm.replace(normalized_marker, replacement)

        # Stage 1: per-run direct, then normalized fallback per-run.
        for t_elem in p_element.iter(qn('w:t')):
            if not t_elem.text:
                continue
            if marker in t_elem.text:
                t_elem.text = _apply_direct(t_elem.text)
            elif normalized_marker in DocxTemplateService._normalize_typography(t_elem.text):
                t_elem.text = _apply_normalized(t_elem.text)

        # Stage 2: paragraph-level merge (for markers straddling multiple runs).
        t_elems = list(p_element.iter(qn('w:t')))
        if not t_elems:
            return
        combined = "".join(t.text or "" for t in t_elems)
        idx = combined.find(marker) if pattern is None else None
        if pattern is not None:
            m = pattern.search(combined)
            idx = m.start() if m else -1
            match_len = (m.end() - m.start()) if m else 0
        else:
            match_len = len(marker)
        if idx is not None and idx != -1:
            # SURGICAL per-run replacement: only modify <w:t>s overlapping
            # the marker span; leave prefix / suffix runs (and any
            # non-<w:t> children of intermediate runs like <w:tab/> or
            # <w:br/>) untouched. The blunt "dump everything into the
            # first <w:t>" path would move tabs to the wrong position
            # and lose per-run formatting outside the marker.
            DocxTemplateService._apply_substitution_to_t_elems(
                t_elems, idx, match_len, replacement,
            )
            return
        # Direct match failed; try typography-normalized fallback.
        # The fallback is intentionally blunt (collapses everything into
        # the first <w:t>) because length-changing folds (`…` → `...`)
        # would drift per-run offsets. Acceptable trade-off because this
        # path only fires when the document's typography differs from
        # the marker's — non-<w:t> elements in that scope are rare in
        # practice.
        norm_combined = DocxTemplateService._normalize_typography(combined)
        norm_result = _apply_normalized(combined)
        if norm_result != norm_combined:
            t_elems[0].text = norm_result
            for t in t_elems[1:]:
                t.text = ""
            return

        # Stage 3: marker spans <w:br/> (soft line break). Used for joint-debtor
        # captions where both names live in ONE paragraph joined by a line break.
        # Stages 1 and 2 miss these because they iterate <w:t> only — the \n
        # that python-docx surfaces in paragraph.text comes from <w:br/>, which
        # never appears in the <w:t> concatenation. This stage only fires when
        # the marker contains \n, keeping non-joint cases on the fast path.
        if "\n" not in marker:
            return
        DocxTemplateService._replace_span_with_breaks(p_element, marker, replacement)

    @staticmethod
    def _apply_substitution_to_t_elems(
        t_elems: list,
        idx: int,
        match_len: int,
        replacement: str,
    ) -> None:
        """Surgically substitute `replacement` for the `match_len`-char
        span starting at offset `idx` in the concatenated text of
        `t_elems`, modifying ONLY the <w:t> elements that overlap the
        span.

        Why: the blunt approach "concat → substitute → dump back into
        t_elems[0]" would move any non-<w:t> elements (tabs, breaks,
        formatting boundaries) to the end of the run sequence because
        the replacement text now occupies the first <w:t>'s slot
        instead of the original split positions. Surgical replacement
        preserves the document's original run layout outside the
        marker's span.

        Behavior:
        - <w:t>s entirely BEFORE the span → untouched (their <w:r>
          siblings, like <w:tab/>, also untouched).
        - <w:t>s entirely AFTER the span → untouched.
        - <w:t>s overlapping the span:
            - First overlapping: keeps any prefix chars before the
              match + the replacement text.
            - Last overlapping (when distinct from first): keeps any
              suffix chars after the match.
            - Intermediate overlapping: emptied.
            - Single overlapping (first == last): keeps prefix +
              replacement + suffix.
        """
        end = idx + match_len
        pos = 0
        first_overlap = None
        last_overlap = None
        intermediates: list = []
        for elem in t_elems:
            text = elem.text or ""
            start = pos
            stop = pos + len(text)
            pos = stop
            if stop <= idx or start >= end:
                continue
            if first_overlap is None:
                first_overlap = (elem, start, stop)
                last_overlap = (elem, start, stop)
            else:
                # The previous last_overlap becomes an intermediate
                # ONLY if it's not the first_overlap itself (which we
                # keep as-is to host the prefix + replacement text).
                if last_overlap is not None and last_overlap[0] is not first_overlap[0]:
                    intermediates.append(last_overlap[0])
                last_overlap = (elem, start, stop)

        if first_overlap is None:
            return  # marker not in any <w:t>; nothing to do

        first_elem, first_start, _ = first_overlap
        last_elem, last_start, _ = last_overlap
        first_text = first_elem.text or ""
        last_text = last_elem.text or ""

        prefix = first_text[: max(0, idx - first_start)]
        suffix = (
            last_text[max(0, end - last_start):]
            if last_elem is not first_elem
            else first_text[max(0, end - first_start):]
        )

        if first_elem is last_elem:
            first_elem.text = prefix + replacement + suffix
        else:
            first_elem.text = prefix + replacement
            last_elem.text = suffix

        # Preserve xml:space=preserve on modified runs (in case prefix
        # or suffix carries leading/trailing whitespace).
        if first_elem.text and (
            first_elem.text.startswith(" ") or first_elem.text.endswith(" ")
        ):
            first_elem.set(qn("xml:space"), "preserve")
        if last_elem is not first_elem and last_elem.text and (
            last_elem.text.startswith(" ") or last_elem.text.endswith(" ")
        ):
            last_elem.set(qn("xml:space"), "preserve")

        for elem in intermediates:
            elem.text = ""

    @staticmethod
    def _replace_span_with_breaks(p_element, marker: str, replacement: str) -> bool:
        """Match `marker` against the paragraph's text including `<w:br/>` line breaks.

        Walks `<w:t>` and `<w:br/>` children in document order, treating each
        `<w:br/>` as `\\n` in the combined text. On match, writes the
        replacement into the first `<w:t>` that overlaps the span (preserving
        surrounding prefix/suffix outside the span), empties any other
        overlapping `<w:t>` elements, and removes `<w:br/>` elements inside
        the span. Typography-normalized fallback is skipped here to avoid
        index drift from length-changing folds (e.g. ellipsis → "...").
        Returns True when a replacement was made.
        """
        items: list[tuple[str, object, str, int, int]] = []
        pos = 0
        for elem in p_element.iter(qn("w:t"), qn("w:br")):
            if elem.tag == qn("w:t"):
                text = elem.text or ""
                items.append(("t", elem, text, pos, pos + len(text)))
                pos += len(text)
            else:
                items.append(("br", elem, "\n", pos, pos + 1))
                pos += 1

        if not items:
            return False

        combined = "".join(text for _, _, text, _, _ in items)
        idx = combined.find(marker)
        if idx == -1:
            return False
        end = idx + len(marker)

        first_t_item = None
        last_t_item = None
        br_in_span: list[object] = []
        t_in_span: list[tuple[object, int, int]] = []
        for kind, elem, _, start, stop in items:
            if stop <= idx or start >= end:
                continue
            if kind == "br":
                br_in_span.append(elem)
                continue
            t_in_span.append((elem, start, stop))
            if first_t_item is None:
                first_t_item = (elem, start, stop)
            last_t_item = (elem, start, stop)

        if first_t_item is None:
            return False  # span covers only <w:br/>; no text host for replacement

        first_elem, first_start, _ = first_t_item
        last_elem, last_start, _ = last_t_item
        first_text = first_elem.text or ""
        last_text = last_elem.text or ""

        prefix = first_text[: max(0, idx - first_start)]
        suffix = last_text[max(0, end - last_start) :] if last_elem is not first_elem else first_text[max(0, end - first_start) :]

        if first_elem is last_elem:
            first_elem.text = prefix + replacement + suffix
        else:
            first_elem.text = prefix + replacement
            last_elem.text = suffix

        for t_elem, _, _ in t_in_span:
            if t_elem is first_elem or t_elem is last_elem:
                continue
            t_elem.text = ""

        for br_elem in br_in_span:
            parent = br_elem.getparent()
            if parent is not None:
                parent.remove(br_elem)

        return True
