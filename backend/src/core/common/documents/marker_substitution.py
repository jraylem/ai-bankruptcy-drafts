"""Two-pass marker substitution: scan → resolve conflicts → mutate.

The legacy `create_template` flow processes markers longest-first but
mutates immediately. When a longer marker fails to match (e.g. cross-
paragraph markers blocked by hyperlinks/field codes/whitespace drift),
its shorter substring sibling can grab the longer marker's territory
and orphan it.

Concrete example: spec has `cos_email_section_1` (marker
`"Michael R. Bakst\n<emails>"`) AND `trustee_name` (marker
`"Michael R. Bakst"`). If the longer marker fails its cross-paragraph
match, `trustee_name` then replaces `"Michael R. Bakst"` in BOTH the
body AND the CoS region. `cos_email_section_1`'s placeholder never
makes it into the docx and gets dropped as an orphan.

The two-pass approach:
  Pass 1  — scan the doc for EVERY (marker, position) candidate
            without mutating.
  Pass 1.5 — resolve conflicts: process longest-first, claim character
             ranges in a global 1-D address space; shorter candidates
             whose span overlaps a longer candidate's claim get
             dropped. Safety net: even when a longer marker produced
             no candidate of its own, if its text contains a shorter
             marker's text AND would match the doc at the implied
             offset, the shorter marker still defers.
  Pass 2  — mutate: each surviving candidate applies its replacement
            in reverse-offset order so earlier candidates in the same
            paragraph aren't shifted by later ones.

Entry point: `substitute_markers(doc, template_spec)`.
"""

import logging
from dataclasses import dataclass, field
from typing import Literal

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.core.agents.types.spec import TemplateVariable

from .docx_template import DocxTemplateService

logger = logging.getLogger(__name__)


# === Data types ===


SpanKind = Literal["plain_para", "soft_break", "cross_para"]


@dataclass(frozen=True)
class ParaLocation:
    """Stable identifier for a paragraph in the doc.

    body paragraph: kind="body", body_idx=N
    table cell paragraph: kind="cell", table_idx=T, row_idx=R, col_idx=C, para_idx=P
    """
    kind: Literal["body", "cell"]
    body_idx: int = -1
    table_idx: int = -1
    row_idx: int = -1
    col_idx: int = -1
    para_idx: int = -1


@dataclass
class ParaInfo:
    """Per-paragraph state cached for scanning + mutation."""
    location: ParaLocation
    para_element: object   # lxml <w:p>
    text: str              # joined <w:t> chars only (no break translation)
    text_with_breaks: str  # joined <w:t> + each <w:br/> → "\n"
    global_start: int      # offset in the flat doc text (text-only view)


@dataclass
class FlatIndex:
    """Cached per-doc state used by every pass."""
    paragraphs: list[ParaInfo]   # body first, then every table cell para
    body_paras: list[ParaInfo]   # subset, for cross-paragraph scan
    flat_text: str               # all paragraphs joined by "\n" (text-only)


@dataclass(frozen=True)
class MatchCandidate:
    """One textual occurrence of a marker in the doc."""
    marker: str
    replacement: str
    variable_index: int
    identifying_text: str | None
    global_start: int            # for collision overlap
    global_end: int
    kind: SpanKind
    # Mutation payload — interpretation depends on kind:
    #   plain_para  : spans = ((loc, idx_in_text, length_in_text),)
    #   soft_break  : spans = ((loc, idx_in_text_with_breaks,
    #                            length_in_text_with_breaks),)
    #   cross_para  : spans = ((loc_first, prefix_in_text), (loc_last,
    #                            suffix_in_text), (loc_intermediate,
    #                            -1), ...)
    spans: tuple


# === Entry point ===


def substitute_markers(doc, template_spec: list[TemplateVariable]) -> None:
    """Replace template_property_marker (and aliases) with
    template_variable_string across `doc`, using the collision-aware
    two-pass algorithm documented at the top of this module.

    Drop-in replacement for the legacy mutate-immediately pass in
    `DocxTemplateService.create_template`. Wraps `doc` in place; no
    return value (callers serialize the mutated `doc`).
    """
    if not template_spec:
        return
    flat_index = _build_flat_index(doc)
    if not flat_index.paragraphs:
        return
    candidates = _scan_all_candidates(flat_index, template_spec)
    survivors = _resolve_conflicts(candidates, template_spec, flat_index)
    _apply_survivors(flat_index, survivors)


# === Pass 0: build flat index ===


def _build_flat_index(doc) -> FlatIndex:
    paragraphs: list[ParaInfo] = []
    body_paras: list[ParaInfo] = []
    global_pos = 0

    for body_idx, para in enumerate(doc.paragraphs):
        info = _make_para_info(
            location=ParaLocation(kind="body", body_idx=body_idx),
            para_element=para._element,
            global_pos=global_pos,
        )
        paragraphs.append(info)
        body_paras.append(info)
        global_pos += len(info.text) + 1

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    info = _make_para_info(
                        location=ParaLocation(
                            kind="cell",
                            table_idx=table_idx,
                            row_idx=row_idx,
                            col_idx=col_idx,
                            para_idx=para_idx,
                        ),
                        para_element=para._element,
                        global_pos=global_pos,
                    )
                    paragraphs.append(info)
                    global_pos += len(info.text) + 1

    flat_text = "\n".join(p.text for p in paragraphs)
    return FlatIndex(paragraphs=paragraphs, body_paras=body_paras, flat_text=flat_text)


def _make_para_info(location: ParaLocation, para_element, global_pos: int) -> ParaInfo:
    text_parts: list[str] = []
    text_with_break_parts: list[str] = []
    for elem in para_element.iter(qn("w:t"), qn("w:br")):
        if elem.tag == qn("w:t"):
            t = elem.text or ""
            text_parts.append(t)
            text_with_break_parts.append(t)
        else:
            text_with_break_parts.append("\n")
    return ParaInfo(
        location=location,
        para_element=para_element,
        text="".join(text_parts),
        text_with_breaks="".join(text_with_break_parts),
        global_start=global_pos,
    )


# === Pass 1: scan ===


def _scan_all_candidates(
    flat_index: FlatIndex, spec: list[TemplateVariable],
) -> list[MatchCandidate]:
    candidates: list[MatchCandidate] = []
    for variable_index, variable in enumerate(spec):
        replacement = variable.template_variable_string
        if not replacement:
            continue
        primary = variable.template_property_marker
        aliases = list(variable.template_property_marker_aliases or [])
        all_markers = [m for m in [primary, *aliases] if m]
        for marker in all_markers:
            clean = _clean_marker(marker)
            if not clean:
                continue
            candidates.extend(
                _scan_one_marker(
                    flat_index, clean, replacement, variable, variable_index,
                )
            )
    return candidates


def _clean_marker(marker: str) -> str:
    """Normalize `\\t` → ` ` and literal `\\n` → real `\\n` so the agent's
    JSON-encoded markers align with the doc's stored text."""
    return marker.replace("\t", " ").replace("\\n", "\n")


def _scan_one_marker(
    flat_index: FlatIndex,
    marker: str,
    replacement: str,
    variable: TemplateVariable,
    variable_index: int,
) -> list[MatchCandidate]:
    results: list[MatchCandidate] = []
    has_newline = "\n" in marker
    identifying_text = variable.template_identifying_text_match

    for para in flat_index.paragraphs:
        if not has_newline:
            for idx in _find_all_with_typography_tolerance(para.text, marker):
                results.append(MatchCandidate(
                    marker=marker, replacement=replacement,
                    variable_index=variable_index,
                    identifying_text=identifying_text,
                    global_start=para.global_start + idx,
                    global_end=para.global_start + idx + len(marker),
                    kind="plain_para",
                    spans=((para.location, idx, len(marker)),),
                ))
            continue
        if "\n" not in para.text_with_breaks:
            continue
        for idx in _find_all_with_typography_tolerance(para.text_with_breaks, marker):
            text_idx, text_len = _project_break_span_to_text(
                para.text_with_breaks, idx, len(marker),
            )
            results.append(MatchCandidate(
                marker=marker, replacement=replacement,
                variable_index=variable_index,
                identifying_text=identifying_text,
                global_start=para.global_start + text_idx,
                global_end=para.global_start + text_idx + text_len,
                kind="soft_break",
                spans=((para.location, idx, len(marker)),),
            ))

    if has_newline and len(flat_index.body_paras) >= 2:
        cross_results = _scan_cross_para(
            flat_index, marker, replacement, identifying_text, variable_index,
        )
        results.extend(cross_results)
        if not cross_results and not any(
            r.kind == "soft_break" for r in results
        ):
            _log_cross_para_miss_diagnostic(flat_index, marker, variable.template_variable)

    return results


def _log_cross_para_miss_diagnostic(
    flat_index: FlatIndex, marker: str, variable_name: str,
) -> None:
    """Log a one-shot diagnostic when a `\\n`-bearing marker produced
    zero candidates (no single-paragraph soft-break match AND no
    cross-paragraph match). Surfaces which paragraph the marker's first
    line WOULD have matched + the first divergence between marker and
    doc — so paralegals + devs can see WHY a long marker silently
    failed instead of having to reverse-engineer it from XML.

    Logged at INFO so it shows up in the composer log stream without
    spamming WARNING-level for what is often a recoverable miss (the
    short substring marker may still claim the region elsewhere).
    """
    first_line = marker.split("\n", 1)[0]
    if not first_line:
        return
    for p in flat_index.body_paras:
        if first_line in p.text:
            # Compare the marker against the would-be cross-para join
            # starting from this paragraph. Find first character that
            # differs.
            body_idx = p.location.body_idx
            tail = "\n".join(b.text for b in flat_index.body_paras[body_idx:])
            divergence = _first_divergence(marker, tail)
            logger.info(
                "marker_substitution: cross-para marker for %r matched its "
                "first line at body para %d but the full marker diverged at "
                "char offset %d (marker[%d:%d]=%r vs doc[%d:%d]=%r); marker "
                "preview=%r; doc preview=%r",
                variable_name,
                body_idx,
                divergence,
                max(0, divergence - 20),
                min(len(marker), divergence + 20),
                marker[max(0, divergence - 20): min(len(marker), divergence + 20)],
                max(0, divergence - 20),
                min(len(tail), divergence + 20),
                tail[max(0, divergence - 20): min(len(tail), divergence + 20)],
                marker[: min(200, len(marker))],
                tail[: min(200, len(tail))],
            )
            return


def _first_divergence(a: str, b: str) -> int:
    for i in range(min(len(a), len(b))):
        if a[i] != b[i]:
            return i
    return min(len(a), len(b))


def _project_break_span_to_text(text_with_breaks: str, idx: int, length: int) -> tuple[int, int]:
    """Translate (idx, length) in `text_with_breaks` to (text_idx, text_len)
    in the same paragraph's `text` (which omits the `<w:br/>` newlines).

    Used to give soft-break candidates a global_start that aligns with
    every other candidate in the doc's text-only address space.
    """
    text_idx = idx - text_with_breaks[:idx].count("\n")
    text_len = length - text_with_breaks[idx : idx + length].count("\n")
    return text_idx, text_len


_LENGTH_PRESERVING_TYPOGRAPHY_MAP = str.maketrans({
    "‘": "'",   # LEFT SINGLE QUOTATION MARK
    "’": "'",   # RIGHT SINGLE QUOTATION MARK (Word-auto apostrophe)
    "“": '"',   # LEFT DOUBLE QUOTATION MARK
    "”": '"',   # RIGHT DOUBLE QUOTATION MARK
    "–": "-",   # EN DASH
    "—": "-",   # EM DASH
    " ": " ",   # NON-BREAKING SPACE
    # Deliberately EXCLUDES `…` (ellipsis → "...") because the 1→3
    # length change would drift downstream offsets. Cross-para scans
    # depending on absolute char positions need exact 1:1 mapping.
})


def _normalize_for_cross_para_scan(text: str) -> str:
    """Length-preserving typography fold for cross-para scanning.

    A subset of DocxTemplateService._normalize_typography that
    GUARANTEES `len(out) == len(text)` so per-paragraph offsets stay
    intact across the fold. Used to catch markers that differ from the
    doc only in curly-quote / dash / NBSP punctuation — common when
    Word's autocorrect mangles emails or contact blocks differently
    than the agent's marker extraction.
    """
    return text.translate(_LENGTH_PRESERVING_TYPOGRAPHY_MAP)


def _scan_cross_para(
    flat_index: FlatIndex,
    marker: str,
    replacement: str,
    identifying_text: str | None,
    variable_index: int,
) -> list[MatchCandidate]:
    body = flat_index.body_paras
    texts = [p.text for p in body]
    offsets: list[int] = []
    pos = 0
    for t in texts:
        offsets.append(pos)
        pos += len(t) + 1
    end_offsets = [offsets[i] + len(texts[i]) for i in range(len(texts))]

    haystacks: list[tuple[list[str], list[int], list[int], str]] = [
        (texts, offsets, end_offsets, marker),
    ]
    stripped_texts = [t.rstrip() for t in texts]
    if stripped_texts != texts:
        stripped_offsets: list[int] = []
        pos = 0
        for st in stripped_texts:
            stripped_offsets.append(pos)
            pos += len(st) + 1
        stripped_end_offsets = [
            stripped_offsets[i] + len(stripped_texts[i])
            for i in range(len(stripped_texts))
        ]
        haystacks.append((stripped_texts, stripped_offsets, stripped_end_offsets, marker))

    # Typography-tolerant fallback: fold curly quotes / dashes / NBSP
    # on BOTH sides so the marker still matches when Word stored
    # autocorrected punctuation that the agent's marker didn't carry.
    # Length-preserving fold keeps per-paragraph offsets intact.
    normalized_texts = [_normalize_for_cross_para_scan(t) for t in texts]
    normalized_marker = _normalize_for_cross_para_scan(marker)
    if normalized_texts != texts or normalized_marker != marker:
        haystacks.append((normalized_texts, offsets, end_offsets, normalized_marker))
    normalized_stripped = [_normalize_for_cross_para_scan(t) for t in stripped_texts]
    if stripped_texts != texts and (
        normalized_stripped != stripped_texts or normalized_marker != marker
    ):
        haystacks.append((
            normalized_stripped, stripped_offsets, stripped_end_offsets, normalized_marker,
        ))

    results: list[MatchCandidate] = []
    seen_spans: set[tuple[int, int]] = set()

    for view_texts, view_offsets, view_end_offsets, view_marker in haystacks:
        combined = "\n".join(view_texts)
        for idx in _find_all_in_string(combined, view_marker):
            end = idx + len(marker)
            spanned: list[int] = []
            for i in range(len(body)):
                if view_end_offsets[i] <= idx:
                    continue
                if view_offsets[i] >= end:
                    break
                spanned.append(i)
            if len(spanned) < 2:
                continue
            first_i, last_i = spanned[0], spanned[-1]
            prefix_len = max(0, idx - view_offsets[first_i])
            suffix_start = max(0, end - view_offsets[last_i])

            # Global range in TEXT (not text_with_breaks) coordinates.
            global_start = body[first_i].global_start + prefix_len
            text_last = body[last_i].text
            # When matching against stripped view, the suffix in the
            # mutation step needs the original trailing whitespace
            # appended; collision detection only cares about textual
            # range so we use the stripped length.
            global_end = body[last_i].global_start + suffix_start
            key = (global_start, global_end)
            if key in seen_spans:
                continue
            seen_spans.add(key)

            # Build mutation spans.
            # First entry: (first_loc, prefix_len_in_text).
            # Last entry:  (last_loc, suffix_start_in_text).
            # Intermediate: (loc, -1).
            spans = [(body[first_i].location, prefix_len)]
            for mid_i in spanned[1:-1]:
                spans.append((body[mid_i].location, -1))
            spans.append((body[last_i].location, suffix_start))

            results.append(MatchCandidate(
                marker=marker, replacement=replacement,
                variable_index=variable_index,
                identifying_text=identifying_text,
                global_start=global_start,
                global_end=global_end,
                kind="cross_para",
                spans=tuple(spans),
            ))
    return results


def _find_all_in_string(hay: str, needle: str) -> list[int]:
    if not needle:
        return []
    results: list[int] = []
    start = 0
    while True:
        idx = hay.find(needle, start)
        if idx == -1:
            return results
        results.append(idx)
        start = idx + 1


def _find_all_word_boundary(hay: str, needle: str) -> list[int]:
    """Find every occurrence of `needle` in `hay`, honoring word-boundary
    rules when both ends of the needle are alphanumeric.

    Mirrors the matching contract of `DocxTemplateService._replace_in_paragraph`
    stage 1 (`\\b...\\b` boundaries when alphanumeric, plain `find`
    otherwise) but returns ALL positions, not just first.
    """
    if not needle:
        return []
    use_boundary = needle[0].isalnum() and needle[-1].isalnum()
    if not use_boundary:
        return _find_all_in_string(hay, needle)

    import re
    pattern = re.compile(r"\b" + re.escape(needle) + r"\b")
    return [m.start() for m in pattern.finditer(hay)]


def _find_all_with_typography_tolerance(hay: str, needle: str) -> list[int]:
    """Find every occurrence of `needle` in `hay` with length-preserving
    typography tolerance.

    Why this exists: Word autocorrects regular spaces to NBSP (\\u00A0)
    inside contact/email blocks, swaps straight quotes for curly ones,
    swaps `-` for en/em dashes, etc. The agent's markers carry the
    ASCII forms because `parse_document_v2` reads `paragraph.text`
    (which gives the stored char). When stored char ≠ marker char,
    exact `find` silently misses — which is the Fleisher CoS bug
    (NBSP between `lippes.com,` and `ecf.alert+`).

    Strategy: try exact match first. If zero hits AND the needle has
    word-boundary-friendly ends, try again on the typography-normalized
    haystack with the normalized needle. The fold is length-preserving
    so offsets returned are valid for the ORIGINAL `hay`.

    Returns deduplicated, sorted positions.
    """
    direct = _find_all_word_boundary(hay, needle)
    if direct:
        return direct
    normalized_hay = _normalize_for_cross_para_scan(hay)
    normalized_needle = _normalize_for_cross_para_scan(needle)
    if normalized_hay == hay and normalized_needle == needle:
        return []
    return _find_all_word_boundary(normalized_hay, normalized_needle)


# === Pass 1.5: resolve conflicts ===


def _resolve_conflicts(
    candidates: list[MatchCandidate],
    spec: list[TemplateVariable],
    flat_index: FlatIndex,
) -> list[MatchCandidate]:
    """Longest-first claim of [global_start, global_end) intervals.

    Two safety nets:
      1. A shorter candidate whose span overlaps an already-claimed
         interval is dropped — the longer marker has rightful claim.
      2. Even if no longer candidate exists for a span, if some longer
         marker in the spec *contains* this marker as a substring AND
         its text would match the doc at the implied offset, the
         shorter candidate is dropped anyway (the user's "would-be
         claimant" rule — handles the case where the longer marker
         failed its own match for an unrelated reason).

    Shared-marker tiebreak: when multiple variables share the same
    marker text AND that marker has multiple matching positions, the
    shared-marker logic is delegated to `_assign_shared_marker_positions`
    after the longest-first claim so each variable's identifying-text
    routing is honored.
    """
    if not candidates:
        return []

    by_marker_text: dict[str, list[MatchCandidate]] = {}
    for c in candidates:
        by_marker_text.setdefault(c.marker, []).append(c)

    shared_marker_pool: list[MatchCandidate] = []
    unique_pool: list[MatchCandidate] = []
    for marker, cs in by_marker_text.items():
        distinct_replacements = {c.replacement for c in cs}
        if len(distinct_replacements) > 1:
            shared_marker_pool.extend(cs)
        else:
            unique_pool.extend(cs)

    # Process longest-first across BOTH pools combined so a shared
    # marker (short) still defers to a unique longer marker.
    ordered = sorted(
        unique_pool + shared_marker_pool,
        key=lambda c: (-len(c.marker), c.variable_index, c.global_start),
    )
    claimed: list[tuple[int, int]] = []
    survivors: list[MatchCandidate] = []
    deferred_shared: dict[str, list[MatchCandidate]] = {}

    spec_markers = _build_spec_marker_set(spec)

    for c in ordered:
        if _overlaps_any(claimed, c.global_start, c.global_end):
            continue
        if _is_contained_in_unclaimed_longer(c, spec_markers, flat_index.flat_text):
            continue

        if c.marker in {m for m, cs in by_marker_text.items()
                        if len({x.replacement for x in cs}) > 1}:
            deferred_shared.setdefault(c.marker, []).append(c)
            claimed.append((c.global_start, c.global_end))
            continue

        survivors.append(c)
        claimed.append((c.global_start, c.global_end))

    survivors.extend(
        _assign_shared_marker_positions(deferred_shared, by_marker_text, flat_index),
    )
    return survivors


def _overlaps_any(claimed: list[tuple[int, int]], start: int, end: int) -> bool:
    for a, b in claimed:
        if start < b and a < end:
            return True
    return False


def _build_spec_marker_set(spec: list[TemplateVariable]) -> list[str]:
    out: list[str] = []
    for v in spec:
        if not v.template_variable_string:
            continue
        if v.template_property_marker:
            out.append(_clean_marker(v.template_property_marker))
        for alias in v.template_property_marker_aliases or []:
            if alias:
                out.append(_clean_marker(alias))
    return out


def _is_contained_in_unclaimed_longer(
    short: MatchCandidate, all_markers: list[str], flat_text: str,
) -> bool:
    """User's would-be-claimant rule.

    For each LONGER marker N in spec whose text contains `short.marker`
    as a substring at offset k: check whether `flat_text` at
    `(short.global_start - k)` equals N. If yes, N is the rightful
    claimant of this region — even if N produced no candidate of its
    own — and `short` must defer.

    Plays the role of a typography-tolerant containment check by
    delegating to `DocxTemplateService._normalize_typography` for the
    equality test so curly-quote / dash variants don't slip past.
    """
    short_text = short.marker
    flat_len = len(flat_text)
    for n in all_markers:
        if len(n) <= len(short_text):
            continue
        offsets = _find_all_substring_offsets(n, short_text)
        for k in offsets:
            implied_start = short.global_start - k
            if implied_start < 0:
                continue
            implied_end = implied_start + len(n)
            if implied_end > flat_len:
                continue
            window = flat_text[implied_start:implied_end]
            if window == n:
                return True
            if DocxTemplateService._normalize_typography(window) == \
                    DocxTemplateService._normalize_typography(n):
                return True
    return False


def _find_all_substring_offsets(hay: str, needle: str) -> list[int]:
    if not needle:
        return []
    results: list[int] = []
    start = 0
    while True:
        idx = hay.find(needle, start)
        if idx == -1:
            return results
        results.append(idx)
        start = idx + 1


def _assign_shared_marker_positions(
    deferred: dict[str, list[MatchCandidate]],
    by_marker_text: dict[str, list[MatchCandidate]],
    flat_index: FlatIndex,
) -> list[MatchCandidate]:
    """For each shared-marker group, distribute survivors among the
    variables that share that marker using `identifying_text` routing.

    Mirrors `_replace_first_in_context`'s disambiguation logic 1:1,
    including the WARNING-level log lines on fallback / no-occurrence
    so spec authors get the same debugging signal they're used to.
    """
    loc_to_text: dict[ParaLocation, str] = {
        p.location: p.text for p in flat_index.paragraphs
    }

    def _para_text_of(position: MatchCandidate) -> str:
        first_loc = position.spans[0][0]
        return loc_to_text.get(first_loc, "")

    def _normalize(s: str) -> str:
        return " ".join(s.lower().split())

    out: list[MatchCandidate] = []
    for marker, position_candidates in deferred.items():
        if not position_candidates:
            continue
        all_for_marker = by_marker_text[marker]
        variables: list[MatchCandidate] = []
        seen_vars: set[int] = set()
        for c in all_for_marker:
            if c.variable_index in seen_vars:
                continue
            seen_vars.add(c.variable_index)
            variables.append(c)
        unused_positions = list(position_candidates)

        for var_template in variables:
            if not unused_positions:
                logger.warning(
                    "Shared-marker placement: marker %r has no unconsumed "
                    "occurrence in the document; placeholder %r will be "
                    "orphaned and dropped by _drop_orphan_variables.",
                    marker, var_template.replacement,
                )
                continue

            picked: MatchCandidate | None = None
            if var_template.identifying_text and var_template.identifying_text.strip():
                needle = _normalize(var_template.identifying_text)
                for pos in unused_positions:
                    hay = _normalize(_para_text_of(pos))
                    if needle in hay or hay in needle:
                        picked = pos
                        break

            if picked is None:
                picked = unused_positions[0]
                if var_template.identifying_text:
                    logger.warning(
                        "Shared-marker disambiguation fallback: "
                        "identifying_text_match %r did not match any "
                        "unconsumed paragraph for marker %r; using first "
                        "unconsumed occurrence.",
                        var_template.identifying_text, marker,
                    )

            unused_positions.remove(picked)
            out.append(MatchCandidate(
                marker=picked.marker,
                replacement=var_template.replacement,
                variable_index=var_template.variable_index,
                identifying_text=var_template.identifying_text,
                global_start=picked.global_start,
                global_end=picked.global_end,
                kind=picked.kind,
                spans=picked.spans,
            ))
    return out


# === Pass 2: mutate ===


def _apply_survivors(flat_index: FlatIndex, survivors: list[MatchCandidate]) -> None:
    """Apply each survivor's replacement.

    Cross-paragraph candidates run FIRST (they may delete entire
    paragraphs, invalidating later indices). Within a paragraph, plain
    + soft-break candidates run in REVERSE-idx order so earlier-offset
    survivors aren't shifted by later mutations.
    """
    cross_para = [s for s in survivors if s.kind == "cross_para"]
    in_para = [s for s in survivors if s.kind != "cross_para"]

    in_para_by_loc: dict[ParaLocation, list[MatchCandidate]] = {}
    for c in in_para:
        loc = c.spans[0][0]
        in_para_by_loc.setdefault(loc, []).append(c)

    loc_to_para_element: dict[ParaLocation, object] = {
        p.location: p.para_element for p in flat_index.paragraphs
    }

    for cs in in_para_by_loc.values():
        cs.sort(key=lambda c: c.spans[0][1], reverse=True)
        for c in cs:
            para_element = loc_to_para_element.get(c.spans[0][0])
            if para_element is None:
                continue
            if c.kind == "plain_para":
                _apply_plain_para(para_element, c)
            else:
                _apply_soft_break(para_element, c)

    for c in cross_para:
        _apply_cross_para(flat_index, c)


def _apply_plain_para(p_element, c: MatchCandidate) -> None:
    _loc, idx, length = c.spans[0]
    t_elems = list(p_element.iter(qn("w:t")))
    if not t_elems:
        return
    DocxTemplateService._apply_substitution_to_t_elems(
        t_elems, idx, length, c.replacement,
    )


def _apply_soft_break(p_element, c: MatchCandidate) -> None:
    """Replace a marker span that contains <w:br/> soft breaks inside
    one paragraph. Mirrors `_replace_span_with_breaks` but uses the
    pre-scanned `(idx_in_text_with_breaks, length_in_text_with_breaks)`
    instead of re-searching.
    """
    _loc, idx, length = c.spans[0]
    items: list[tuple[str, object, str, int, int]] = []
    pos = 0
    for elem in p_element.iter(qn("w:t"), qn("w:br")):
        if elem.tag == qn("w:t"):
            text = elem.text or ""
            items.append(("t", elem, text, pos, pos + len(text)))
            pos += len(text)
        else:
            items.append(("br", elem, "\n", pos, pos + 1))
            pos += 1

    end = idx + length
    first_t = None
    last_t = None
    br_in_span: list[object] = []
    t_in_span: list[object] = []
    for kind, elem, _, start, stop in items:
        if stop <= idx or start >= end:
            continue
        if kind == "br":
            br_in_span.append(elem)
            continue
        t_in_span.append(elem)
        if first_t is None:
            first_t = (elem, start)
        last_t = (elem, start)

    if first_t is None:
        return

    first_elem, first_start = first_t
    last_elem, last_start = last_t
    first_text = first_elem.text or ""
    last_text = last_elem.text or ""

    prefix = first_text[: max(0, idx - first_start)]
    suffix = (
        last_text[max(0, end - last_start):]
        if last_elem is not first_elem
        else first_text[max(0, end - first_start):]
    )

    if first_elem is last_elem:
        first_elem.text = prefix + c.replacement + suffix
    else:
        first_elem.text = prefix + c.replacement
        last_elem.text = suffix

    for t_elem in t_in_span:
        if t_elem is first_elem or t_elem is last_elem:
            continue
        t_elem.text = ""

    for br_elem in br_in_span:
        parent = br_elem.getparent()
        if parent is not None:
            parent.remove(br_elem)


def _apply_cross_para(flat_index: FlatIndex, c: MatchCandidate) -> None:
    """Apply a cross-paragraph replacement.

    spans = (
        (first_loc, prefix_len_in_text),
        (mid_loc, -1),     # zero or more intermediates to delete entirely
        ...
        (last_loc, suffix_start_in_text),
    )

    Behavior mirrors `_replace_across_paragraphs`'s mutation tail:
      - Write `prefix + replacement` into the first paragraph's first
        <w:t>; clear its other <w:t>s.
      - Remove every intermediate paragraph entirely.
      - Write `suffix` (or stripped + trailing whitespace) into the
        last paragraph's first <w:t>; remove the paragraph entirely
        when suffix is empty.
    """
    if len(c.spans) < 2:
        return

    loc_to_para = {p.location: p for p in flat_index.paragraphs}
    first_loc, prefix_len = c.spans[0]
    last_loc, suffix_start = c.spans[-1]
    intermediate_locs = [loc for loc, _ in c.spans[1:-1]]

    first_para = loc_to_para.get(first_loc)
    last_para = loc_to_para.get(last_loc)
    if first_para is None or last_para is None:
        return

    prefix = first_para.text[:prefix_len]
    suffix = last_para.text[suffix_start:]

    first_t_elems = list(first_para.para_element.iter(qn("w:t")))
    if first_t_elems:
        first_t_elems[0].text = prefix + c.replacement
        first_t_elems[0].set(qn("xml:space"), "preserve")
        for t in first_t_elems[1:]:
            t.text = ""

    if last_loc != first_loc:
        last_t_elems = list(last_para.para_element.iter(qn("w:t")))
        if last_t_elems:
            last_t_elems[0].text = suffix
            last_t_elems[0].set(qn("xml:space"), "preserve")
            for t in last_t_elems[1:]:
                t.text = ""
        if not suffix:
            parent = last_para.para_element.getparent()
            if parent is not None:
                parent.remove(last_para.para_element)

    for mid_loc in intermediate_locs:
        mid_para = loc_to_para.get(mid_loc)
        if mid_para is None:
            continue
        parent = mid_para.para_element.getparent()
        if parent is not None:
            parent.remove(mid_para.para_element)
