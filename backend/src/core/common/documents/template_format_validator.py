"""Template format validator — detects formatting drift introduced by
the substitution engine.

The substitution engine is SUPPOSED to preserve every char outside the
marker / placeholder spans (tabs, line breaks, paragraph splits,
leading whitespace, typography). Bugs in the substitution engine can
drop those — visible as collapsed lines, lost tabs, content from
adjacent paragraphs appearing inline.

This module catches drift at TWO points in the template lifecycle:

  1. **Composer-time** — `validate_template_format(original, template, spec)`
     runs after `create_template` produces the placeholder-marked
     template.docx, before R2 upload. Reconstructs by substituting
     placeholders back to known markers; expects the result to
     appear in the original.

  2. **Fill-time** — `validate_fill_format(template, filled, resolved_values)`
     runs after `fill_template` produces the dry-run / draft docx.
     Reconstructs the expected filled paragraph by substituting
     placeholders forward to resolved values; expects per-paragraph
     equality with the actual filled output.

Both validators are Tier 1 — deterministic, no LLM. Typography-tolerant
via `marker_substitution._normalize_for_cross_para_scan` (length-
preserving NBSP / curly quote / dash fold).

Used behind `TEMPLATE_FORMAT_AUTOFIX_V2` env flag. Tier 1 only LOGS
drift; the LLM-driven Tier 2 fixer will be added in a follow-up PR.
"""

import logging
import re
from dataclasses import dataclass
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from src.core.agents.types.spec import TemplateVariable

from .docx_template import DocxTemplateService
from .marker_substitution import _normalize_for_cross_para_scan


def _paragraph_text_with_breaks(paragraph) -> str:
    """Read a paragraph's text including `<w:br/>` soft breaks as
    `\\n` chars AND `<w:tab/>` elements as `\\t` chars.

    python-docx's `.text` property's break / tab handling varies by
    version; the validator + fixer need guaranteed break-aware AND
    tab-aware reading to compare against the symbolic forms that
    `_apply_corrected_text` emits.

    Mirrors the helper used inside `marker_substitution._make_para_info`.
    """
    parts: list[str] = []
    for elem in paragraph._element.iter(qn("w:t"), qn("w:br"), qn("w:tab")):
        if elem.tag == qn("w:t"):
            parts.append(elem.text or "")
        elif elem.tag == qn("w:br"):
            parts.append("\n")
        else:
            parts.append("\t")
    return "".join(parts)

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class DriftedParagraph:
    """One template paragraph whose reconstructed text could not be
    located in the original doc — likely formatting drift caused by
    the substitution engine."""
    template_paragraph_index: int
    template_text: str
    reconstructed_text: str
    drift_reason: str


@dataclass(frozen=True)
class FormatValidationResult:
    ok: bool
    drifted_paragraphs: list[DriftedParagraph]
    summary: str


def validate_template_format(
    original_bytes: bytes,
    template_bytes: bytes,
    template_spec: list[TemplateVariable],
) -> FormatValidationResult:
    """Detect format drift introduced by `create_template`.

    Returns `ok=True` when every template paragraph either:
      - Has no placeholders and matches an original paragraph verbatim
        (typography-tolerant), OR
      - Has placeholders that, when substituted back to their known
        markers, produce text found in the original (single paragraph
        or cross-paragraph join).

    Returns `ok=False` with a `DriftedParagraph` per template
    paragraph that couldn't be matched — the suspect set the caller
    should log, surface to the paralegal, or feed into a Tier 2 LLM
    fixer.

    Designed conservatively: false-positives (validator flags drift
    that's actually fine) are preferable to false-negatives (validator
    misses real drift). The caller decides what to do with the result
    — Tier 1 callers typically log + continue.
    """
    if not template_spec:
        return FormatValidationResult(ok=True, drifted_paragraphs=[], summary="No spec; skipping validation")

    orig_doc = Document(BytesIO(original_bytes))
    tmpl_doc = Document(BytesIO(template_bytes))

    placeholder_to_markers = _build_placeholder_marker_map(template_spec)
    if not placeholder_to_markers:
        return FormatValidationResult(ok=True, drifted_paragraphs=[], summary="No placeholders in spec; skipping validation")

    orig_paras = list(DocxTemplateService._iter_all_paragraphs(orig_doc))
    tmpl_paras = list(DocxTemplateService._iter_all_paragraphs(tmpl_doc))

    orig_texts = [p.text for p in orig_paras]
    orig_normalized = [_normalize_for_cross_para_scan(t) for t in orig_texts]

    drifted: list[DriftedParagraph] = []

    for tmpl_idx, tmpl_para in enumerate(tmpl_paras):
        tmpl_text = tmpl_para.text
        if not tmpl_text.strip():
            continue

        reconstructed, placeholders_substituted = _reconstruct_paragraph(
            tmpl_text, placeholder_to_markers,
        )
        reconstructed_normalized = _normalize_for_cross_para_scan(reconstructed)

        if _matches_original(
            reconstructed_normalized, orig_normalized,
        ):
            continue

        # Aliases retry — if we used the primary marker and missed,
        # try with each alias once.
        matched_via_alias = False
        if placeholders_substituted:
            for alt_reconstructed, _ in _reconstruct_with_aliases(
                tmpl_text, placeholder_to_markers,
            ):
                alt_normalized = _normalize_for_cross_para_scan(alt_reconstructed)
                if _matches_original(alt_normalized, orig_normalized):
                    matched_via_alias = True
                    break
        if matched_via_alias:
            continue

        drifted.append(DriftedParagraph(
            template_paragraph_index=tmpl_idx,
            template_text=tmpl_text,
            reconstructed_text=reconstructed,
            drift_reason=(
                "Reconstructed paragraph (placeholders restored to markers) "
                "not found in original — substitution likely lost a tab, "
                "line break, paragraph boundary, or surrounding whitespace."
            ),
        ))

    if drifted:
        summary = (
            f"Template format drift detected: {len(drifted)} paragraph(s) "
            f"could not be reconstructed from original (of {len(tmpl_paras)} total)."
        )
        return FormatValidationResult(ok=False, drifted_paragraphs=drifted, summary=summary)

    return FormatValidationResult(
        ok=True,
        drifted_paragraphs=[],
        summary=f"No drift detected across {len(tmpl_paras)} template paragraphs",
    )


def _matches_original(
    reconstructed_normalized: str,
    orig_normalized: list[str],
) -> bool:
    """A template paragraph "matches the original" when its
    placeholder-restored text equals (typography-tolerant) either:
      - One single original paragraph, OR
      - The newline-join of a contiguous slice of original paragraphs
        (covers cross-paragraph marker substitutions that collapsed
        N paragraphs into 1).

    EQUALITY is required — not substring containment. Substring matching
    would silently accept template paragraphs that lost surrounding
    text or merged with adjacent paragraphs.
    """
    if reconstructed_normalized in orig_normalized:
        return True
    if "\n" not in reconstructed_normalized:
        return False
    target = reconstructed_normalized
    n = len(orig_normalized)
    for start in range(n):
        for end in range(start + 2, n + 1):
            joined = "\n".join(orig_normalized[start:end])
            if joined == target:
                return True
            if len(joined) > len(target) + 256:
                break
    return False


def _build_placeholder_marker_map(
    spec: list[TemplateVariable],
) -> dict[str, list[str]]:
    """Map each `template_variable_string` to its list of markers
    (primary first, then aliases). Cleaned (`\\t` → space, literal
    `\\n` → real `\\n`)."""
    result: dict[str, list[str]] = {}
    for v in spec:
        if not v.template_variable_string:
            continue
        markers: list[str] = []
        if v.template_property_marker:
            markers.append(_clean_marker(v.template_property_marker))
        for alias in v.template_property_marker_aliases or []:
            if alias:
                markers.append(_clean_marker(alias))
        if markers:
            result[v.template_variable_string] = markers
    return result


def _clean_marker(marker: str) -> str:
    return marker.replace("\t", " ").replace("\\n", "\n")


def _reconstruct_paragraph(
    template_text: str,
    placeholder_to_markers: dict[str, list[str]],
) -> tuple[str, list[str]]:
    """Replace each placeholder in `template_text` with its PRIMARY
    marker. Returns (reconstructed_text, list_of_placeholders_substituted).

    Replacements are PROCESSED LONGEST-PLACEHOLDER-FIRST so a
    placeholder that is a substring of another (`[[case]]` vs
    `[[case_number]]`) doesn't accidentally consume part of the longer
    one. Mirrors `create_template`'s ordering contract.
    """
    result = template_text
    substituted: list[str] = []
    ordered_placeholders = sorted(
        placeholder_to_markers.keys(), key=len, reverse=True,
    )
    for placeholder in ordered_placeholders:
        if placeholder in result:
            primary_marker = placeholder_to_markers[placeholder][0]
            result = result.replace(placeholder, primary_marker)
            substituted.append(placeholder)
    return result, substituted


def _reconstruct_with_aliases(
    template_text: str,
    placeholder_to_markers: dict[str, list[str]],
):
    """Yield alternate reconstructions where each placeholder gets
    substituted by every one of its aliases (one alias per yield).

    Bounded: only yields N alternates where N is the total alias count
    across all placeholders in template_text — cheap.
    """
    ordered_placeholders = sorted(
        placeholder_to_markers.keys(), key=len, reverse=True,
    )
    for placeholder in ordered_placeholders:
        if placeholder not in template_text:
            continue
        markers = placeholder_to_markers[placeholder]
        for alias in markers[1:]:
            result = template_text
            # Substitute everyone with primary except this placeholder
            # with the alias.
            for p in ordered_placeholders:
                if p not in result:
                    continue
                use = alias if p == placeholder else placeholder_to_markers[p][0]
                result = result.replace(p, use)
            yield result, placeholder


# === Fill-time validation (Phase 2 dry-run, Phase 3 draft) ===


def validate_fill_format(
    template_bytes: bytes,
    filled_bytes: bytes,
    resolved_values: dict[str, str],
) -> FormatValidationResult:
    """Detect format drift introduced by `fill_template`.

    For each template paragraph, derive the EXPECTED filled paragraph
    by substituting each `[[placeholder]]` with its resolved value
    (in both "as-is" and " and "-joined forms, since
    `_substitute_placeholder` picks one per occurrence based on
    paragraph shape). The actual filled paragraph at the same index
    should equal one of the expected forms.

    Returns `ok=True` when every filled paragraph matches its
    template counterpart under at least one substitution form.

    Designed conservatively: any equality check that succeeds (raw,
    typography-folded, or alternate substitution form) lets the
    paragraph pass. False-positives (validator flags drift that's
    actually fine) preferred over false-negatives — Tier 1 only logs.

    `resolved_values` is keyed by VARIABLE NAME (not placeholder
    string) — same shape `fill_template` receives. The validator
    derives `[[name]]` placeholders internally.
    """
    if not resolved_values:
        return FormatValidationResult(
            ok=True, drifted_paragraphs=[],
            summary="No resolved values; skipping fill validation",
        )

    tmpl_doc = Document(BytesIO(template_bytes))
    filled_doc = Document(BytesIO(filled_bytes))

    tmpl_paras = list(DocxTemplateService._iter_all_paragraphs(tmpl_doc))
    filled_paras = list(DocxTemplateService._iter_all_paragraphs(filled_doc))

    # Fill preserves paragraph count (multi-line values render as
    # `<w:br/>` inside the SAME `<w:p>`, never split into new ones).
    # A mismatch here is itself a drift signal worth surfacing.
    if len(tmpl_paras) != len(filled_paras):
        return FormatValidationResult(
            ok=False,
            drifted_paragraphs=[
                DriftedParagraph(
                    template_paragraph_index=-1,
                    template_text=f"template has {len(tmpl_paras)} paragraphs",
                    reconstructed_text=f"filled has {len(filled_paras)} paragraphs",
                    drift_reason=(
                        "Paragraph count differs between template and filled — "
                        "fill_template should never add or remove paragraphs."
                    ),
                )
            ],
            summary=(
                f"Paragraph count mismatch: template={len(tmpl_paras)}, "
                f"filled={len(filled_paras)}"
            ),
        )

    placeholder_to_value = {f"[[{name}]]": val for name, val in resolved_values.items()}

    drifted: list[DriftedParagraph] = []
    for idx, (tmpl_para, filled_para) in enumerate(zip(tmpl_paras, filled_paras)):
        tmpl_text = _paragraph_text_with_breaks(tmpl_para)
        filled_text = _paragraph_text_with_breaks(filled_para)

        if not tmpl_text.strip() and not filled_text.strip():
            continue

        if not _filled_paragraph_matches(
            tmpl_para, tmpl_text, filled_text, placeholder_to_value,
        ):
            drifted.append(DriftedParagraph(
                template_paragraph_index=idx,
                template_text=tmpl_text,
                reconstructed_text=filled_text,
                drift_reason=(
                    "Filled paragraph does not equal template paragraph with "
                    "placeholders substituted to resolved values — fill_template "
                    "may have lost a tab, line break, or surrounding text."
                ),
            ))

    if drifted:
        return FormatValidationResult(
            ok=False,
            drifted_paragraphs=drifted,
            summary=(
                f"Fill format drift detected: {len(drifted)} paragraph(s) "
                f"diverged from template (of {len(tmpl_paras)} total)."
            ),
        )
    return FormatValidationResult(
        ok=True,
        drifted_paragraphs=[],
        summary=f"No fill drift detected across {len(tmpl_paras)} paragraphs",
    )


def _filled_paragraph_matches(
    tmpl_para,
    tmpl_text: str,
    filled_text: str,
    placeholder_to_value: dict[str, str],
) -> bool:
    """A filled paragraph "matches" its template counterpart when its
    text equals the template text with each placeholder substituted to
    the SAME form `_substitute_placeholder` would have picked.

    Replicates `fill_template`'s per-placeholder decision so the
    validator catches "wrong rendering form" drift (e.g. a caption-
    shape paragraph that should have rendered a multi-line value with
    `<w:br/>` but instead got " and "-joined).

    Decision rule (mirrors `_substitute_placeholder`):
      - Single-line value → drop in verbatim.
      - Multi-line value AND caption-shape paragraph → preserve `\\n`
        (will render as `<w:br/>` in the docx).
      - Multi-line value AND inline-shape paragraph → join lines with
        " and ".

    Typography-tolerant.
    """
    # Cheap exact-match path before doing any substitution.
    if tmpl_text == filled_text and not any(
        ph in tmpl_text for ph in placeholder_to_value
    ):
        return True

    expected = _render_expected_fill(tmpl_para, tmpl_text, placeholder_to_value)
    if expected == filled_text:
        return True
    if _normalize_for_cross_para_scan(expected) == _normalize_for_cross_para_scan(filled_text):
        return True
    return False


def _render_expected_fill(
    tmpl_para,
    tmpl_text: str,
    placeholder_to_value: dict[str, str],
) -> str:
    """Per-placeholder, build the expected filled text using a
    SMARTER caption-vs-inline decision than `_substitute_placeholder`'s.

    `_substitute_placeholder` calls `_is_caption_shape_paragraph` which
    removes ONLY the current placeholder and checks if the residue has
    alphanumerics. That mis-classifies paragraphs that contain MULTIPLE
    placeholders — e.g. `[[cos_section_1]]\\n[[cos_section_2]]` —
    because the OTHER placeholder's `[[…]]` token contributes alphanum
    chars to the residue. Misclassified as inline → multi-line values
    get joined with `" and "` (the Fleisher CoS bug).

    The validator's `_paragraph_is_caption_shape_for_predictor` strips
    ALL `[[…]]` tokens before the residue check, so multi-placeholder
    paragraphs that consist purely of placeholders + whitespace are
    correctly predicted as caption shape. When validator's prediction
    differs from `_substitute_placeholder`'s, the format-fixer fires
    and Sonnet corrects the docx.

    Longest-placeholder-first so a placeholder that is a substring of
    another doesn't accidentally consume part of the longer one.
    """
    is_caption_paragraph = _paragraph_is_caption_shape_for_predictor(tmpl_text)
    result = tmpl_text
    for placeholder, value in sorted(
        placeholder_to_value.items(), key=lambda kv: -len(kv[0])
    ):
        if placeholder not in result:
            continue
        if "\n" in value:
            rendered = value if is_caption_paragraph else " and ".join(value.split("\n"))
        else:
            rendered = value
        result = result.replace(placeholder, rendered)
    return result


_PLACEHOLDER_TOKEN_RE = re.compile(r"\[\[[A-Za-z0-9_]+\]\]")


def _paragraph_is_caption_shape_for_predictor(tmpl_text: str) -> bool:
    """Smarter caption-shape test for the format-validator.

    Strip ALL `[[…]]` tokens from the paragraph text, then check if
    any alphanumeric chars remain. Returns True (caption shape) when
    the paragraph consists ONLY of placeholders + whitespace +
    punctuation — meaning multi-line resolved values SHOULD render
    with `<w:br/>` between lines, not with `" and "` joining.

    Differs from `_is_caption_shape_paragraph` (in `docx_template`)
    which only strips the CURRENT placeholder; that function
    over-classifies multi-placeholder paragraphs as inline because the
    sibling placeholder's `[[name]]` letters count as alphanum residue.
    """
    stripped_of_placeholders = _PLACEHOLDER_TOKEN_RE.sub("", tmpl_text)
    return not any(ch.isalnum() for ch in stripped_of_placeholders)
