"""parse_document_v2 — flatten an uploaded .docx into text for TemplateAgentV2.

Mirrors v1's `parse_document()` 1:1; uses DocxTemplateService's
table-row extractor (read-only import — pure utility under the
no-touch rule's carve-out).
"""

import io

from docx import Document

from src.core.common.documents.docx_template import DocxTemplateService

from .schemas import DocumentParseResponseV2


async def parse_document_v2(filename: str, file_content: bytes) -> DocumentParseResponseV2:
    """Flatten a .docx upload into a single parsed-text string plus
    metadata for TemplateAgentV2.

    Body paragraphs join with `\\n`; tables flatten to row strings
    (cells joined by ` | `) via DocxTemplateService.extract_table_rows.
    """
    file_stream = io.BytesIO(file_content)
    doc = Document(file_stream)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        paragraphs.extend(DocxTemplateService.extract_table_rows(table))

    content = "\n".join(paragraphs)

    return DocumentParseResponseV2(
        document_id=filename,
        parsed=True,
        content=content,
        metadata={
            "format": "docx",
            "filename": filename,
            "content_length": len(content),
            "paragraph_count": len(paragraphs),
        },
    )
