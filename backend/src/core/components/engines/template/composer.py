"""Template composer: parse document → generate template → compose AgentConfig.

Three public operations, all used by the composer flow that turns a raw
uploaded .docx into a DraftTemplate row with a committed AgentConfig:

    parse_document(filename, file_content) -> DocumentParseResponse
        Flatten a .docx to text for the template agent's LLM prompt.

    generate_template(name, parsed_document, file_content) -> TemplateGenerateResponse
        Run the template agent on the parsed text, upload original + template
        docx to R2, and persist a DraftTemplate row with the extracted spec.

    compose_agent_config(template_id, template_spec) -> AgentConfig
        Validate + compose the user-filled template_spec into an AgentConfig;
        persist it on the DraftTemplate row. Also exposed internally as
        `build_agent_config` without the persist step for dry-run use.
"""

import io
import logging
import re
import uuid

from docx import Document
from docx.oxml.ns import qn
from fastapi import HTTPException

from src.core.agents.llm.template import MergeInstruction, TemplateAgent
from src.core.agents.types.sources import (
    AutoDerivedSourceParams,
    FieldSource,
    GroupDropdownComposite,
    GroupDropdownSourceParams,
)
from src.core.agents.types.spec import AgentConfig, TemplateField, TemplateVariable
from src.core.common.documents.docx_template import DocxTemplateService
from src.core.common.storage.database import DraftTemplateRepository
from src.core.common.storage.r2 import r2_service

from .schemas import (
    DocumentParseResponse,
    MergeOperation,
    RemovedVariableReason,
    TemplateGenerateResponse,
    TemplateRegenerateDiff,
)
from .validators import (
    GROUP_DROPDOWN_ANCHOR_SOURCES,
    assert_child_only_has_no_user_input,
    partner_variable_names,
    validate_template_spec_source_map,
)

logger = logging.getLogger(__name__)


async def parse_document(filename: str, file_content: bytes) -> DocumentParseResponse:
    """Flatten a .docx upload into a single parsed-text string plus metadata for the template agent."""
    file_stream = io.BytesIO(file_content)
    doc = Document(file_stream)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        paragraphs.extend(DocxTemplateService.extract_table_rows(table))

    content = "\n".join(paragraphs)

    return DocumentParseResponse(
        document_id=filename,
        parsed=True,
        content=content,
        metadata={
            "format": "docx",
            "filename": filename,
            "content_length": len(content),
            "paragraph_count": len(paragraphs),
        },
    )


async def _run_agent_and_upload_template_docx(
    template_id: str,
    parsed_content: str,
    file_content: bytes,
    ignored_texts: list[str] | None = None,
    merges: list[MergeOperation] | None = None,
    regeneration_instruction: str | None = None,
    previous_spec: list[TemplateVariable] | None = None,
) -> tuple[list[TemplateVariable], str]:
    """Run the template agent, build the template docx, upload to R2, and return `(template_spec, template_doc_url)`.

    Used by both the initial generate and the regenerate flows. The
    regenerate flow passes `previous_spec` (the author's confirmed
    baseline) so the agent preserves user-authored entries verbatim
    unless an explicit signal (merges / ignored_texts / regen_instruction)
    overrides; initial-generate passes `None`.
    """
    merge_instructions = (
        [
            MergeInstruction(
                new_variable_name=m.resolve_variable_name(),
                source_variables=m.source_variables,
                description=m.description,
            )
            for m in merges
        ]
        if merges
        else None
    )
    llm_result = await TemplateAgent.run(
        parsed_content,
        ignored_texts=ignored_texts,
        merges=merge_instructions,
        regeneration_instruction=regeneration_instruction,
        previous_spec=previous_spec,
    )

    cleaned_spec = _enforce_merges(llm_result.template_spec, merges or [])
    cleaned_spec = _ensure_joint_debtor_variable(file_content, cleaned_spec)

    template_doc_content = DocxTemplateService.create_template(
        file_content=file_content,
        template_spec=cleaned_spec,
    )
    cleaned_spec = _drop_orphan_variables(template_doc_content, cleaned_spec)

    await r2_service.upload_file(
        file_content=template_doc_content,
        template_id=template_id,
        filename="template.docx",
        prefix="template",
    )
    template_doc_url = await r2_service.get_presigned_url(
        template_id, "template.docx", prefix="template"
    )
    return cleaned_spec, template_doc_url


def _drop_orphan_variables(
    template_bytes: bytes,
    template_spec: list[TemplateVariable],
) -> list[TemplateVariable]:
    """Drop variables whose [[placeholder]] does not appear in the rendered template.docx.

    Catches LLM hallucinations and sub-tokens swallowed by a longer marker
    (e.g. 'her' extracted as its own variable when it lived inside another
    variable's marker — create_template's longest-first pass consumes the
    parent marker, so '[[her_var]]' never makes it into the template docx).

    Carve-outs:
    - **Virtual variables** (`kind == "virtual"`, i.e. `template_variable_string is None`)
      have no placeholder by design — they exist to power auto_derive children
      via their resolved value, never render in the docx. Skip the orphan check.
    - **Physical variables referenced as `dependent_variable`** by an
      AUTO_DERIVED_FROM_VARIABLE child are kept even when their own marker
      didn't make it into the rendered docx (e.g. swallowed by a longer
      marker). Defense-in-depth: dropping a parent breaks every child's
      derivation, so only drop if no child depends on it.
    """
    auto_derive_parent_names: set[str] = set()
    for var in template_spec:
        if var.source != FieldSource.AUTO_DERIVED_FROM_VARIABLE:
            continue
        params = var.source_params
        if isinstance(params, AutoDerivedSourceParams):
            auto_derive_parent_names.add(params.dependent_variable)

    placeholders = [
        v.template_variable_string
        for v in template_spec
        if v.kind == "physical" and v.template_variable_string
    ]
    missing = DocxTemplateService.find_missing_placeholders(template_bytes, placeholders)
    if not missing:
        return template_spec

    logger.warning(
        "Dropping %d orphan variable(s) from template_spec — placeholders not present in template.docx: %s",
        len(missing),
        sorted(missing),
    )
    return [
        v for v in template_spec
        if v.kind == "virtual"
        or v.template_variable in auto_derive_parent_names
        or v.template_variable_string not in missing
    ]


_IN_RE_PATTERN = re.compile(r"^\s*in\s*re\b", re.IGNORECASE)
_DEBTOR_ROLE_PATTERN = re.compile(r"\bdebtor(?:s|\(s\))?\.?\s*$", re.IGNORECASE)
_CASE_NO_PATTERN = re.compile(r"\bcase\s*(?:no|number)\b", re.IGNORECASE)
_CHAPTER_PATTERN = re.compile(r"^\s*chapter\s+\d+\s*$", re.IGNORECASE)
_NAME_SHAPE_PATTERN = re.compile(r"^[A-Za-z][A-Za-z.\-'\s]{1,78}[A-Za-z.,]$")


def _looks_like_name(line: str) -> bool:
    """Return True if `line` looks like a debtor name candidate.

    Rejects case-metadata lines (Case No., Chapter N) and anything that looks
    numeric, email-ish, or too long. Accepts hyphenated / apostrophed / middle-
    initial names. Intentionally conservative — a false negative means the
    detector falls through to the LLM; a false positive would corrupt the marker.
    """
    if not line or _CASE_NO_PATTERN.search(line) or _CHAPTER_PATTERN.match(line):
        return False
    if any(ch.isdigit() for ch in line):
        return False
    if "@" in line or "/" in line:
        return False
    return bool(_NAME_SHAPE_PATTERN.match(line))


def _strip_debtor_role_suffix(line: str) -> str:
    """If `line` ends with 'Debtor', 'Debtors.', 'Debtor(s).', etc., return it without that suffix.

    Real-world captions sometimes glue the role marker onto the last name line
    ('Robert Creswell, Debtors.'). Detector strips the role so the name alone
    goes into the marker.
    """
    stripped = _DEBTOR_ROLE_PATTERN.sub("", line).rstrip(" ,.")
    return stripped


def _extract_joint_debtor_names(file_content: bytes) -> list[str] | None:
    """Scan the source docx for an 'In re:' caption block with 2+ debtor names.

    Flattens every body paragraph AND every table-cell paragraph on BOTH `\\n`
    (python-docx surfaces `<w:br/>` as `\\n` in `p.text`) AND `\\t` (captions
    often glue 'In re:' to 'Case No:' via tab) so all of the observed caption
    shapes resolve to the same flat line list:

      - standalone 'In re:' line + multiple name paragraphs
      - 'In re:\\tCase No: ...' header + 'Chapter N' line + names
      - both names in ONE paragraph separated by `<w:br/>`
      - glued role marker on last name ('Robert Creswell, Debtors.')

    Returns the ordered list of name strings (preserving trailing comma / period
    from the source) or None if no joint caption is found.
    """
    doc = Document(io.BytesIO(file_content))

    def _paragraph_text_with_breaks(p) -> str:
        """Reconstruct paragraph text from XML walk so <w:br/> contributes `\\n`.

        python-docx's `p.text` in some versions strips <w:br/> entirely —
        which would make joint-debtor paragraphs (two names joined by a soft
        break) look like 'LoriCreswellRobertCreswell,' and hide the joint-ness
        from the detector. Walking <w:t> and <w:br/> in document order is the
        canonical way to get the right text.
        """
        parts: list[str] = []
        for elem in p._element.iter(qn("w:t"), qn("w:br")):
            if elem.tag == qn("w:t"):
                parts.append(elem.text or "")
            else:
                parts.append("\n")
        return "".join(parts)

    def _flatten(paragraphs) -> list[str]:
        lines: list[str] = []
        for p in paragraphs:
            text = _paragraph_text_with_breaks(p)
            for chunk in re.split(r"[\n\t]", text):
                stripped = chunk.strip()
                if stripped:
                    lines.append(stripped)
        return lines

    def _scan(lines: list[str]) -> list[str] | None:
        for i, text in enumerate(lines):
            if not _IN_RE_PATTERN.match(text):
                continue
            names: list[str] = []
            for j in range(i + 1, min(i + 15, len(lines))):
                candidate = lines[j]
                if _DEBTOR_ROLE_PATTERN.search(candidate):
                    head = _strip_debtor_role_suffix(candidate)
                    if head and _looks_like_name(head):
                        names.append(head + "," if not head.endswith(",") else head)
                    break
                if _CASE_NO_PATTERN.search(candidate) or _CHAPTER_PATTERN.match(candidate):
                    continue
                if _looks_like_name(candidate):
                    names.append(candidate)
                    continue
                break
            if len(names) >= 2:
                logger.info(
                    "Joint-caption scan matched after 'In re:' at line %d; collected %d names: %s",
                    i, len(names), names,
                )
                return names
        return None

    body_lines = _flatten(doc.paragraphs)
    found = _scan(body_lines)
    if found:
        return found

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_lines = _flatten(cell.paragraphs)
                if not cell_lines:
                    continue
                found = _scan(cell_lines)
                if found:
                    logger.info(
                        "Joint-caption scan matched in table %d row %d cell %d",
                        t_idx, r_idx, c_idx,
                    )
                    return found
    return None


def _ensure_joint_debtor_variable(
    file_content: bytes,
    template_spec: list[TemplateVariable],
) -> list[TemplateVariable]:
    """Guarantee a correct `debtor_name` variable when the source is a joint caption.

    The TemplateAgent frequently misses or mangles joint-filing captions —
    especially the common shape where two names share one paragraph separated
    by a `<w:br/>`. This post-processor runs a deterministic scan of the
    source docx; if a joint caption is detected, it replaces any existing
    `debtor_name` variable with one carrying a `\\n`-joined marker so the
    docx engine's stage-3 span-with-break replacement path fires.

    Idempotent: if an existing `debtor_name` already has the correct marker,
    the spec is returned unchanged. Solo documents are no-ops.
    """
    names = _extract_joint_debtor_names(file_content)
    if not names or len(names) < 2:
        return template_spec

    joined_marker = "\n".join(names)
    existing = next(
        (v for v in template_spec if v.template_variable == "debtor_name"),
        None,
    )
    if existing and existing.template_property_marker == joined_marker:
        return template_spec

    logger.info(
        "Joint-debtor post-processor: synthesizing debtor_name variable with %d-line marker",
        len(names),
    )

    new_var = TemplateVariable(
        template_variable="debtor_name",
        template_index=existing.template_index if existing else 0,
        template_property_marker=joined_marker,
        template_variable_string="[[debtor_name]]",
        template_identifying_text_match="In re:\n" + joined_marker + "\n  Debtors.",
        description=(
            f"Debtor name(s) in the 'In re:' caption. Joint filing with "
            f"{len(names)} debtors; renders as soft line breaks at fill time."
        ),
        source=None,
        source_params=None,
    )

    if existing:
        return [new_var if v.template_variable == "debtor_name" else v for v in template_spec]
    return [new_var, *template_spec]


def _enforce_merges(
    template_spec: list[TemplateVariable],
    merges: list[MergeOperation],
) -> list[TemplateVariable]:
    """Enforce that each merge's source variables are absent and the merged variable is present in the agent output.

    Drops any stale source variables the agent failed to remove (belt-and-
    suspenders). Raises 400 if a merged variable is missing — the agent
    didn't comply with the merge instruction and the caller should retry
    or adjust the merge."""
    if not merges:
        return template_spec

    errors: list[str] = []
    spec = list(template_spec)

    for merge in merges:
        names_present = {v.template_variable for v in spec}
        resolved_name = merge.resolve_variable_name()

        if resolved_name not in names_present:
            errors.append(
                f"Merge: expected merged variable '{resolved_name}' "
                f"in agent output but it's missing. Agent may not have followed the "
                f"merge instruction; try again or adjust the source_variables."
            )
            continue

        stale_sources = [s for s in merge.source_variables if s in names_present]
        if stale_sources:
            spec = [v for v in spec if v.template_variable not in stale_sources]

    if errors:
        raise HTTPException(status_code=422, detail={"merge_errors": errors})

    return spec


def _normalize_for_match(text: str) -> str:
    """Whitespace-collapsed, lowercased text for fuzzy fragment matching."""
    return " ".join(text.lower().split())


def _compute_regenerate_diff(
    baseline: list[TemplateVariable],
    new_spec: list[TemplateVariable],
    merges: list[MergeOperation],
    ignored_texts: list[str],
) -> TemplateRegenerateDiff:
    """Partition baseline + new_spec names into added / removed / preserved.

    For each removed name (in baseline but not in new spec), classify the
    reason so the FE can render an appropriate annotation:

      - `merged`     → name appears in any merge's `source_variables` list.
        `merged_into` is the merge's resolved target name.
      - `ignored`    → baseline entry's `template_identifying_text_match`
        overlaps any `ignored_texts` fragment after whitespace + casing
        normalization (substring match in either direction). Catches the
        common case where the author flagged a row in the regenerate
        modal and the agent honored the ignore directive.
      - `unexpected` → neither applies. Surfaces as drift in the UI so
        the author can audit whether to re-add via the next iteration.
    """
    baseline_names = [v.template_variable for v in baseline]
    new_names = {v.template_variable for v in new_spec}

    added = sorted(n for n in new_names if n not in set(baseline_names))
    preserved = sorted(n for n in baseline_names if n in new_names)

    # Build lookup tables once.
    merge_target_by_source: dict[str, str] = {}
    for merge in merges:
        target = merge.resolve_variable_name()
        for source_name in merge.source_variables:
            merge_target_by_source[source_name] = target

    normalized_ignored = [_normalize_for_match(t) for t in ignored_texts if t and t.strip()]
    baseline_by_name = {v.template_variable: v for v in baseline}

    removed: list[RemovedVariableReason] = []
    for name in baseline_names:
        if name in new_names:
            continue
        if name in merge_target_by_source:
            removed.append(RemovedVariableReason(
                name=name,
                reason="merged",
                merged_into=merge_target_by_source[name],
            ))
            continue
        identifying = (baseline_by_name[name].template_identifying_text_match or "").strip()
        if identifying and normalized_ignored:
            needle = _normalize_for_match(identifying)
            if any(needle in haystack or haystack in needle for haystack in normalized_ignored):
                removed.append(RemovedVariableReason(name=name, reason="ignored"))
                continue
        removed.append(RemovedVariableReason(name=name, reason="unexpected"))

    return TemplateRegenerateDiff(
        added=added,
        removed=removed,
        preserved=preserved,
    )


async def generate_template(
    template_name: str,
    parsed_document: DocumentParseResponse,
    file_content: bytes,
) -> TemplateGenerateResponse:
    """Run the template agent on a parsed doc, upload the original + placeholder docx to R2, and persist a DraftTemplate row."""
    template_id = str(uuid.uuid4())

    await r2_service.upload_file(
        file_content=file_content,
        template_id=template_id,
        filename="original.docx",
        prefix="template",
    )

    template_spec, template_doc_url = await _run_agent_and_upload_template_docx(
        template_id=template_id,
        parsed_content=parsed_document.content,
        file_content=file_content,
    )

    original_doc_url = await r2_service.get_presigned_url(
        template_id, "original.docx", prefix="template"
    )

    template_spec_dict = [var.model_dump() for var in template_spec]
    await DraftTemplateRepository.create(
        template_id=template_id,
        name=template_name,
        original_doc_url=original_doc_url,
        template_doc_url=template_doc_url,
        template_spec=template_spec_dict,
    )

    return TemplateGenerateResponse(
        template_id=template_id,
        template_name=template_name,
        template_spec=template_spec,
        generated=True,
        original_doc_url=original_doc_url,
        template_doc_url=template_doc_url,
    )


async def regenerate_template(
    template_id: str,
    ignored_texts: list[str],
    merges: list[MergeOperation] | None = None,
    regeneration_instruction: str | None = None,
) -> TemplateGenerateResponse:
    """Re-run the template agent on the existing original.docx with optional post-generation transforms.

    Transforms supported: ignore specific boilerplate text fragments,
    collapse existing variables into merged ones, and/or steer the
    re-extraction with a free-form `regeneration_instruction`.
    Overwrites template_spec + template.docx on the same DraftTemplate
    row and clears the previously-composed agent_config because the
    variable set changed.
    """
    existing = await DraftTemplateRepository.get(template_id)
    if existing is None:
        raise HTTPException(status_code=404, detail=f"Template '{template_id}' not found")

    # Load the author's confirmed baseline from the existing row so the
    # agent preserves verbatim entries unless explicitly signalled to
    # change. None for templates that never had a successful first
    # generation; the agent extracts from scratch in that case.
    previous_spec = (
        [TemplateVariable(**v) for v in existing.template_spec]
        if existing.template_spec
        else None
    )

    file_content = await r2_service.download_file(
        template_id, "original.docx", prefix="template"
    )
    parsed = await parse_document("original.docx", file_content)

    template_spec, template_doc_url = await _run_agent_and_upload_template_docx(
        template_id=template_id,
        parsed_content=parsed.content,
        file_content=file_content,
        ignored_texts=ignored_texts,
        merges=merges,
        regeneration_instruction=regeneration_instruction,
        previous_spec=previous_spec,
    )

    diff = (
        _compute_regenerate_diff(
            baseline=previous_spec,
            new_spec=template_spec,
            merges=merges or [],
            ignored_texts=ignored_texts or [],
        )
        if previous_spec
        else None
    )

    original_doc_url = await r2_service.get_presigned_url(
        template_id, "original.docx", prefix="template"
    )
    template_spec_dict = [var.model_dump() for var in template_spec]
    await DraftTemplateRepository.update(
        template_id=template_id,
        template_spec=template_spec_dict,
        clear_agent_config=True,
    )

    return TemplateGenerateResponse(
        template_id=template_id,
        template_name=existing.name,
        template_spec=template_spec,
        generated=True,
        original_doc_url=original_doc_url,
        template_doc_url=template_doc_url,
        diff=diff,
    )


async def build_agent_config(
    template_id: str,
    template_spec: list[TemplateVariable],
) -> AgentConfig:
    """Validate + compose template_spec → AgentConfig (no persist).

    Shared by compose_agent_config (persists) and the dry-run flow (in-memory
    only). Collapses group-dropdown sibling pairs into a single composite
    TemplateField so downstream consumers see one entry per dropdown.
    """
    await validate_template_spec_source_map(template_spec)

    by_name = {var.template_variable: var for var in template_spec}
    partner_names = partner_variable_names(template_spec)

    template_fields: list[TemplateField] = []
    for var in template_spec:
        if var.template_variable in partner_names:
            continue  # collapsed into its anchor's composite below

        if var.source in GROUP_DROPDOWN_ANCHOR_SOURCES:
            anchor_params = var.source_params
            if not isinstance(anchor_params, GroupDropdownSourceParams):
                continue  # validator already errored; defensive
            partner = by_name.get(anchor_params.right_partner_variable)
            if partner is None:
                continue  # validator already errored
            composite = GroupDropdownComposite(
                subject_query=anchor_params.subject_query,
                body_query=anchor_params.body_query,
                group_label=anchor_params.group_label,
                left_variable=var.template_variable,
                left_label=anchor_params.left_label,
                left_template_variable_string=var.template_variable_string or "",
                right_variable=partner.template_variable,
                right_label=anchor_params.right_label,
                right_template_variable_string=partner.template_variable_string or "",
            )
            template_fields.append(TemplateField(
                property_name=f"{var.template_variable}__{partner.template_variable}",
                source=var.source,
                source_params=composite,
                instruction=var.instruction,
                output_instruction=var.output_instruction,
                template_variable_string=None,
                template_property_marker=var.template_property_marker,
                template_identifying_text_match=var.template_identifying_text_match,
            ))
            continue

        template_fields.append(TemplateField(
            property_name=var.template_variable,
            source=var.source,
            # Preserve source_params for case_vector — the optional
            # `CaseVectorSourceParams.text_query` is read by both the
            # pgvector retrieval handler (overrides property-name auto-
            # derivation) AND the CaseVectorVisionAgent prompt builder
            # (surfaced as the authoritative topical query). Stripping
            # it here silently breaks both paths whenever the author
            # types a query.
            source_params=var.source_params,
            instruction=var.instruction,
            output_instruction=var.output_instruction,
            template_variable_string=var.template_variable_string,
            template_property_marker=var.template_property_marker,
            template_identifying_text_match=var.template_identifying_text_match,
        ))

    return AgentConfig(
        template_id=template_id,
        template_fields=template_fields,
    )


async def compose_agent_config(
    template_id: str,
    template_spec: list[TemplateVariable],
) -> AgentConfig:
    """Validate + compose the user-filled template_spec into an AgentConfig and persist it; raise 404 on missing/deleted template."""
    existing = await DraftTemplateRepository.get(template_id)
    if not existing:
        raise HTTPException(status_code=404, detail=f"Template '{template_id}' not found")

    if existing.bundle_role == "child_only":
        assert_child_only_has_no_user_input(template_id, template_spec)

    agent_config = await build_agent_config(template_id, template_spec)

    template_spec_dict = [var.model_dump() for var in template_spec]
    agent_config_dict = agent_config.model_dump()

    await DraftTemplateRepository.update(
        template_id=template_id,
        template_spec=template_spec_dict,
        agent_config=agent_config_dict,
    )

    return agent_config
