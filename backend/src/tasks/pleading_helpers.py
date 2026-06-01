"""
Shared helpers and constants for pleading tasks.

This module contains utilities used by both pleading extraction and generation tasks.
"""

import logging
import re
import shutil
import traceback
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from .orchestrator import get_document_generator

logger = logging.getLogger(__name__)


# Called by: _store_motion_in_vectorstore() (same file)
def _extract_text_from_docx(docx_path: Path) -> str:
    """Extract all text content from a DOCX file, including tables."""
    try:
        from docx import Document
        doc = Document(str(docx_path))

        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning(f"Failed to extract text from {docx_path}: {e}")
        return ""


# Called by: _generate_motion_doc() (same file), _generate_service_doc() (same file)
def _store_motion_in_vectorstore(session_id: str, doc_type: str, docx_path: Path, filename: str):
    """Extract text from DOCX and store in vectorstore for agent access."""
    try:
        from ..chatbot.vectorestore import store_generated_motion

        content = _extract_text_from_docx(docx_path)
        if not content:
            logger.warning(f"No content extracted from {docx_path}, skipping vectorstore storage")
            return

        result = store_generated_motion(
            session_id=session_id,
            doc_type=doc_type,
            content=content,
            filename=filename
        )
        if result.get("success"):
            logger.info(f"Stored {doc_type} in vectorstore: {result.get('stored_count')} chunks")
        else:
            logger.warning(f"Failed to store {doc_type} in vectorstore: {result.get('error')}")
    except Exception as e:
        logger.exception(f"Error storing {doc_type} in vectorstore: {e}")

_SLUG_STRIP_RE = re.compile(r"[^\w\s\-]")
_SLUG_SPACE_RE = re.compile(r"\s+")
_CASE_NUMBER_RE = re.compile(r"\d{2}-\d{5}")


def _slugify(s: str) -> str:
    return _SLUG_SPACE_RE.sub("_", _SLUG_STRIP_RE.sub("", s).strip())


def build_download_filename(prefix: str, case_name: str | None, case_number: str | None, ext: str) -> str:
    """Build a user-friendly download filename: {prefix}_{case_name}_{case_number}.{ext}.
    Falls back to {prefix}.{ext} when both case fields are absent.
    Case number is normalized to XX-XXXXX format (judge initials and other suffixes are stripped).
    """
    parts = [prefix]
    if case_name:
        parts.append(_slugify(case_name))
    if case_number:
        m = _CASE_NUMBER_RE.search(case_number)
        if m:
            parts.append(m.group(0))

    return f"{'_'.join(parts)}.{ext.lstrip('.')}"


PDF_ERROR_LOG = Path(__file__).parent.parent / "pdf_errors.log"

GENERATED_PDF_DIR = Path(__file__).parent.parent.parent / 'generated_pdf'
GENERATED_DOCX_DIR = Path(__file__).parent.parent.parent / 'generated_docx'
GENERATED_PDF_DIR.mkdir(parents=True, exist_ok=True)
GENERATED_DOCX_DIR.mkdir(parents=True, exist_ok=True)

OUT_DIR = Path(__file__).parent.parent / "motion_filling" / "out"

# Called by: _generate_motion_doc(), _copy_to_legacy_dir(), _check_existing_documents() (same file)
MOTION_TYPE_PREFIX_MAP = {
    'extend': 'motion_to_extend',
    'modify': 'motion_to_modify',
    'value': 'motion_to_value',
    'loe': 'motion_to_loe',
    'withdraw': 'motion_to_withdraw',
    'waive': 'motion_to_waive',
    'claim': 'motion_to_claim',
    'delay': 'motion_to_delay',
    'reinstate': 'motion_to_reinstate',
    'service': 'certificate_of_service',
    'certificate-of-service': 'standalone_certificate_of_service',
    'suggestion': 'motion_to_suggestion',
    'objection-sustain': 'order_sustaining_objection',
    'order-extend': 'order_on_motion_to_extend',
    'order-waive': 'order_on_motion_to_waive',
    'order-withdraw': 'order_on_motion_to_withdraw',
    'order-reinstate': 'order_on_motion_to_reinstate',
    'order-extension': 'order_on_motion_for_extension',
    'ex-parte-extension': 'ex_parte_motion_for_extension',
    'notice-withdraw': 'notice_to_withdraw',
}


# Called by: _generate_motion_doc(), _generate_service_doc() (same file)
def log_pdf_error(context: str, error: Exception):
    """Log PDF generation errors to a dedicated file for easy debugging."""
    timestamp = datetime.now().isoformat()
    with open(PDF_ERROR_LOG, "a") as f:
        f.write(f"\n{'='*60}\n")
        f.write(f"[{timestamp}] {context}\n")
        f.write(f"Error: {error}\n")
        f.write(f"Traceback:\n{traceback.format_exc()}\n")


# Called by: _generate_motion_doc(), _generate_service_doc() (same file)
def _copy_to_legacy_dir(source_path: Path, motion_type: str, session_id: str) -> Optional[Path]:
    """
    Copy generated file to legacy directory with correct naming.
    This enables the /api/motions/existence API to find worker-generated documents.
    """
    if not source_path or not source_path.exists():
        return None

    prefix = MOTION_TYPE_PREFIX_MAP.get(motion_type, f"motion_to_{motion_type}")
    ext = source_path.suffix

    if ext == '.pdf':
        dest_dir = GENERATED_PDF_DIR
    elif ext == '.docx':
        dest_dir = GENERATED_DOCX_DIR
    else:
        return None

    dest_filename = f"{prefix}_{session_id}{ext}"
    dest_path = dest_dir / dest_filename

    try:
        shutil.copy2(source_path, dest_path)
        logger.info(f"Copied {source_path.name} -> {dest_path}")
        return dest_path
    except Exception as e:
        logger.error(f"Failed to copy to legacy dir: {e}")
        return None


# Called by: pleading_tasks.extract_pleading_payload() (tasks/pleading_tasks.py)
def _check_existing_documents(
    session_id: str,
    motion_type: str,
    include_cos: bool = False,
) -> Optional[dict[str, Any]]:
    """Check if documents already exist for this session + motion type."""
    prefix = MOTION_TYPE_PREFIX_MAP.get(motion_type, f"motion_to_{motion_type}")

    docx_filename = f"{prefix}_{session_id}.docx"
    pdf_filename = f"{prefix}_{session_id}.pdf"

    docx_path = GENERATED_DOCX_DIR / docx_filename
    pdf_path = GENERATED_PDF_DIR / pdf_filename

    docx_exists = docx_path.exists()
    pdf_exists = pdf_path.exists()

    if docx_exists or pdf_exists:
        docx_motion_matches = sorted(
            OUT_DIR.glob(f"{motion_type}_{session_id}*.docx"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )
        pdf_motion_matches = sorted(
            OUT_DIR.glob(f"{motion_type}_{session_id}*.pdf"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )

        result = {
            "motion": {
                "docx_url": f"/api/pleadings/files/{docx_motion_matches[0].name}" if docx_motion_matches else None,
                "pdf_url": f"/api/pleadings/files/{pdf_motion_matches[0].name}" if pdf_motion_matches else None,
            }
        }

        if include_cos:
            cos_docx_matches = sorted(
                OUT_DIR.glob(f"cos_{motion_type}_{session_id}_*.docx"),
                key=lambda p: p.stat().st_mtime,
                reverse=True,
            )
            cos_pdf_matches = sorted(
                OUT_DIR.glob(f"cos_{motion_type}_{session_id}_*.pdf"),
                key=lambda p: p.stat().st_mtime,
                reverse=True,
            )
            if cos_docx_matches or cos_pdf_matches:
                result["certificate_of_service"] = {
                    "docx_url": f"/api/pleadings/files/{cos_docx_matches[0].name}" if cos_docx_matches else None,
                    "pdf_url": f"/api/pleadings/files/{cos_pdf_matches[0].name}" if cos_pdf_matches else None,
                }

        return result

    return None


# Called by: _generate_documents() (tasks/pleading_tasks.py)
def _generate_motion_doc(generator_func, payload: dict, name_slug: str, session_id: str, motion_type: str) -> dict[str, str]:
    """Generate DOCX and PDF for a motion and copy outputs to the legacy discovery directory.

    Args:
        generator_func: Dict with "docx" and "pdf" callable keys from get_document_generator().
        payload: Field dict to pass to each generator callable.
        name_slug: Debtor-name slug used as part of the output filename.
        session_id: Active session identifier.
        motion_type: Motion type key (used to derive the filename prefix).

    Returns:
        Dict with keys "docx_url" and "pdf_url" (either a relative URL string or None).
    """
    print(f"\n{'='*60}")
    print(f">>> _generate_motion_doc CALLED for {name_slug}")
    print(f"{'='*60}\n")

    docx_url = None
    pdf_url = None

    try:
        print(">>> Generating DOCX...")
        docx_path = generator_func["docx"](payload, name_slug)
        if docx_path:
            docx_url = f"/api/pleadings/files/{docx_path.name}"
            print(f">>> DOCX SUCCESS: {docx_path.name}")
            _copy_to_legacy_dir(docx_path, motion_type, session_id)
            _store_motion_in_vectorstore(session_id, motion_type, docx_path, docx_path.name)
    except Exception as e:
        print(f">>> DOCX FAILED: {e}")
        logger.exception(f"DOCX generation failed: {e}")

    try:
        print(">>> Generating PDF...")
        pdf_path = generator_func["pdf"](payload, name_slug)
        if pdf_path:
            pdf_url = f"/api/pleadings/files/{pdf_path.name}"
            print(f">>> PDF SUCCESS: {pdf_path.name}")
            _copy_to_legacy_dir(pdf_path, motion_type, session_id)
        else:
            print(f">>> PDF returned None (no error raised)")
    except Exception as e:
        print(f">>> PDF FAILED: {e}")
        logger.exception(f"PDF generation failed: {e}")
        log_pdf_error(f"Motion PDF generation failed for {name_slug}", e)

    print(f">>> Returning: docx_url={docx_url}, pdf_url={pdf_url}")
    return {"docx_url": docx_url, "pdf_url": pdf_url}


# Called by: _generate_documents() (tasks/pleading_tasks.py)
def _generate_service_doc(payload: dict, name_slug: str, session_id: str) -> dict[str, str]:
    """Generate DOCX and PDF for a Certificate of Service.

    Args:
        payload: CertificateOfServicePayload fields as a dict.
        name_slug: Debtor-name slug used as part of the output filename.
        session_id: Active session identifier.

    Returns:
        Dict with keys "docx_url" and "pdf_url" (either a relative URL string or None).
    """
    from ..motion_filling.fill_motion_service import (
        generate_docx_from_payload,
        generate_pdf_from_payload
    )

    docx_url = None
    pdf_url = None

    try:
        docx_path = generate_docx_from_payload(payload, name_slug)
        if docx_path:
            docx_url = f"/api/pleadings/files/{docx_path.name}"
            logger.info(f"Service DOCX generated successfully: {docx_path.name}")
            _copy_to_legacy_dir(docx_path, "service", session_id)
            _store_motion_in_vectorstore(session_id, "certificate_of_service", docx_path, docx_path.name)
    except Exception as e:
        logger.exception(f"Service DOCX generation failed: {e}")

    try:
        pdf_path = generate_pdf_from_payload(payload, name_slug)
        if pdf_path:
            pdf_url = f"/api/pleadings/files/{pdf_path.name}"
            logger.info(f"Service PDF generated successfully: {pdf_path.name}")
            _copy_to_legacy_dir(pdf_path, "service", session_id)
    except Exception as e:
        logger.exception(f"Service PDF generation failed: {e}")
        log_pdf_error(f"Service PDF generation failed for {name_slug}", e)

    return {"docx_url": docx_url, "pdf_url": pdf_url}


# Called by: pleading_tasks.generate_pleading_documents() (tasks/pleading_tasks.py)
def _merge_user_input(payload: dict, user_input: dict) -> dict:
    """Merge user-supplied field values into an extracted payload dict.

    Applies a set of key alias mappings so that user input keys that differ
    from the Pydantic field names are also written to the canonical field names.

    Args:
        payload: Extracted payload dict from the LLM extraction phase.
        user_input: Dict of user-supplied overrides (from AWAITING_INPUT submission).

    Returns:
        Merged dict with user input applied over the extracted payload.
    """
    if payload is None:
        return user_input

    merged = payload.copy()

    field_mappings = {
        "debtor_address": ["DebtorAddress", "debtor_current_address"],
        "reason_for_extension": ["DismissalReason", "dismissal_reason"],
        "change_in_circumstances": ["ChangeInCircumstances"],
        "reason_for_withdrawal": ["WithdrawReason", "withdraw_reason"],
        "delinquent_reason": ["DelinquentReason"],
        "basis_for_objection": ["Basis"],
        "explanation": ["Explanation"],
        "employment_explanation": ["EmploymentExplanation"],
    }

    for input_key, input_value in user_input.items():
        merged[input_key] = input_value

        if input_key in field_mappings:
            for mapped_key in field_mappings[input_key]:
                merged[mapped_key] = input_value

    return merged


# Called by: pleading_tasks.generate_pleading_documents() (tasks/pleading_tasks.py)
def _build_order_sustaining_payload(claim_payload: dict) -> dict:
    """Build Order on Objection (objection-sustain) payload from claim payload."""
    return {
        "DebtorName": claim_payload.get("DebtorName", ""),
        "CaseNo": claim_payload.get("CaseNumber", ""),
        "Chapter": claim_payload.get("Chapter", "N/A"),
        "CalendarDate": "N/A",
        "SlotNumb": claim_payload.get("Slot", ""),
        "Creditor": claim_payload.get("ClaimantName", ""),
        "Docket": "N/A",
    }
