"""Tests for `parse_document_v2` — the .docx upload flattener that
feeds TemplateAgentV2. Verifies the body-paragraph + table-row
extraction without needing a full upload pipeline.
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from src.core.studio_v2.services.composer.parse import parse_document_v2


def _make_docx_with(paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> bytes:
    """Build a minimal in-memory .docx with the given paragraphs +
    optional 2D table rows. Returns the serialized bytes the upload
    handler would receive."""
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = len(table_data[0]) if rows > 0 else 0
            tbl = doc.add_table(rows=rows, cols=cols)
            for r, row in enumerate(table_data):
                for c, cell in enumerate(row):
                    tbl.cell(r, c).text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.unit
@pytest.mark.asyncio
async def test_parse_document_returns_paragraph_text():
    """Each non-empty paragraph appears in the joined content."""
    docx_bytes = _make_docx_with(["First line.", "Second line.", "Third line."])
    result = await parse_document_v2("motion.docx", docx_bytes)
    assert result.parsed is True
    assert result.document_id == "motion.docx"
    assert "First line." in result.content
    assert "Second line." in result.content
    assert "Third line." in result.content
    assert result.metadata["filename"] == "motion.docx"
    assert result.metadata["format"] == "docx"
    assert result.metadata["paragraph_count"] == 3
    assert result.metadata["content_length"] == len(result.content)


@pytest.mark.unit
@pytest.mark.asyncio
async def test_parse_document_strips_empty_paragraphs():
    """Blank / whitespace-only paragraphs are dropped (matches v1's
    parse_document behavior)."""
    docx_bytes = _make_docx_with(["Real content", "", "   ", "More content"])
    result = await parse_document_v2("doc.docx", docx_bytes)
    assert result.metadata["paragraph_count"] == 2
    assert "Real content" in result.content
    assert "More content" in result.content


@pytest.mark.unit
@pytest.mark.asyncio
async def test_parse_document_includes_table_rows():
    """Tables are extracted via DocxTemplateService.extract_table_rows
    and flattened into the paragraph stream."""
    docx_bytes = _make_docx_with(
        paragraphs=["Header paragraph"],
        tables=[
            [
                ["Debtor", "Amount"],
                ["Acme Bank", "$1,200"],
                ["Wells Fargo", "$3,400"],
            ],
        ],
    )
    result = await parse_document_v2("schedule.docx", docx_bytes)
    assert result.parsed is True
    assert "Header paragraph" in result.content
    # Row content surfaces (exact format depends on
    # DocxTemplateService.extract_table_rows; cell-by-cell strings
    # MUST appear regardless of separator).
    assert "Acme Bank" in result.content
    assert "$1,200" in result.content
    assert "Wells Fargo" in result.content
    assert "$3,400" in result.content


@pytest.mark.unit
@pytest.mark.asyncio
async def test_parse_document_empty_docx():
    """A docx with no paragraphs or tables still returns a valid
    response (empty content, paragraph_count=0)."""
    docx_bytes = _make_docx_with(paragraphs=[])
    result = await parse_document_v2("empty.docx", docx_bytes)
    assert result.parsed is True
    assert result.content == ""
    assert result.metadata["paragraph_count"] == 0
    assert result.metadata["content_length"] == 0
