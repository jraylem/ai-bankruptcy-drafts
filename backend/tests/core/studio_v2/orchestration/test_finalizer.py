"""Tests for finalize_run_v2 — heal ordering, fill behavior, R2
upload, warnings shaping. We patch DocxTemplateService primitives +
r2_service so no real files are written."""

from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document

from src.core.studio_v2.resolvers.date_healing import DateHealingResolverV2
from src.core.studio_v2.agents.heal import UserInputHealAgentV2
from src.core.studio_v2.orchestration.finalizer import (
    _build_warnings,
    _fill_template_v2,
    finalize_run_v2,
)
from src.core.studio_v2.types.fields import TemplateFieldV2, TemplateSpecV2
from src.core.studio_v2.types.resolution import ResolvedTemplateValueV2
from src.core.studio_v2.types.wizard_sources import SourceKind, WizardSourceParams


_TEMPLATE_UUID = "00000000-0000-0000-0000-000000000099"
_FIELD_UUID = "00000000-0000-0000-0000-000000000001"


def _field(name):
    return TemplateFieldV2(
        id=_FIELD_UUID, template_id=_TEMPLATE_UUID,
        template_variable=name,
        params=WizardSourceParams(source=SourceKind.CURRENT_DATE),
    )


def _spec(fields):
    return TemplateSpecV2(template_id=_TEMPLATE_UUID, fields=fields)


def _docx_bytes_with_text(text: str) -> bytes:
    """Build a minimal docx with a single paragraph containing `text`."""
    doc = Document()
    doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── _fill_template_v2 ───────────────────────────────────────────────


@pytest.mark.unit
def test_fill_substitutes_placeholder():
    template = _docx_bytes_with_text(
        "The Debtor, [[debtor_name]], moves to extend.",
    )
    spec = _spec([_field("debtor_name")])
    resolved = [ResolvedTemplateValueV2(template_variable="debtor_name", value="Jane Doe")]
    filled_bytes, unresolved = _fill_template_v2(
        template_bytes=template, spec=spec, resolved_values=resolved,
    )
    # Re-open the filled docx and confirm the paragraph reads correctly.
    doc = Document(BytesIO(filled_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jane Doe" in text
    assert "[[debtor_name]]" not in text
    assert unresolved == []


@pytest.mark.unit
def test_fill_marks_unresolved_when_no_value():
    template = _docx_bytes_with_text("Hi [[name]] bye [[other]]")
    spec = _spec([_field("name"), _field("other")])
    resolved = [ResolvedTemplateValueV2(template_variable="name", value="Jane")]
    _, unresolved = _fill_template_v2(
        template_bytes=template, spec=spec, resolved_values=resolved,
    )
    assert "[[other]]" in unresolved


@pytest.mark.unit
def test_fill_marks_unresolved_when_empty_value():
    template = _docx_bytes_with_text("Hi [[name]]")
    spec = _spec([_field("name")])
    resolved = [ResolvedTemplateValueV2(template_variable="name", value="")]
    _, unresolved = _fill_template_v2(
        template_bytes=template, spec=spec, resolved_values=resolved,
    )
    assert "[[name]]" in unresolved


@pytest.mark.unit
def test_fill_template_v2_invokes_format_validator_when_flag_enabled():
    """`_fill_template_v2` runs the SYNC `_maybe_run_fill_validator`
    gate for logging — even though the LLM-driven autofix runs
    AFTER fill (awaited from `finalize_run_v2`). The sync gate
    surfaces drift signals in composer logs even when the async
    fixer never runs (e.g. flag off, finalize_run skipped, async
    layer disabled). Verifies both: gate is called AND the
    returned bytes propagate to the caller."""
    template = _docx_bytes_with_text("Hi [[name]]")
    spec = _spec([_field("name")])
    resolved = [ResolvedTemplateValueV2(template_variable="name", value="Jane")]

    with patch(
        "src.core.studio_v2.orchestration.finalizer.DocxTemplateService._maybe_run_fill_validator",
        return_value=b"GATE_OUTPUT_BYTES",
    ) as gate:
        filled, _ = _fill_template_v2(
            template_bytes=template, spec=spec, resolved_values=resolved,
        )

    assert gate.called
    assert filled == b"GATE_OUTPUT_BYTES"
    args, _kwargs = gate.call_args
    assert args[0] == template
    assert args[2] == {"name": "Jane"}


@pytest.mark.unit
@pytest.mark.asyncio
async def test_finalize_run_v2_invokes_async_autofix_after_fill():
    """`finalize_run_v2` must await `maybe_autofix_fill_async` AFTER
    `_fill_template_v2` returns, so the LLM-driven Tier 2 fixer runs
    in the same event loop the caller is in (avoids the httpx
    connection-pool error caused by thread-bridged event loops)."""
    template = _docx_bytes_with_text("Hi [[name]]")
    spec = _spec([_field("name")])
    resolved = [ResolvedTemplateValueV2(template_variable="name", value="Jane")]

    fake_url = "https://r2/test.docx"
    fake_async = AsyncMock(return_value=b"AUTOFIXED_BYTES")

    with patch(
        "src.core.common.documents.docx_template.DocxTemplateService.maybe_autofix_fill_async",
        new=fake_async,
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.upload_file",
        new=AsyncMock(),
    ) as upload, patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.get_presigned_url",
        new=AsyncMock(return_value=fake_url),
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.UserInputHealAgentV2.heal_resolved_values",
        new=AsyncMock(return_value=resolved),
    ):
        result = await finalize_run_v2(
            template_id=_TEMPLATE_UUID,
            case_id="case-1",
            resource_key="case-1",
            spec=spec,
            all_resolved=resolved,
            output_prefix="dry_run/x",
            template_bytes=template,
        )

    fake_async.assert_awaited_once()
    # The autofix output bytes should have been the ones uploaded.
    upload.assert_awaited_once()
    upload_kwargs = upload.call_args.kwargs
    assert upload_kwargs["file_content"] == b"AUTOFIXED_BYTES"
    # And the FinalizedRunV2 should carry those same bytes.
    assert result.filled_bytes == b"AUTOFIXED_BYTES"
    assert result.generated_doc_url == fake_url


@pytest.mark.unit
@pytest.mark.asyncio
async def test_finalize_run_v2_invokes_grammar_autofix_after_format_autofix():
    """`finalize_run_v2` must await `maybe_autofix_grammar_async`
    AFTER `maybe_autofix_fill_async`. The grammar fixer sees the
    cleaned-layout bytes the format fixer produced, so its agreement
    pass operates on consistent paragraph shapes. Output bytes from
    the grammar fixer are what land in R2 + FinalizedRunV2."""
    template = _docx_bytes_with_text("Hi [[name]]")
    spec = _spec([_field("name")])
    resolved = [ResolvedTemplateValueV2(template_variable="name", value="Jane")]

    fake_url = "https://r2/test.docx"
    fake_format = AsyncMock(return_value=b"FORMAT_FIXED_BYTES")
    fake_grammar = AsyncMock(return_value=(b"GRAMMAR_FIXED_BYTES", []))

    with patch(
        "src.core.common.documents.docx_template.DocxTemplateService.maybe_autofix_fill_async",
        new=fake_format,
    ), patch(
        "src.core.common.documents.docx_template.DocxTemplateService.maybe_autofix_grammar_async",
        new=fake_grammar,
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.upload_file",
        new=AsyncMock(),
    ) as upload, patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.get_presigned_url",
        new=AsyncMock(return_value=fake_url),
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.UserInputHealAgentV2.heal_resolved_values",
        new=AsyncMock(return_value=resolved),
    ):
        result = await finalize_run_v2(
            template_id=_TEMPLATE_UUID,
            case_id="case-1",
            resource_key="case-1",
            spec=spec,
            all_resolved=resolved,
            output_prefix="dry_run/x",
            template_bytes=template,
        )

    fake_format.assert_awaited_once()
    fake_grammar.assert_awaited_once()
    # The grammar fixer must be called with the format fixer's output,
    # NOT the raw fill bytes — confirms ordering.
    grammar_kwargs = fake_grammar.call_args.kwargs
    assert grammar_kwargs["filled_bytes"] == b"FORMAT_FIXED_BYTES"
    # The grammar fixer's output is what lands in R2 + FinalizedRunV2.
    upload_kwargs = upload.call_args.kwargs
    assert upload_kwargs["file_content"] == b"GRAMMAR_FIXED_BYTES"
    assert result.filled_bytes == b"GRAMMAR_FIXED_BYTES"
    # No repairs in this scenario since the mocked grammar fixer
    # returned an empty list.
    assert result.grammar_repairs == []


@pytest.mark.unit
@pytest.mark.asyncio
async def test_finalize_run_v2_surfaces_grammar_repairs_on_response():
    """When the grammar fixer applies swaps, the FinalizedRunV2 must
    carry them so the FE Resolution Log can display them."""
    from src.core.common.documents.template_grammar_fixer import (
        GrammarRepairRecord,
    )
    template = _docx_bytes_with_text("Hi [[name]]")
    spec = _spec([_field("name")])
    resolved = [ResolvedTemplateValueV2(template_variable="name", value="Jane")]

    fake_repairs = [
        GrammarRepairRecord(
            paragraph_index=7,
            original_word="Debtors",
            replacement_word="Debtor",
            occurrences=2,
            paragraph_preview="The Debtor, Jane, has filed",
            reason="single debtor",
        ),
    ]

    with patch(
        "src.core.common.documents.docx_template.DocxTemplateService.maybe_autofix_fill_async",
        new=AsyncMock(return_value=b"FORMAT_FIXED_BYTES"),
    ), patch(
        "src.core.common.documents.docx_template.DocxTemplateService.maybe_autofix_grammar_async",
        new=AsyncMock(return_value=(b"GRAMMAR_FIXED_BYTES", fake_repairs)),
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.upload_file",
        new=AsyncMock(),
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.get_presigned_url",
        new=AsyncMock(return_value="https://r2/test.docx"),
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.UserInputHealAgentV2.heal_resolved_values",
        new=AsyncMock(return_value=resolved),
    ):
        result = await finalize_run_v2(
            template_id=_TEMPLATE_UUID,
            case_id="case-1",
            resource_key="case-1",
            spec=spec,
            all_resolved=resolved,
            output_prefix="dry_run/x",
            template_bytes=template,
        )

    assert len(result.grammar_repairs) == 1
    repair = result.grammar_repairs[0]
    assert repair.paragraph_index == 7
    assert repair.original_word == "Debtors"
    assert repair.replacement_word == "Debtor"
    assert repair.occurrences == 2
    assert "Jane" in repair.paragraph_preview
    assert repair.reason == "single debtor"


# ─── _build_warnings ─────────────────────────────────────────────────


@pytest.mark.unit
def test_warnings_include_unresolved_placeholders():
    warnings = _build_warnings(
        resolved_values=[],
        unresolved=["[[a]]", "[[b]]"],
    )
    assert any("[[a]]" in w for w in warnings)
    assert any("[[b]]" in w for w in warnings)


@pytest.mark.unit
def test_warnings_include_low_confidence_rows():
    rows = [
        ResolvedTemplateValueV2(template_variable="x", value="v", confidence="high"),
        ResolvedTemplateValueV2(
            template_variable="y", value="z", confidence="low",
            note="LLM unsure",
        ),
    ]
    warnings = _build_warnings(rows, unresolved=[])
    assert not any("'x'" in w for w in warnings)
    assert any("'y'" in w and "Low-confidence" in w for w in warnings)


@pytest.mark.unit
def test_warnings_include_none_confidence_rows():
    rows = [
        ResolvedTemplateValueV2(
            template_variable="x", value="", confidence="none", note="not found",
        ),
    ]
    warnings = _build_warnings(rows, unresolved=[])
    assert any("None-confidence" in w for w in warnings)


# ─── finalize_run_v2 — orchestration ─────────────────────────────────


@pytest.mark.unit
@pytest.mark.asyncio
async def test_finalize_calls_heal_then_fill_then_upload_in_order():
    """End-to-end orchestration test: confirm date heal runs before
    prose heal, prose heal runs before fill, then upload + presigned
    URL are emitted."""
    template = _docx_bytes_with_text("Today is [[doc_date]] [[name]]")
    spec = _spec([_field("doc_date"), _field("name")])
    resolved = [
        ResolvedTemplateValueV2(template_variable="doc_date", value="01/21/2026"),
        ResolvedTemplateValueV2(template_variable="name", value="Jane Doe"),
    ]

    call_order: list[str] = []
    original_date_apply = DateHealingResolverV2.apply  # capture pre-patch reference

    def date_heal_apply(vs):
        call_order.append("date_heal")
        return original_date_apply(vs)

    async def prose_heal(*, template_bytes, template_fields, resolved_values):
        call_order.append("prose_heal")
        return resolved_values

    async def fake_upload(**_):
        call_order.append("upload")
        return "key"

    async def fake_presign(**_):
        call_order.append("presign")
        return "https://r2/signed.docx"

    with patch.object(
        DateHealingResolverV2, "apply", side_effect=date_heal_apply,
    ), patch.object(
        UserInputHealAgentV2, "heal_resolved_values",
        new=AsyncMock(side_effect=prose_heal),
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.upload_file",
        new=AsyncMock(side_effect=fake_upload),
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.get_presigned_url",
        new=AsyncMock(side_effect=fake_presign),
    ):
        result = await finalize_run_v2(
            template_id=_TEMPLATE_UUID,
            case_id="case-1",
            spec=spec,
            all_resolved=resolved,
            template_bytes=template,
        )
    # Critical ordering invariant: date_heal → prose_heal → upload.
    assert call_order.index("date_heal") < call_order.index("prose_heal")
    assert call_order.index("prose_heal") < call_order.index("upload")
    assert result.generated_doc_url == "https://r2/signed.docx"
    # Date was healed.
    by_name = {rv.template_variable: rv for rv in result.resolved_values}
    assert by_name["doc_date"].value == "January 21, 2026"
    # Docx was filled — open the (zipped) docx and confirm text.
    filled_doc = Document(BytesIO(result.filled_bytes))
    filled_text = "\n".join(p.text for p in filled_doc.paragraphs)
    assert "Jane Doe" in filled_text
    assert "January 21, 2026" in filled_text


@pytest.mark.unit
@pytest.mark.asyncio
async def test_finalize_downloads_template_bytes_when_not_supplied():
    """When no template_bytes passed, finalize fetches from R2."""
    template = _docx_bytes_with_text("Hi [[name]]")
    spec = _spec([_field("name")])
    resolved = [ResolvedTemplateValueV2(template_variable="name", value="Jane")]

    download_calls: list[dict] = []

    async def fake_download(**kwargs):
        download_calls.append(kwargs)
        return template

    with patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.download_file",
        new=AsyncMock(side_effect=fake_download),
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.upload_file",
        new=AsyncMock(return_value="key"),
    ), patch(
        "src.core.studio_v2.orchestration.finalizer.r2_service.get_presigned_url",
        new=AsyncMock(return_value="https://r2/x.docx"),
    ), patch.object(
        UserInputHealAgentV2, "heal_resolved_values",
        new=AsyncMock(side_effect=lambda **kw: kw["resolved_values"]),
    ):
        await finalize_run_v2(
            template_id=_TEMPLATE_UUID, case_id="case-1",
            spec=spec, all_resolved=resolved,
        )
    # Confirm download was called for the template_v2 prefix.
    assert len(download_calls) == 1
    assert download_calls[0]["filename"] == "template.docx"
    assert download_calls[0]["prefix"] == "template_v2"


# ─── _apply_web_enhance ──────────────────────────────────────────────


def _docx_with(text: str) -> bytes:
    """Build a minimal .docx so find_paragraph_containing has something
    to scan."""
    doc = Document()
    doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _field_with_params(name: str, params: WizardSourceParams):
    return TemplateFieldV2(
        id=_FIELD_UUID, template_id=_TEMPLATE_UUID,
        template_variable=name,
        template_property_marker=name.upper(),
        params=params,
    )


@pytest.mark.unit
@pytest.mark.asyncio
async def test_apply_web_enhance_skips_when_no_instruction():
    """Fields without a web_enhance_instruction pass through verbatim
    — no WebEnhanceAgentV2 call."""
    from src.core.studio_v2.orchestration.finalizer import _apply_web_enhance

    field = _field_with_params(
        "circuit",
        WizardSourceParams(
            source=SourceKind.GMAIL,
            extraction_prompt="x",
            web_enhance_instruction=None,
        ),
    )
    spec = TemplateSpecV2(template_id=_TEMPLATE_UUID, fields=[field])
    rv = ResolvedTemplateValueV2(
        template_variable="circuit", value="17", raw_context="", confidence="high",
    )
    run_mock = AsyncMock()
    with patch(
        "src.core.studio_v2.orchestration.finalizer.WebEnhanceAgentV2.run",
        new=run_mock,
    ):
        out = await _apply_web_enhance(
            spec=spec, resolved_values=[rv],
            template_bytes=_docx_with("Circuit [[circuit]]."),
        )
    run_mock.assert_not_awaited()
    assert out == [rv]


@pytest.mark.unit
@pytest.mark.asyncio
async def test_apply_web_enhance_skips_when_value_empty():
    """Empty resolved value short-circuits — nothing to enhance."""
    from src.core.studio_v2.orchestration.finalizer import _apply_web_enhance

    field = _field_with_params(
        "circuit",
        WizardSourceParams(
            source=SourceKind.GMAIL,
            extraction_prompt="x",
            web_enhance_instruction="confirm circuit",
        ),
    )
    spec = TemplateSpecV2(template_id=_TEMPLATE_UUID, fields=[field])
    rv = ResolvedTemplateValueV2(
        template_variable="circuit", value="", raw_context="", confidence="none",
    )
    run_mock = AsyncMock()
    with patch(
        "src.core.studio_v2.orchestration.finalizer.WebEnhanceAgentV2.run",
        new=run_mock,
    ):
        out = await _apply_web_enhance(
            spec=spec, resolved_values=[rv],
            template_bytes=_docx_with("Circuit [[circuit]]."),
        )
    run_mock.assert_not_awaited()
    assert out == [rv]


@pytest.mark.unit
@pytest.mark.asyncio
async def test_apply_web_enhance_swaps_value_and_appends_note():
    """Happy path: agent returns a polished value → resolved row gets
    the new value + the 'web-enhanced' note appended."""
    from src.core.studio_v2.orchestration.finalizer import _apply_web_enhance

    field = _field_with_params(
        "circuit",
        WizardSourceParams(
            source=SourceKind.GMAIL,
            extraction_prompt="x",
            web_enhance_instruction="confirm circuit",
            output_expectation="use ordinal",
        ),
    )
    spec = TemplateSpecV2(template_id=_TEMPLATE_UUID, fields=[field])
    rv = ResolvedTemplateValueV2(
        template_variable="circuit", value="17", raw_context="",
        confidence="high", note="picked by paralegal",
    )
    with patch(
        "src.core.studio_v2.orchestration.finalizer.WebEnhanceAgentV2.run",
        new=AsyncMock(return_value="17TH JUDICIAL CIRCUIT"),
    ) as run_mock:
        out = await _apply_web_enhance(
            spec=spec, resolved_values=[rv],
            template_bytes=_docx_with("Court for [[circuit]], Florida."),
        )
    run_mock.assert_awaited_once()
    kwargs = run_mock.await_args.kwargs
    assert kwargs["variable_name"] == "circuit"
    assert kwargs["current_value"] == "17"
    assert kwargs["web_enhance_instruction"] == "confirm circuit"
    assert kwargs["output_expectation"] == "use ordinal"
    assert kwargs["template_property_marker"] == "CIRCUIT"
    assert out[0].value == "17TH JUDICIAL CIRCUIT"
    assert "web-enhanced" in out[0].note
    assert "picked by paralegal" in out[0].note


@pytest.mark.unit
@pytest.mark.asyncio
async def test_apply_web_enhance_passes_through_when_agent_returns_unchanged():
    """Soft-fail: agent returns the same value → no swap, no note
    append, row passes through verbatim."""
    from src.core.studio_v2.orchestration.finalizer import _apply_web_enhance

    field = _field_with_params(
        "circuit",
        WizardSourceParams(
            source=SourceKind.GMAIL,
            extraction_prompt="x",
            web_enhance_instruction="confirm circuit",
        ),
    )
    spec = TemplateSpecV2(template_id=_TEMPLATE_UUID, fields=[field])
    rv = ResolvedTemplateValueV2(
        template_variable="circuit", value="17", raw_context="",
        confidence="high", note="picked by paralegal",
    )
    with patch(
        "src.core.studio_v2.orchestration.finalizer.WebEnhanceAgentV2.run",
        new=AsyncMock(return_value="17"),
    ):
        out = await _apply_web_enhance(
            spec=spec, resolved_values=[rv],
            template_bytes=_docx_with("Court for [[circuit]]."),
        )
    assert out == [rv]
    assert "web-enhanced" not in out[0].note


@pytest.mark.unit
@pytest.mark.asyncio
async def test_apply_web_enhance_handles_field_without_params():
    """Field rows with `params is None` (paralegal hasn't bound a
    source yet) skip enhancement without crashing."""
    from src.core.studio_v2.orchestration.finalizer import _apply_web_enhance

    field = TemplateFieldV2(
        id=_FIELD_UUID, template_id=_TEMPLATE_UUID,
        template_variable="x", template_property_marker="X", params=None,
    )
    spec = TemplateSpecV2(template_id=_TEMPLATE_UUID, fields=[field])
    rv = ResolvedTemplateValueV2(
        template_variable="x", value="value", raw_context="", confidence="high",
    )
    run_mock = AsyncMock()
    with patch(
        "src.core.studio_v2.orchestration.finalizer.WebEnhanceAgentV2.run",
        new=run_mock,
    ):
        out = await _apply_web_enhance(
            spec=spec, resolved_values=[rv],
            template_bytes=_docx_with("Body [[x]]."),
        )
    run_mock.assert_not_awaited()
    assert out == [rv]
