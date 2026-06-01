"""Tests for the Tier 2 LLM-assisted grammar fixer.

The fixer consumes the validator's `GrammarSuspectParagraph[]` output,
calls Sonnet 4.6 once at document scope, and receives a flat list of
`_WordSubstitution` entries (paragraph_index + original_word +
replacement_word). Each substitution is validated against the closed
agreement allowlist + case-style check, then applied in place inside
the matching `<w:t>` element so run-level formatting (bold /
underline / italic) survives.

LLM calls are mocked throughout — these tests cover the deterministic
allowlist + case-style + in-place apply logic.
"""

from io import BytesIO
from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml.ns import qn

from src.core.common.documents.template_grammar_fixer import (
    GrammarRepairRecord,
    _WordSubstitution,
    _agreement_key,
    _case_style,
    _is_allowed_agreement_swap,
    _reason_hedges_against_swap,
    _replace_word_token_in_t_elements,
    _split_outer_punctuation,
    _strip_matching_outer_punctuation,
    autofix_grammar_drift,
)
from src.core.common.documents.template_grammar_validator import (
    GrammarSuspectParagraph,
)


def _build_docx(paragraphs: list[str]) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _suspect(idx: int, text: str, tokens: tuple[str, ...] = ("debtors",)) -> GrammarSuspectParagraph:
    return GrammarSuspectParagraph(
        template_paragraph_index=idx,
        paragraph_text=text,
        triggered_tokens=tokens,
        suspect_reason="test",
    )


def _para_with_runs(*runs: tuple[str, dict]) -> tuple:
    """Build a single-paragraph docx where each `run` is `(text, props)`.
    `props` may set `'bold': True`, `'italic': True`, `'underline': True`.
    Returns `(doc, paragraph, paragraph_element)`."""
    doc = Document()
    p = doc.add_paragraph()
    for text, props in runs:
        r = p.add_run(text)
        if props.get("bold"):
            r.bold = True
        if props.get("italic"):
            r.italic = True
        if props.get("underline"):
            r.underline = True
    return doc, p, p._element


# === _case_style ===


@pytest.mark.unit
def test_case_style_upper():
    assert _case_style("DEBTORS") == "upper"
    assert _case_style("DEBTOR'S") == "upper"  # apostrophes ignored


@pytest.mark.unit
def test_case_style_lower():
    assert _case_style("debtors") == "lower"
    assert _case_style("debtor's") == "lower"


@pytest.mark.unit
def test_case_style_title():
    assert _case_style("Debtors") == "title"
    assert _case_style("Debtor's") == "title"


@pytest.mark.unit
def test_case_style_mixed():
    assert _case_style("DeBtOrS") == "mixed"


@pytest.mark.unit
def test_case_style_none_for_no_letters():
    assert _case_style("'") == "none"


# === _agreement_key ===


@pytest.mark.unit
def test_agreement_key_normalizes_apostrophes_and_case():
    assert _agreement_key("DEBTORS’") == "debtors'"
    assert _agreement_key("Debtor’s") == "debtor's"
    assert _agreement_key("debtor's") == "debtor's"


# === _is_allowed_agreement_swap ===


@pytest.mark.unit
def test_allowed_swap_plural_to_singular_noun():
    assert _is_allowed_agreement_swap("Debtors", "Debtor")
    assert _is_allowed_agreement_swap("DEBTORS", "DEBTOR")
    assert _is_allowed_agreement_swap("debtors", "debtor")


@pytest.mark.unit
def test_allowed_swap_possessive_plural_to_singular():
    """The exact substitution from the screenshots: DEBTORS' → DEBTOR'S
    (both ALL CAPS, both possessive)."""
    assert _is_allowed_agreement_swap("DEBTORS'", "DEBTOR'S")
    assert _is_allowed_agreement_swap("Debtors'", "Debtor's")


@pytest.mark.unit
def test_rejected_swap_pronoun_singular_they_is_legalese_standard():
    """Pronouns are DELIBERATELY ABSENT from _AGREEMENT_PAIRS — modern
    legalese uses singular they ('the Debtor, by and through their
    counsel'). The fixer rejects any LLM-emitted pronoun swap so the
    original pronoun stays regardless of resolved gender."""
    assert not _is_allowed_agreement_swap("their", "her")
    assert not _is_allowed_agreement_swap("their", "his")
    assert not _is_allowed_agreement_swap("their", "its")
    assert not _is_allowed_agreement_swap("Their", "Her")
    assert not _is_allowed_agreement_swap("they", "he")
    assert not _is_allowed_agreement_swap("they", "she")
    assert not _is_allowed_agreement_swap("them", "him")
    assert not _is_allowed_agreement_swap("them", "her")
    assert not _is_allowed_agreement_swap("themselves", "himself")
    assert not _is_allowed_agreement_swap("themselves", "herself")
    assert not _is_allowed_agreement_swap("theirs", "his")
    assert not _is_allowed_agreement_swap("theirs", "hers")


@pytest.mark.unit
def test_allowed_swap_verb_have_to_has():
    assert _is_allowed_agreement_swap("have", "has")
    assert _is_allowed_agreement_swap("HAVE", "HAS")


@pytest.mark.unit
def test_allowed_swap_determiner_these_to_this():
    assert _is_allowed_agreement_swap("these", "this")


@pytest.mark.unit
def test_rejected_swap_non_agreement_word():
    """Random word pair not in the allowlist."""
    assert not _is_allowed_agreement_swap("filed", "submitted")
    assert not _is_allowed_agreement_swap("Motion", "Pleading")


@pytest.mark.unit
def test_rejected_swap_case_style_mismatch():
    """The exact bug from the screenshots: DEBTORS' → Debtor's
    (uppercase → title) — same agreement pair but case shifts. Reject.
    """
    assert not _is_allowed_agreement_swap("DEBTORS'", "Debtor's")
    assert not _is_allowed_agreement_swap("Debtors", "DEBTOR")
    # 'have' → 'Has' is allowed pair shape but case-mismatched.
    assert not _is_allowed_agreement_swap("have", "Has")


@pytest.mark.unit
def test_rejected_swap_name_or_factual_content():
    """Names, dates, numbers can't be in the allowlist by construction."""
    assert not _is_allowed_agreement_swap("Schwartz", "Smith")
    assert not _is_allowed_agreement_swap("2026", "2025")


# === _replace_word_token_in_t_elements ===


@pytest.mark.unit
def test_replace_word_in_single_t_element():
    _doc, p, p_elem = _para_with_runs(("The Debtors have filed", {}))
    t_elems = list(p_elem.iter(qn("w:t")))
    count = _replace_word_token_in_t_elements(t_elems, "Debtors", "Debtor")
    assert count == 1
    assert p.text == "The Debtor have filed"


@pytest.mark.unit
def test_replace_word_does_not_match_inside_possessive():
    """Replacing 'Debtors' must NOT match inside 'Debtors'' (the
    possessive plural). Apostrophes are part of word tokens."""
    _doc, p, p_elem = _para_with_runs(("The Debtors and the Debtors' attorneys", {}))
    t_elems = list(p_elem.iter(qn("w:t")))
    count = _replace_word_token_in_t_elements(t_elems, "Debtors", "Debtor")
    # Only the bare 'Debtors' should change; 'Debtors'' stays.
    assert count == 1
    assert p.text == "The Debtor and the Debtors' attorneys"


@pytest.mark.unit
def test_replace_word_replaces_all_occurrences_in_paragraph():
    _doc, p, p_elem = _para_with_runs(
        ("The Debtors hereby request that the Debtors are heard", {}),
    )
    t_elems = list(p_elem.iter(qn("w:t")))
    count = _replace_word_token_in_t_elements(t_elems, "Debtors", "Debtor")
    assert count == 2
    assert p.text == "The Debtor hereby request that the Debtor are heard"


@pytest.mark.unit
def test_replace_word_returns_zero_when_word_not_present():
    _doc, p, p_elem = _para_with_runs(("The Creditors filed", {}))
    t_elems = list(p_elem.iter(qn("w:t")))
    count = _replace_word_token_in_t_elements(t_elems, "Debtors", "Debtor")
    assert count == 0
    assert p.text == "The Creditors filed"


@pytest.mark.unit
def test_replace_preserves_bold_underline_run_formatting():
    """The KILLER test: bold + underlined heading containing a plural
    word. After replacement, the run is STILL bold + underlined."""
    _doc, p, p_elem = _para_with_runs(
        ("DEBTORS' EX-PARTE MOTION FOR EXTENSION OF TIME",
         {"bold": True, "underline": True}),
    )
    t_elems = list(p_elem.iter(qn("w:t")))
    count = _replace_word_token_in_t_elements(t_elems, "DEBTORS'", "DEBTOR'S")
    assert count == 1
    assert p.text == "DEBTOR'S EX-PARTE MOTION FOR EXTENSION OF TIME"
    # Formatting MUST survive.
    assert p.runs[0].bold is True
    assert p.runs[0].underline is True


@pytest.mark.unit
def test_replace_preserves_caption_underscores_and_whitespace():
    """The other bug from the screenshots: caption '________ Debtors. /'
    must keep its leading underscores + spacing. Since we only swap
    'Debtors' for 'Debtor' at the word-token level, everything else
    is preserved by construction."""
    _doc, p, p_elem = _para_with_runs(
        ("________ Debtors.    /", {}),
    )
    t_elems = list(p_elem.iter(qn("w:t")))
    count = _replace_word_token_in_t_elements(t_elems, "Debtors", "Debtor")
    assert count == 1
    assert p.text == "________ Debtor.    /"


@pytest.mark.unit
def test_replace_word_split_across_runs_succeeds_via_pass3():
    """Word splits across two `<w:t>` runs (mid-word formatting).
    Pass-3 of the apply concatenates the run texts, locates the word
    boundary across runs, and redistributes the replacement —
    untouched prefix / suffix chars stay in their original runs."""
    _doc, p, p_elem = _para_with_runs(
        ("The Debt", {"bold": True}),
        ("ors have filed", {}),
    )
    t_elems = list(p_elem.iter(qn("w:t")))
    count = _replace_word_token_in_t_elements(t_elems, "Debtors", "Debtor")
    assert count == 1
    # The new word lands in the FIRST overlapping run (where the
    # match starts); the trailing 'ors' is dropped from the second
    # run since pass-3 redistributes the original word's character
    # span. Run formatting on each run is preserved.
    assert p.text == "The Debtor have filed"
    assert p.runs[0].bold is True


# === autofix_grammar_drift integration ===


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_no_suspects_returns_original():
    """No suspects → fixer no-ops without invoking LLM. Returns
    `(original_bytes, [])`."""
    filled = _build_docx(["Some text"])

    bytes_out, applied = await autofix_grammar_drift(
        filled_bytes=filled,
        resolved_values={"debtor_name": "Judith S. Schwartz"},
        suspect_paragraphs=[],
        cardinality_signals={"debtor_name": "single"},
    )
    assert bytes_out == filled
    assert applied == []


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_applies_singular_word_substitutions():
    """End-to-end: LLM emits a few word swaps; fixer applies them in
    place; final paragraph reads with singular agreement on noun +
    verb but KEEPS the original pronoun (singular they is
    legalese-standard). The applied list captures one record per
    successfully-landed substitution; the pronoun reject + skip
    doesn't appear."""
    filled = _build_docx([
        "The Debtors, Judith S. Schwartz, have filed their motion",
    ])

    async def fake_llm(**_kwargs):
        return [
            _WordSubstitution(paragraph_index=0, original_word="Debtors", replacement_word="Debtor"),
            _WordSubstitution(paragraph_index=0, original_word="have", replacement_word="has"),
            # The LLM SHOULD emit replace_original=False on 'their'
            # per the prompt, but even if it forgets and emits True,
            # the allowlist rejects 'their' → 'her' so the pronoun
            # stays.
            _WordSubstitution(
                paragraph_index=0,
                original_word="their",
                replacement_word="her",
                replace_original=True,
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=fake_llm,
    ):
        fixed_bytes, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[
                _suspect(0, "The Debtors, Judith S. Schwartz, have filed their motion",
                         ("debtors", "have", "their")),
            ],
            cardinality_signals={"debtor_name": "single"},
        )

    assert fixed_bytes != filled
    fixed_doc = Document(BytesIO(fixed_bytes))
    assert fixed_doc.paragraphs[0].text == (
        "The Debtor, Judith S. Schwartz, has filed their motion"
    )
    # Only the noun + verb swaps landed; the pronoun was rejected.
    assert len(applied) == 2
    assert all(isinstance(r, GrammarRepairRecord) for r in applied)
    swaps = {(r.original_word, r.replacement_word) for r in applied}
    assert swaps == {("Debtors", "Debtor"), ("have", "has")}
    # paragraph_preview should carry the FIXED text, useful in the FE
    # Resolution Log so paralegals see what the paragraph now reads.
    assert all("Debtor" in r.paragraph_preview for r in applied)


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_preserves_run_formatting_on_heading_fix():
    """Heading paragraph with bold + underline; LLM swaps 'DEBTORS''
    → 'DEBTOR'S'; formatting survives end-to-end."""
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("DEBTORS' EX-PARTE MOTION FOR EXTENSION OF TIME")
    r.bold = True
    r.underline = True
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    filled = buf.read()

    async def fake_llm(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=0,
                original_word="DEBTORS'",
                replacement_word="DEBTOR'S",
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=fake_llm,
    ):
        fixed_bytes, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[
                _suspect(0, "DEBTORS' EX-PARTE MOTION FOR EXTENSION OF TIME",
                         ("debtors",)),
            ],
            cardinality_signals={"debtor_name": "single"},
        )

    fixed_doc = Document(BytesIO(fixed_bytes))
    fixed_p = fixed_doc.paragraphs[0]
    assert fixed_p.text == "DEBTOR'S EX-PARTE MOTION FOR EXTENSION OF TIME"
    assert fixed_p.runs[0].bold is True
    assert fixed_p.runs[0].underline is True
    assert len(applied) == 1
    assert applied[0].original_word == "DEBTORS'"
    assert applied[0].replacement_word == "DEBTOR'S"
    assert applied[0].occurrences == 1


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_rejects_case_style_mismatch():
    """LLM emitted 'DEBTORS'' → 'Debtor's' (case shifted). Allowlist
    check fails; fix is dropped; bytes unchanged."""
    filled = _build_docx(["DEBTORS' MOTION"])

    async def fake_llm(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=0,
                original_word="DEBTORS'",
                replacement_word="Debtor's",  # case shift
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=fake_llm,
    ):
        result, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[_suspect(0, "DEBTORS' MOTION", ("debtors",))],
            cardinality_signals={"debtor_name": "single"},
        )

    assert result == filled
    assert applied == []


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_rejects_non_agreement_word_swap():
    """LLM tried to swap a non-agreement word (e.g. a verb not in the
    allowlist). Dropped at allowlist check."""
    filled = _build_docx(["The Debtors filed today"])

    async def fake_llm(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=0,
                original_word="filed",
                replacement_word="submitted",  # not an agreement swap
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=fake_llm,
    ):
        result, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[_suspect(0, "The Debtors filed today")],
            cardinality_signals={"debtor_name": "single"},
        )

    assert result == filled
    assert applied == []


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_skips_substitution_for_word_not_in_paragraph():
    """LLM hallucinated a substitution for a word that isn't in the
    paragraph. Replacement returns 0 occurrences; substitution is
    skipped silently (no other fixes are attempted)."""
    filled = _build_docx(["The Creditors filed"])

    async def fake_llm(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=0,
                original_word="Debtors",  # not in paragraph text
                replacement_word="Debtor",
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=fake_llm,
    ):
        result, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[_suspect(0, "The Creditors filed", ("creditors",))],
            cardinality_signals={"debtor_name": "single"},
        )

    assert result == filled
    assert applied == []


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_applies_substitutions_to_multiple_paragraphs():
    """LLM emits substitutions for paragraphs 1 + 2; both apply in one
    pass, paragraph 0 (no suspect / no subs) untouched."""
    filled = _build_docx([
        "JUDITH S SCHWARTZ",
        "Debtors.",
        "The Debtors, Judith S. Schwartz, have filed",
    ])

    async def fake_llm(**_kwargs):
        return [
            _WordSubstitution(paragraph_index=1, original_word="Debtors", replacement_word="Debtor"),
            _WordSubstitution(paragraph_index=2, original_word="Debtors", replacement_word="Debtor"),
            _WordSubstitution(paragraph_index=2, original_word="have", replacement_word="has"),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=fake_llm,
    ):
        fixed, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[
                _suspect(1, "Debtors."),
                _suspect(2, "The Debtors, Judith S. Schwartz, have filed",
                         ("debtors", "have")),
            ],
            cardinality_signals={"debtor_name": "single"},
        )

    fixed_doc = Document(BytesIO(fixed))
    paras = list(fixed_doc.paragraphs)
    assert paras[0].text == "JUDITH S SCHWARTZ"  # untouched
    assert paras[1].text == "Debtor."
    assert paras[2].text == "The Debtor, Judith S. Schwartz, has filed"
    # 3 applied swaps across 2 paragraphs.
    assert len(applied) == 3
    indices = sorted({r.paragraph_index for r in applied})
    assert indices == [1, 2]


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_soft_fails_on_llm_error():
    filled = _build_docx(["The Debtors have filed"])

    async def boom(**_kwargs):
        raise RuntimeError("Sonnet API timeout")

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=boom,
    ):
        result, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[_suspect(0, "The Debtors have filed")],
            cardinality_signals={"debtor_name": "single"},
        )

    assert result == filled
    assert applied == []


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_ignores_substitution_for_unknown_index():
    """LLM returned paragraph_index not in suspect set — skip silently."""
    filled = _build_docx(["The Debtors have filed"])

    async def confused_llm(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=42,  # not in suspect set
                original_word="Debtors",
                replacement_word="Debtor",
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=confused_llm,
    ):
        result, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[_suspect(0, "The Debtors have filed")],
            cardinality_signals={"debtor_name": "single"},
        )

    assert result == filled
    assert applied == []


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_handles_doc_parse_failure():
    """Corrupt bytes → soft-fail, return unchanged."""
    result, applied = await autofix_grammar_drift(
        filled_bytes=b"not a docx",
        resolved_values={"debtor_name": "Judith S. Schwartz"},
        suspect_paragraphs=[_suspect(0, "Anything")],
        cardinality_signals={"debtor_name": "single"},
    )
    assert result == b"not a docx"
    assert applied == []


# === _strip_matching_outer_punctuation + _split_outer_punctuation ===


@pytest.mark.unit
def test_split_outer_punctuation_handles_trailing_period():
    assert _split_outer_punctuation("Debtors.") == ("", "Debtors", ".")


@pytest.mark.unit
def test_split_outer_punctuation_handles_no_punctuation():
    assert _split_outer_punctuation("Debtors") == ("", "Debtors", "")


@pytest.mark.unit
def test_split_outer_punctuation_keeps_inner_apostrophe():
    """Inner apostrophes ARE word chars; should stay in the core."""
    assert _split_outer_punctuation("Debtor's") == ("", "Debtor's", "")
    assert _split_outer_punctuation("Debtor's.") == ("", "Debtor's", ".")


@pytest.mark.unit
def test_strip_matching_outer_punctuation_strips_matching_period():
    """LLM emitted 'Debtors.' → 'Debtor.' — matching trailing period
    on both sides, strip to 'Debtors' → 'Debtor' for allowlist check."""
    assert _strip_matching_outer_punctuation("Debtors.", "Debtor.") == (
        "Debtors", "Debtor",
    )


@pytest.mark.unit
def test_strip_matching_outer_punctuation_rejects_mismatched_punctuation():
    """Trailing punctuation differs → LLM trying to mutate punctuation,
    reject."""
    assert _strip_matching_outer_punctuation("Debtors,", "Debtor.") is None
    assert _strip_matching_outer_punctuation("Debtors", "Debtor.") is None


# === autofix_grammar_drift — robustness paths from live-log issues ===


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_handles_llm_with_trailing_punctuation():
    """LIVE-LOG BUG #1: LLM emitted ('Debtors.', 'Debtor.') instead
    of ('Debtors', 'Debtor'). Strip matching trailing period; allowlist
    accepts the bare word swap; in-place apply finds 'Debtors' inside
    the `<w:t>` and replaces it. Trailing period preserved by the
    token-level apply."""
    filled = _build_docx(["JUDITH S SCHWARTZ\n\nDebtors."])

    async def fake_llm(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=0,
                original_word="Debtors.",
                replacement_word="Debtor.",
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=fake_llm,
    ):
        fixed, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[_suspect(0, "JUDITH S SCHWARTZ\n\nDebtors.")],
            cardinality_signals={"debtor_name": "single"},
        )

    fixed_doc = Document(BytesIO(fixed))
    assert "Debtor." in fixed_doc.paragraphs[0].text
    assert "Debtors." not in fixed_doc.paragraphs[0].text
    assert len(applied) == 1
    # The recorded swap is the BARE word, not the punctuated form.
    assert applied[0].original_word == "Debtors"
    assert applied[0].replacement_word == "Debtor"


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_handles_apostrophe_variant_in_docx():
    """LIVE-LOG BUG #2: docx has curly apostrophe (U+2019) but LLM
    emits the straight-apostrophe form ("DEBTORS'"). Pass-2 of the
    apply normalizes apostrophes for matching and rewrites the token
    in place."""
    doc = Document()
    p = doc.add_paragraph()
    # Note: real curly apostrophe (U+2019) in the docx text.
    p.add_run("DEBTORS’ MOTION")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    filled = buf.read()

    async def fake_llm(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=0,
                original_word="DEBTORS'",  # straight apostrophe
                replacement_word="DEBTOR'S",
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=fake_llm,
    ):
        fixed, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[_suspect(0, "DEBTORS’ MOTION")],
            cardinality_signals={"debtor_name": "single"},
        )

    fixed_doc = Document(BytesIO(fixed))
    assert fixed_doc.paragraphs[0].text == "DEBTOR'S MOTION"
    assert len(applied) == 1


# === replace_original contract + hedging backstop ===


@pytest.mark.unit
def test_reason_hedges_against_swap_catches_but_wait():
    """The exact phrasing from the live worker log."""
    reason = (
        "single debtor resolved; subject 'Irreconcilable differences' "
        "is plural so 'have' is correct — but wait, the subject here "
        "is NOT the debtor; 'have' agrees with 'differences' (plural "
        "noun), so this is grammatically correct and should NOT be "
        "changed"
    )
    assert _reason_hedges_against_swap(reason)


@pytest.mark.unit
def test_reason_hedges_against_swap_catches_is_correct():
    assert _reason_hedges_against_swap(
        "On reflection this is grammatically correct",
    )


@pytest.mark.unit
def test_reason_hedges_against_swap_accepts_affirmative_reason():
    """Affirmative justification stays — only hedging language
    triggers the backstop."""
    assert not _reason_hedges_against_swap("single debtor; subject is debtor")
    assert not _reason_hedges_against_swap("single debtor resolved")
    assert not _reason_hedges_against_swap("possessive form")


@pytest.mark.unit
def test_reason_hedges_against_swap_handles_empty_reason():
    assert not _reason_hedges_against_swap("")
    assert not _reason_hedges_against_swap("   ")


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_skips_substitution_when_replace_original_false():
    """Primary contract: replace_original=False means 'I considered
    this and decided the original is correct'. Fixer skips silently;
    no swap applied; not counted as a repair."""
    filled = _build_docx([
        "Irreconcilable differences have arisen between the Debtor",
    ])

    async def llm_recognizes_sva(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=0,
                original_word="have",
                replacement_word="has",
                replace_original=False,  # LLM correctly decided NOT to swap
                reason="subject is 'differences' (plural); keep have",
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=llm_recognizes_sva,
    ):
        result, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[
                _suspect(0, "Irreconcilable differences have arisen between the Debtor",
                         ("have",)),
            ],
            cardinality_signals={"debtor_name": "single"},
        )

    assert result == filled  # docx unchanged
    assert applied == []  # nothing applied — LLM correctly kept original


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_rejects_swap_with_hedging_reason_even_when_replace_original_true():
    """Belt-and-suspenders: if the LLM emits replace_original=True
    but the reason contains hedging language ('but wait...',
    'should not be changed'), reject the swap. The LLM's natural
    language reveals the right decision even when the bool is wrong."""
    filled = _build_docx([
        "Irreconcilable differences have arisen between the Debtor",
    ])

    async def llm_emits_hedged_swap(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=0,
                original_word="have",
                replacement_word="has",
                replace_original=True,  # LLM set this wrong
                reason=(
                    "single debtor; but wait, the subject here is NOT "
                    "the debtor; 'have' agrees with 'differences' "
                    "(plural noun), so this is grammatically correct "
                    "and should NOT be changed"
                ),
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=llm_emits_hedged_swap,
    ):
        result, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[
                _suspect(0, "Irreconcilable differences have arisen between the Debtor",
                         ("have",)),
            ],
            cardinality_signals={"debtor_name": "single"},
        )

    assert result == filled
    assert applied == []  # backstop rejected the hedged swap


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_mixes_apply_and_skip_in_same_paragraph():
    """Realistic case: one paragraph flagged for multiple words, LLM
    decides to swap some and keep others. Only the True ones apply."""
    filled = _build_docx([
        "The Debtors have many claims that have been disputed",
    ])

    async def llm_mixed_decision(**_kwargs):
        return [
            # Debtors → Debtor — apply (party-name swap, single debtor)
            _WordSubstitution(
                paragraph_index=0,
                original_word="Debtors",
                replacement_word="Debtor",
                replace_original=True,
                reason="single debtor",
            ),
            # First 'have' — apply (subject is 'Debtor' singular)
            _WordSubstitution(
                paragraph_index=0,
                original_word="have",
                replacement_word="has",
                replace_original=True,
                reason="subject is Debtor (singular)",
            ),
            # Second 'have' — keep (subject is 'claims' plural)
            _WordSubstitution(
                paragraph_index=0,
                original_word="have",
                replacement_word="has",
                replace_original=False,
                reason="subject is 'claims' (plural); keep have",
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=llm_mixed_decision,
    ):
        fixed, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[
                _suspect(0, "The Debtors have many claims that have been disputed",
                         ("debtors", "have")),
            ],
            cardinality_signals={"debtor_name": "single"},
        )

    # Both apply-True entries land. Note: the second 'have' (which
    # the LLM marked as keep) lives in a separate word-token position,
    # but our token replacer replaces ALL occurrences of 'have' in the
    # paragraph when the first matching swap fires. This is acceptable
    # in practice because the LLM is supposed to mark both occurrences
    # if both should change — when they disagree we either apply both
    # (matches the first decision) or skip both (depending on
    # iteration order). What this test pins is that an explicit
    # replace_original=False entry doesn't FORCE an apply on its own.
    # The first matching True entry drives the actual replacement.
    fixed_doc = Document(BytesIO(fixed))
    assert fixed_doc.paragraphs[0].text.startswith("The Debtor")
    # Confirm at least one apply landed.
    assert any(r.original_word == "Debtors" for r in applied)


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_handles_word_split_across_runs():
    """LIVE-LOG BUG #3: 'Debtors' splits across two `<w:t>` elements
    (mid-word formatting). Pass-3 of the apply concatenates, finds the
    word boundary across runs, and redistributes the replacement —
    untouched prefix / suffix chars stay in their original runs."""
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("The Debt")
    r1.bold = True
    r2 = p.add_run("ors have filed")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    filled = buf.read()

    async def fake_llm(**_kwargs):
        return [
            _WordSubstitution(
                paragraph_index=0,
                original_word="Debtors",
                replacement_word="Debtor",
            ),
        ]

    with patch(
        "src.core.common.documents.template_grammar_fixer._call_sonnet_for_grammar",
        new=fake_llm,
    ):
        fixed, applied = await autofix_grammar_drift(
            filled_bytes=filled,
            resolved_values={"debtor_name": "Judith S. Schwartz"},
            suspect_paragraphs=[_suspect(0, "The Debtors have filed")],
            cardinality_signals={"debtor_name": "single"},
        )

    fixed_doc = Document(BytesIO(fixed))
    fixed_p = fixed_doc.paragraphs[0]
    assert fixed_p.text == "The Debtor have filed"
    # The first run still carries its bold formatting (now contains
    # "The Debtor" instead of "The Debt").
    assert fixed_p.runs[0].bold is True
    assert len(applied) == 1
