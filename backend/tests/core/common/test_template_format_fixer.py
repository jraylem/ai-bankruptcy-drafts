"""Tests for the Tier 2 LLM-assisted template format fixer.

The fixer takes the validator's `DriftedParagraph[]` output, calls
Sonnet 4.6 per drift, content-equality-checks the response, and
rebuilds the paragraph XML deterministically. All LLM calls are
mocked — these tests cover the surrounding deterministic logic.
"""

from io import BytesIO
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document

from src.core.agents.types.spec import TemplateVariable
from src.core.common.documents.template_format_fixer import (
    _content_equivalent,
    _decode_escape_sequences,
    _normalize_for_equality,
    _apply_corrected_text,
    autofix_fill_drift,
)
from src.core.common.documents.template_format_validator import (
    DriftedParagraph,
    _paragraph_text_with_breaks,
    validate_fill_format,
)


def _build_docx(paragraphs: list[str]) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# === _content_equivalent guardrail ===


@pytest.mark.unit
def test_content_equivalent_accepts_pure_whitespace_change():
    """Whitespace flavors (` ` vs `\\n` vs `\\t`) collapse to the same
    canonical form. Words remain. Equivalence holds."""
    assert _content_equivalent(
        "Robin R. Weiner\tauto-forward-ecf@ch13weiner.com",
        "Robin R. Weiner\nauto-forward-ecf@ch13weiner.com",
    )


@pytest.mark.unit
def test_content_equivalent_accepts_tab_added():
    assert _content_equivalent("In re Case No. 26-1", "In re\tCase No. 26-1")


@pytest.mark.unit
def test_content_equivalent_rejects_content_change():
    """LLM tried to change 'Robin' to 'Robyn' — content drift, reject."""
    assert not _content_equivalent(
        "Robin R. Weiner",
        "Robyn R. Weiner",
    )


@pytest.mark.unit
def test_content_equivalent_rejects_added_word():
    """LLM added a stray word — reject."""
    assert not _content_equivalent(
        "Trustee for the Debtor",
        "Trustee for the Honorable Debtor",
    )


@pytest.mark.unit
def test_content_equivalent_rejects_dropped_word():
    """LLM dropped a word — reject."""
    assert not _content_equivalent(
        "Trustee for the Debtor",
        "Trustee for Debtor",
    )


@pytest.mark.unit
def test_normalize_for_equality_strips_whitespace_punctuation_lowercase():
    assert _normalize_for_equality("Robin R. Weiner") == "robinrweiner"
    assert _normalize_for_equality("  Robin R. Weiner  ") == "robinrweiner"
    assert _normalize_for_equality("ROBIN R. WEINER") == "robinrweiner"


# === _decode_escape_sequences ===


@pytest.mark.unit
def test_decode_escape_sequences_converts_literal_backslash_n_to_newline():
    """Claude sometimes echoes our prompt's literal `\\n` (two chars:
    backslash + n) instead of using a real newline. The defensive
    decoder restores the intent."""
    literal = "Robin\\nemail@x.com"  # actually 'R','o','b','i','n','\\','n', ...
    assert _decode_escape_sequences(literal) == "Robin\nemail@x.com"


@pytest.mark.unit
def test_decode_escape_sequences_converts_literal_backslash_t_to_tab():
    literal = "In re\\tCase No. 26-1"
    assert _decode_escape_sequences(literal) == "In re\tCase No. 26-1"


@pytest.mark.unit
def test_decode_escape_sequences_idempotent_on_real_newlines():
    real = "Robin\nemail@x.com"  # real newline
    assert _decode_escape_sequences(real) == real


@pytest.mark.unit
def test_decode_escape_sequences_handles_mixed_real_and_literal():
    """If the LLM mixes real `\\n` (proper) and literal `\\n` (improper)
    in the same string, both end up as real newlines after decode."""
    mixed = "First\nSecond\\nThird"
    assert _decode_escape_sequences(mixed) == "First\nSecond\nThird"


# === _apply_corrected_text rebuild ===


@pytest.mark.unit
def test_apply_corrected_text_rebuilds_with_soft_break_for_newline():
    """`\\n` in corrected text → `<w:br/>` element in the rebuilt run."""
    doc = Document()
    p = doc.add_paragraph("wrong content")
    _apply_corrected_text(p._element, "Line one\nLine two")
    text = _paragraph_text_with_breaks(p)
    assert text == "Line one\nLine two"


@pytest.mark.unit
def test_apply_corrected_text_rebuilds_with_tab():
    """`\\t` in corrected text → `<w:tab/>` element."""
    doc = Document()
    p = doc.add_paragraph("wrong content")
    _apply_corrected_text(p._element, "In re\tCase No. 26-1")
    text = _paragraph_text_with_breaks(p)
    assert text == "In re\tCase No. 26-1"


@pytest.mark.unit
def test_apply_corrected_text_preserves_pPr():
    """Paragraph properties (style/alignment) survive the rebuild."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    p = doc.add_paragraph("wrong content")
    pPr = OxmlElement("w:pPr")
    p._element.insert(0, pPr)
    _apply_corrected_text(p._element, "fixed content")
    assert p._element.find(qn("w:pPr")) is not None


# === autofix_fill_drift orchestration (mocked LLM) ===


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_no_drifts_returns_original_bytes():
    original = _build_docx(["anything"])
    result = await autofix_fill_drift(
        filled_bytes=original,
        template_bytes=original,
        resolved_values={},
        drifted_paragraphs=[],
    )
    assert result == original


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_happy_path_replaces_drifted_paragraph():
    """LLM proposes a content-equivalent fix → fixer applies it."""
    template = _build_docx(["[[cos_block]]"])
    filled_wrong = _build_docx(["Robin R. Weiner and auto@ch13.com"])
    resolved = {"cos_block": "Robin R. Weiner\nauto@ch13.com"}

    drift_result = validate_fill_format(template, filled_wrong, resolved)
    assert not drift_result.ok

    async def fake_llm_fix(**_kwargs):
        return "Robin R. Weiner\nauto@ch13.com"

    with patch(
        "src.core.common.documents.template_format_fixer._call_sonnet_for_fix",
        new=fake_llm_fix,
    ):
        fixed_bytes = await autofix_fill_drift(
            filled_bytes=filled_wrong,
            template_bytes=template,
            resolved_values=resolved,
            drifted_paragraphs=drift_result.drifted_paragraphs,
        )

    # The fixed bytes should now match the template + values correctly.
    fixed_doc = Document(BytesIO(fixed_bytes))
    assert _paragraph_text_with_breaks(fixed_doc.paragraphs[0]) == (
        "Robin R. Weiner\nauto@ch13.com"
    )

    # And the re-run validator should be clean.
    re_result = validate_fill_format(template, fixed_bytes, resolved)
    assert re_result.ok, re_result.summary


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_rejects_llm_content_drift():
    """LLM tried to rewrite content (not just formatting) — guardrail
    rejects, original bytes returned unchanged."""
    template = _build_docx(["[[name]]"])
    filled = _build_docx(["Robin R. Weiner"])
    resolved = {"name": "Robin R. Weiner"}

    drift = [
        DriftedParagraph(
            template_paragraph_index=0,
            template_text="[[name]]",
            reconstructed_text="Robin R. Weiner",
            drift_reason="test",
        ),
    ]

    async def hallucinating_llm(**_kwargs):
        return "Robyn R. Weiner"  # 'i' → 'y' content drift

    with patch(
        "src.core.common.documents.template_format_fixer._call_sonnet_for_fix",
        new=hallucinating_llm,
    ):
        result = await autofix_fill_drift(
            filled_bytes=filled,
            template_bytes=template,
            resolved_values=resolved,
            drifted_paragraphs=drift,
        )

    # Bytes unchanged — guardrail rejected the LLM fix.
    assert result == filled


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_soft_fails_on_llm_error():
    """LLM call raises → fixer logs warning, returns original bytes."""
    template = _build_docx(["[[name]]"])
    filled = _build_docx(["Robin R. Weiner"])
    drift = [
        DriftedParagraph(
            template_paragraph_index=0,
            template_text="[[name]]",
            reconstructed_text="Robin R. Weiner",
            drift_reason="test",
        ),
    ]

    async def boom(**_kwargs):
        raise RuntimeError("Sonnet API timeout")

    with patch(
        "src.core.common.documents.template_format_fixer._call_sonnet_for_fix",
        new=boom,
    ):
        result = await autofix_fill_drift(
            filled_bytes=filled,
            template_bytes=template,
            resolved_values={"name": "Robin R. Weiner"},
            drifted_paragraphs=drift,
        )

    assert result == filled


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_skips_when_llm_returns_none():
    """LLM returns None → skip this drift, leave bytes unchanged."""
    template = _build_docx(["[[name]]"])
    filled = _build_docx(["Robin R. Weiner"])
    drift = [
        DriftedParagraph(
            template_paragraph_index=0,
            template_text="[[name]]",
            reconstructed_text="Robin R. Weiner",
            drift_reason="test",
        ),
    ]

    async def llm_returns_none(**_kwargs):
        return None

    with patch(
        "src.core.common.documents.template_format_fixer._call_sonnet_for_fix",
        new=llm_returns_none,
    ):
        result = await autofix_fill_drift(
            filled_bytes=filled,
            template_bytes=template,
            resolved_values={"name": "Robin R. Weiner"},
            drifted_paragraphs=drift,
        )

    assert result == filled


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_skips_when_llm_returns_unchanged_text():
    """LLM returns the SAME filled text (says 'no fix needed') → skip."""
    template = _build_docx(["[[name]]"])
    filled = _build_docx(["Robin R. Weiner"])
    drift = [
        DriftedParagraph(
            template_paragraph_index=0,
            template_text="[[name]]",
            reconstructed_text="Robin R. Weiner",
            drift_reason="test",
        ),
    ]

    async def no_change(**_kwargs):
        return "Robin R. Weiner"  # identical

    with patch(
        "src.core.common.documents.template_format_fixer._call_sonnet_for_fix",
        new=no_change,
    ):
        result = await autofix_fill_drift(
            filled_bytes=filled,
            template_bytes=template,
            resolved_values={"name": "Robin R. Weiner"},
            drifted_paragraphs=drift,
        )

    assert result == filled


@pytest.mark.unit
@pytest.mark.asyncio
async def test_autofix_handles_out_of_range_paragraph_index():
    """Drift entry references a paragraph index that doesn't exist —
    skip gracefully, no crash."""
    template = _build_docx(["short doc"])
    filled = _build_docx(["short doc"])
    drift = [
        DriftedParagraph(
            template_paragraph_index=999,
            template_text="[[name]]",
            reconstructed_text="Robin R. Weiner",
            drift_reason="test",
        ),
    ]

    result = await autofix_fill_drift(
        filled_bytes=filled,
        template_bytes=template,
        resolved_values={"name": "Robin R. Weiner"},
        drifted_paragraphs=drift,
    )

    assert result == filled
