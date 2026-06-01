"""Tests for the template format validator (Tier 1, deterministic).

The validator detects drift between the original docx and the
placeholder-marked template produced by `create_template`. Drift =
any change to formatting (tabs, line breaks, paragraph splits,
leading whitespace) OUTSIDE the marker spans. The substitution
engine should preserve all such formatting; drift means something
went wrong.

Tier 1 only LOGS drift; it doesn't fix anything. These tests confirm
the detection logic — particularly that the validator:
  - Doesn't false-positive on healthy templates.
  - Catches drift the substitution engine could introduce.
  - Tolerates aliases (typo variants).
  - Tolerates typography (NBSP / curly quotes).
"""

from io import BytesIO

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.core.agents.types.spec import TemplateVariable
from src.core.common.documents.template_format_validator import (
    validate_fill_format,
    validate_template_format,
)


def _build_docx_from_paragraphs(paragraphs: list[str]) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# === Healthy cases — must NOT false-positive ===


@pytest.mark.unit
def test_no_drift_when_template_was_produced_cleanly():
    original = _build_docx_from_paragraphs([
        "The Debtor, John Smith, has filed.",
        "Case No. 25-12345.",
    ])
    template = _build_docx_from_paragraphs([
        "The Debtor, [[debtor_name]], has filed.",
        "Case No. [[case_number]].",
    ])
    spec = [
        TemplateVariable(
            template_variable="debtor_name", template_index=0,
            template_property_marker="John Smith",
            template_variable_string="[[debtor_name]]",
        ),
        TemplateVariable(
            template_variable="case_number", template_index=1,
            template_property_marker="25-12345",
            template_variable_string="[[case_number]]",
        ),
    ]
    result = validate_template_format(original, template, spec)
    assert result.ok, result.summary
    assert result.drifted_paragraphs == []


@pytest.mark.unit
def test_no_drift_for_paragraphs_without_placeholders():
    """A paragraph with no placeholder should match the original verbatim."""
    original = _build_docx_from_paragraphs([
        "Static header — does not change.",
        "Body text mentioning John Smith.",
    ])
    template = _build_docx_from_paragraphs([
        "Static header — does not change.",
        "Body text mentioning [[debtor_name]].",
    ])
    spec = [
        TemplateVariable(
            template_variable="debtor_name", template_index=0,
            template_property_marker="John Smith",
            template_variable_string="[[debtor_name]]",
        ),
    ]
    assert validate_template_format(original, template, spec).ok


@pytest.mark.unit
def test_no_drift_with_aliases():
    """Body uses one spelling, header another — single variable with aliases."""
    original = _build_docx_from_paragraphs([
        "Caption: Judith S Schwartz, Debtor.",
        "Body: Judith S. Schwartz signed the motion.",
    ])
    template = _build_docx_from_paragraphs([
        "Caption: [[debtor_name]], Debtor.",
        "Body: [[debtor_name]] signed the motion.",
    ])
    spec = [
        TemplateVariable(
            template_variable="debtor_name", template_index=0,
            template_property_marker="Judith S. Schwartz",
            template_property_marker_aliases=["Judith S Schwartz"],
            template_variable_string="[[debtor_name]]",
        ),
    ]
    assert validate_template_format(original, template, spec).ok


@pytest.mark.unit
def test_no_drift_when_original_has_curly_quotes_template_has_straight():
    """Typography fold: original has curly apostrophe, template carries
    the marker's straight one. Should still validate clean."""
    # Original: Debtor’s position (curly U+2019)
    original = _build_docx_from_paragraphs([
        "Given that the Debtor’s position requires care.",
    ])
    template = _build_docx_from_paragraphs([
        "Given that [[X]] requires care.",
    ])
    spec = [
        TemplateVariable(
            template_variable="X", template_index=0,
            # Marker uses straight apostrophe — what the LLM extracted.
            template_property_marker="the Debtor's position",
            template_variable_string="[[X]]",
        ),
    ]
    assert validate_template_format(original, template, spec).ok


# === Drift cases — must catch ===


@pytest.mark.unit
def test_detects_drift_when_template_lost_surrounding_text():
    """Template paragraph collapses — text after the marker is missing."""
    original = _build_docx_from_paragraphs([
        "The Debtor, John Smith, has filed today.",
    ])
    # Template lost "has filed today." — represents a substitution bug.
    template = _build_docx_from_paragraphs([
        "The Debtor, [[debtor_name]],",
    ])
    spec = [
        TemplateVariable(
            template_variable="debtor_name", template_index=0,
            template_property_marker="John Smith",
            template_variable_string="[[debtor_name]]",
        ),
    ]
    result = validate_template_format(original, template, spec)
    assert not result.ok
    assert len(result.drifted_paragraphs) == 1
    assert "John Smith" in result.drifted_paragraphs[0].reconstructed_text


@pytest.mark.unit
def test_detects_drift_when_paragraphs_merged_into_one():
    """Two original paragraphs collapsed into one in the template — like
    the Fleisher CoS bug where tab/break was lost."""
    original = _build_docx_from_paragraphs([
        "First paragraph mentioning John Smith.",
        "Second standalone paragraph.",
    ])
    # Template merged into a single paragraph (drift).
    template = _build_docx_from_paragraphs([
        "First paragraph mentioning [[debtor_name]].Second standalone paragraph.",
    ])
    spec = [
        TemplateVariable(
            template_variable="debtor_name", template_index=0,
            template_property_marker="John Smith",
            template_variable_string="[[debtor_name]]",
        ),
    ]
    result = validate_template_format(original, template, spec)
    assert not result.ok


@pytest.mark.unit
def test_detects_drift_when_static_paragraph_changed():
    """Paragraph with no placeholders should match original — if it
    differs, drift."""
    original = _build_docx_from_paragraphs([
        "Static header that should never change.",
    ])
    template = _build_docx_from_paragraphs([
        "Static header that SOMEONE EDITED.",
    ])
    result = validate_template_format(original, template, [
        TemplateVariable(
            template_variable="debtor_name", template_index=0,
            template_property_marker="anything",
            template_variable_string="[[debtor_name]]",
        ),
    ])
    assert not result.ok


# === Edge cases ===


@pytest.mark.unit
def test_empty_spec_skips_validation():
    """No spec → nothing to validate against → ok."""
    original = _build_docx_from_paragraphs(["anything"])
    template = _build_docx_from_paragraphs(["something completely different"])
    result = validate_template_format(original, template, [])
    assert result.ok
    assert "Skipping" in result.summary or "skipping" in result.summary


@pytest.mark.unit
def test_no_placeholders_in_spec_skips():
    """Spec exists but no template_variable_string set → skip."""
    original = _build_docx_from_paragraphs(["anything"])
    template = _build_docx_from_paragraphs(["something completely different"])
    spec = [
        TemplateVariable(
            template_variable="x", template_index=0,
            template_property_marker="anything",
            template_variable_string=None,
        ),
    ]
    result = validate_template_format(original, template, spec)
    assert result.ok


@pytest.mark.unit
def test_handles_cross_paragraph_marker_substitution():
    """A marker with `\\n` spans two original paragraphs but lands in
    one template paragraph. The validator should accept the cross-para
    join match rather than flagging it as drift."""
    original = _build_docx_from_paragraphs([
        "Michael R. Bakst",
        "trustee@bakst.example",
    ])
    template = _build_docx_from_paragraphs([
        "[[cos_block]]",
    ])
    spec = [
        TemplateVariable(
            template_variable="cos_block", template_index=0,
            template_property_marker="Michael R. Bakst\ntrustee@bakst.example",
            template_variable_string="[[cos_block]]",
        ),
    ]
    result = validate_template_format(original, template, spec)
    assert result.ok, result.summary


@pytest.mark.unit
def test_validator_does_not_modify_inputs():
    """Pure function — neither bytes argument is mutated."""
    original = _build_docx_from_paragraphs(["The Debtor, John Smith."])
    template = _build_docx_from_paragraphs(["The Debtor, [[debtor_name]]."])
    original_copy = bytes(original)
    template_copy = bytes(template)
    spec = [
        TemplateVariable(
            template_variable="debtor_name", template_index=0,
            template_property_marker="John Smith",
            template_variable_string="[[debtor_name]]",
        ),
    ]
    validate_template_format(original, template, spec)
    assert original == original_copy
    assert template == template_copy


# === Fill-time validation ===


@pytest.mark.unit
def test_fill_no_drift_single_value():
    template = _build_docx_from_paragraphs(["The Debtor, [[debtor_name]], has filed."])
    filled = _build_docx_from_paragraphs(["The Debtor, John Smith, has filed."])
    result = validate_fill_format(template, filled, {"debtor_name": "John Smith"})
    assert result.ok, result.summary


@pytest.mark.unit
def test_fill_no_drift_multiple_placeholders():
    template = _build_docx_from_paragraphs([
        "Case No. [[case_number]] — Chapter [[chapter]]",
    ])
    filled = _build_docx_from_paragraphs(["Case No. 25-12345 — Chapter 7"])
    resolved = {"case_number": "25-12345", "chapter": "7"}
    result = validate_fill_format(template, filled, resolved)
    assert result.ok


@pytest.mark.unit
def test_fill_no_drift_inline_multiline_value_joined_with_and():
    """Multi-line value rendered inline as `A and B` — should still
    validate clean."""
    template = _build_docx_from_paragraphs([
        "The Debtors [[debtor_name]] filed jointly.",
    ])
    filled = _build_docx_from_paragraphs([
        "The Debtors Lori Creswell and Robert Creswell filed jointly.",
    ])
    resolved = {"debtor_name": "Lori Creswell\nRobert Creswell"}
    result = validate_fill_format(template, filled, resolved)
    assert result.ok


@pytest.mark.unit
def test_fill_no_drift_caption_multiline_value_rendered_with_breaks():
    """Multi-line value in a caption-shape paragraph keeps its `\\n`."""
    template = _build_docx_from_paragraphs(["[[debtor_name]]"])
    filled = _build_docx_from_paragraphs([])
    # Build the filled doc by hand so we can inject the soft break that
    # _render_lines_into_placeholder would produce in practice.
    from docx import Document as DocxDocument
    doc = DocxDocument()
    p = doc.add_paragraph()
    run = p.add_run("Lori Creswell")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "line")
    run._element.append(br)
    t = OxmlElement("w:t")
    t.text = "Robert Creswell"
    t.set(qn("xml:space"), "preserve")
    run._element.append(t)
    buf = BytesIO()
    doc.save(buf)
    filled = buf.getvalue()

    resolved = {"debtor_name": "Lori Creswell\nRobert Creswell"}
    result = validate_fill_format(template, filled, resolved)
    assert result.ok, result.summary


@pytest.mark.unit
def test_fill_detects_drift_when_text_around_placeholder_lost():
    template = _build_docx_from_paragraphs([
        "The Debtor, [[debtor_name]], has filed today.",
    ])
    filled = _build_docx_from_paragraphs([
        "The Debtor, John Smith,",  # lost "has filed today."
    ])
    result = validate_fill_format(template, filled, {"debtor_name": "John Smith"})
    assert not result.ok
    assert "diverged from template" in result.summary


@pytest.mark.unit
def test_fill_detects_drift_on_paragraph_count_mismatch():
    template = _build_docx_from_paragraphs(["[[debtor_name]]", "Static line."])
    filled = _build_docx_from_paragraphs(["John Smith"])  # static line lost
    result = validate_fill_format(template, filled, {"debtor_name": "John Smith"})
    assert not result.ok
    assert "Paragraph count mismatch" in result.summary


@pytest.mark.unit
def test_fill_empty_resolved_values_skips():
    template = _build_docx_from_paragraphs(["anything"])
    filled = _build_docx_from_paragraphs(["something completely different"])
    result = validate_fill_format(template, filled, {})
    assert result.ok
    assert "skipping" in result.summary.lower()


@pytest.mark.unit
def test_fill_detects_drift_when_multi_placeholder_paragraph_rendered_with_and_join():
    """Fleisher CoS bug variant — TWO placeholders share one paragraph
    (the CoS recipients block), each resolving to a multi-line
    `Name\\nEmail` value. `_substitute_placeholder` mis-classifies this
    paragraph as inline (because the OTHER placeholder's `[[name]]`
    token contributes alphanum residue), then both values get joined
    with `" and "`. The validator's smarter caption-shape predictor
    strips ALL `[[…]]` before the residue check, predicts CAPTION
    shape (correct), and catches the drift."""
    template = _build_docx_from_paragraphs([
        "[[cos_email_section_1]]\n[[cos_email_section_2]]",
    ])
    filled = _build_docx_from_paragraphs([
        "Robin R. Weiner and auto@ch13.com\n"
        "Office of the US Trustee and USTPRegion21@usdoj.gov",
    ])
    resolved = {
        "cos_email_section_1": "Robin R. Weiner\nauto@ch13.com",
        "cos_email_section_2": "Office of the US Trustee\nUSTPRegion21@usdoj.gov",
    }
    result = validate_fill_format(template, filled, resolved)
    assert not result.ok, result.summary


@pytest.mark.unit
def test_fill_detects_drift_when_caption_paragraph_rendered_with_and_join():
    """The user's CoS bug: template paragraph is caption-shape
    `[[cos_email_section_2]]` (placeholder is the whole paragraph).
    Resolved value is multi-line `Name\\nEmail`. Correct rendering
    uses `<w:br/>` so name and email appear on separate lines. Buggy
    rendering used " and "-join — visible as
    `Robin R. Weiner and auto-forward-ecf@ch13weiner.com`.

    Validator must detect this — the inline `" and "` form is
    invalid for a caption-shape paragraph."""
    template = _build_docx_from_paragraphs(["[[cos_email_section_2]]"])
    # Filled with the BUG — inline " and " join.
    filled = _build_docx_from_paragraphs([
        "Robin R. Weiner and auto-forward-ecf@ch13weiner.com",
    ])
    resolved = {
        "cos_email_section_2": "Robin R. Weiner\nauto-forward-ecf@ch13weiner.com",
    }
    result = validate_fill_format(template, filled, resolved)
    assert not result.ok, result.summary
    assert "diverged from template" in result.summary


@pytest.mark.unit
def test_fill_tolerates_typography_drift_in_filled_output():
    """Filled docx has curly apostrophe where the resolved value
    carried a straight one. Should still validate clean."""
    template = _build_docx_from_paragraphs(["The [[case]] is open."])
    filled = _build_docx_from_paragraphs(["The Debtor’s case is open."])  # curly
    resolved = {"case": "Debtor's case"}  # straight
    result = validate_fill_format(template, filled, resolved)
    assert result.ok
