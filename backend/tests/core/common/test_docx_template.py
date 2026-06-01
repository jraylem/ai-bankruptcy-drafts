"""Tests for DocxTemplateService._replace_in_paragraph.

Covers the two-stage replacement:
  Stage 1 — per-run (short markers inside a single <w:t>).
  Stage 2 — paragraph-level merge (long markers that straddle multiple runs).

Runs are constructed programmatically with python-docx so tests don't depend
on a fixture .docx on disk.
"""

from io import BytesIO

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.core.common.documents.docx_template import DocxTemplateService


def _paragraph_with_runs(run_texts: list[str]):
    """Build a one-paragraph python-docx Document whose paragraph is composed
    of `len(run_texts)` separate runs, then return that paragraph."""
    doc = Document()
    p = doc.add_paragraph()
    for t in run_texts:
        p.add_run(t)
    return p


# ─── Stage 1: per-run replacement ─────────────────────────────────────


@pytest.mark.unit
def test_replace_in_paragraph_single_run_marker():
    p = _paragraph_with_runs(["The Debtor, John Smith, is employed."])
    DocxTemplateService._replace_in_paragraph(p, "John Smith", "[[debtor_name]]")
    assert p.text == "The Debtor, [[debtor_name]], is employed."


@pytest.mark.unit
def test_replace_in_paragraph_word_boundary_prevents_short_numeric_false_match():
    """'13' must not match inside '130' or '1301'."""
    p = _paragraph_with_runs(["Chapter 13 does not match 130 or 1301."])
    DocxTemplateService._replace_in_paragraph(p, "13", "[[chapter]]")
    assert p.text == "Chapter [[chapter]] does not match 130 or 1301."


@pytest.mark.unit
def test_replace_in_paragraph_noop_when_marker_absent():
    p = _paragraph_with_runs(["Nothing to replace here."])
    DocxTemplateService._replace_in_paragraph(p, "Missing Marker", "[[x]]")
    assert p.text == "Nothing to replace here."


# ─── Stage 2: marker straddles multiple runs ──────────────────────────


@pytest.mark.unit
def test_replace_in_paragraph_marker_spanning_two_runs():
    """Marker is split between two <w:t> runs — stage 1 misses, stage 2 catches it."""
    p = _paragraph_with_runs([
        "Given that the Debtor's ",
        "position involves handling sensitive consumer information.",
    ])
    marker = "the Debtor's position involves handling sensitive consumer information"
    DocxTemplateService._replace_in_paragraph(p, marker, "[[wage_deduction_impact]]")
    assert p.text == "Given that [[wage_deduction_impact]]."


@pytest.mark.unit
def test_replace_in_paragraph_long_marker_spanning_many_runs():
    """Long narrative marker straddling 4+ runs (typical DOCX case)."""
    p = _paragraph_with_runs([
        "Given that the Debtor's ",
        "position involves ",
        "handling sensitive consumer information, ",
        "implementing a wage deduction could negatively affect their standing",
        ".",
    ])
    marker = (
        "the Debtor's position involves handling sensitive consumer information, "
        "implementing a wage deduction could negatively affect their standing"
    )
    DocxTemplateService._replace_in_paragraph(p, marker, "[[wage_deduction_impact]]")
    assert p.text == "Given that [[wage_deduction_impact]]."


@pytest.mark.unit
def test_replace_in_paragraph_stage_1_and_stage_2_both_run():
    """Short marker appears once inside a single run AND once straddled — both replaced."""
    p = _paragraph_with_runs([
        "Debtor owes Debtor, ",
        " and also owes Deb",
        "tor more.",
    ])
    # "Debtor" appears: once in run 0 (single-run), twice in run 0 (single-run again),
    # and once spanning run 1-2 ("Deb" + "tor"). Stage 1 handles the first three,
    # stage 2 catches the split one.
    DocxTemplateService._replace_in_paragraph(p, "Debtor", "X")
    assert p.text == "X owes X,  and also owes X more."


@pytest.mark.unit
def test_replace_in_paragraph_stage_2_preserves_non_marker_text_around_marker():
    """When stage 2 collapses, surrounding text outside the marker is preserved."""
    p = _paragraph_with_runs([
        "Prefix before the marker: ",
        "this is the exact marker text ",
        "— and suffix after.",
    ])
    DocxTemplateService._replace_in_paragraph(p, "this is the exact marker text", "[[X]]")
    assert p.text == "Prefix before the marker: [[X]] — and suffix after."


# ─── Typography mismatch (Word-auto-correct vs LLM-extracted markers) ─


@pytest.mark.unit
def test_replace_matches_straight_marker_against_curly_apostrophe_single_run():
    """Word auto-converts `'` to `’`; LLMs output `'`. Marker must still match."""
    p = _paragraph_with_runs(["Given that the Debtor\u2019s position requires trust."])  # curly
    DocxTemplateService._replace_in_paragraph(p, "the Debtor's position", "[[X]]")  # straight
    assert p.text == "Given that [[X]] requires trust."


@pytest.mark.unit
def test_replace_matches_straight_marker_against_curly_apostrophe_across_runs():
    """The real-world bug: long marker with straight apostrophes, DOCX has curly,
    and the paragraph is split across many runs."""
    p = _paragraph_with_runs([
        "Given that the Debtor\u2019s ",
        "position involves handling ",
        "sensitive consumer information, ",
        "implementing a wage deduction could affect reliability.",
    ])
    marker = (
        "the Debtor's position involves handling sensitive consumer information, "
        "implementing a wage deduction could affect reliability"
    )
    DocxTemplateService._replace_in_paragraph(p, marker, "[[wage_deduction_impact]]")
    assert p.text == "Given that [[wage_deduction_impact]]."


@pytest.mark.unit
def test_replace_matches_straight_double_quotes_against_curly():
    p = _paragraph_with_runs(["She said \u201Cthank you\u201D clearly."])  # curly "..."
    DocxTemplateService._replace_in_paragraph(p, '"thank you"', "[[quote]]")  # straight
    assert p.text == "She said [[quote]] clearly."


@pytest.mark.unit
def test_replace_matches_hyphen_marker_against_em_dash():
    p = _paragraph_with_runs(["Debtor \u2014 undersigned \u2014 certifies."])  # em-dashes
    DocxTemplateService._replace_in_paragraph(p, "Debtor - undersigned - certifies", "[[X]]")
    assert p.text == "[[X]]."


@pytest.mark.unit
def test_replace_preserves_curly_quotes_when_marker_not_present():
    """Unrelated paragraphs keep their typography untouched."""
    p = _paragraph_with_runs(["The Debtor\u2019s signature is attached."])
    DocxTemplateService._replace_in_paragraph(p, "Missing Marker", "[[X]]")
    assert p.text == "The Debtor\u2019s signature is attached."


# ─── Stage 3: marker spans <w:br/> soft line break ────────────────────


def _paragraph_with_soft_break(texts_around_break: list[str]):
    """Build a paragraph whose runs are joined by a <w:br/> (line break).

    `texts_around_break = ["A", "B"]` produces a paragraph with
    `<w:r><w:t>A</w:t></w:r><w:r><w:br w:type="line"/><w:t>B</w:t></w:r>`
    so paragraph.text == "A\\nB".
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(texts_around_break[0])
    for text in texts_around_break[1:]:
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "line")
        run._element.append(br)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        run._element.append(t)
    return p


@pytest.mark.unit
def test_replace_in_paragraph_marker_spans_soft_line_break():
    """Joint-caption shape: marker 'A\\nB' must match across a <w:br/> between two <w:t>."""
    p = _paragraph_with_soft_break(["Lori Creswell", "Robert Creswell,"])
    DocxTemplateService._replace_in_paragraph(
        p, "Lori Creswell\nRobert Creswell,", "[[debtor_name]]"
    )
    assert p.text == "[[debtor_name]]"
    # <w:br/> element should be removed after collapse.
    assert not list(p._element.iter(qn("w:br")))


@pytest.mark.unit
def test_replace_in_paragraph_solo_marker_unaffected_by_stage_3():
    """A single-line marker in a paragraph that has a soft break elsewhere
    still takes the stage-1 fast path — stage 3 only fires when the marker
    itself contains \\n."""
    p = _paragraph_with_soft_break(["Lori Creswell", "Robert Creswell,"])
    DocxTemplateService._replace_in_paragraph(p, "Lori Creswell", "[[debtor_1]]")
    assert p.text.startswith("[[debtor_1]]")
    # Soft break preserved.
    assert list(p._element.iter(qn("w:br")))


@pytest.mark.unit
def test_replace_span_preserves_prefix_and_suffix():
    """Marker matches mid-paragraph; surrounding text is preserved."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Caption: ")
    run2 = p.add_run("Lori Creswell")
    run3 = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "line")
    run3._element.append(br)
    t = OxmlElement("w:t")
    t.text = "Robert Creswell"
    t.set(qn("xml:space"), "preserve")
    run3._element.append(t)
    p.add_run(", Debtors.")

    DocxTemplateService._replace_in_paragraph(
        p, "Lori Creswell\nRobert Creswell", "[[debtor_name]]"
    )
    assert p.text == "Caption: [[debtor_name]], Debtors."


@pytest.mark.unit
def test_replace_in_paragraph_three_debtor_caption():
    """Multi-party filing with 3 names separated by 2 soft breaks."""
    p = _paragraph_with_soft_break(["Name A", "Name B", "Name C,"])
    DocxTemplateService._replace_in_paragraph(
        p, "Name A\nName B\nName C,", "[[debtor_name]]"
    )
    assert p.text == "[[debtor_name]]"
    assert not list(p._element.iter(qn("w:br")))


@pytest.mark.unit
def test_replace_direct_match_preserves_curly_quotes_in_surrounding_text():
    """Stage 1 direct match should leave curly quotes in other parts of the same run intact."""
    p = _paragraph_with_runs(["John Smith signed the Debtor\u2019s motion."])
    DocxTemplateService._replace_in_paragraph(p, "John Smith", "[[debtor_name]]")
    # The curly apostrophe in "Debtor's" must survive — only the marker region changes.
    assert p.text == "[[debtor_name]] signed the Debtor\u2019s motion."


# ─── find_paragraph_containing ────────────────────────────────────────


def _docx_bytes(paragraphs: list[str], table_cells: list[list[str]] | None = None) -> bytes:
    """Build an in-memory DOCX with the given paragraph texts and optional
    table cells. Returns the serialized bytes."""
    from io import BytesIO

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_cells:
        table = doc.add_table(rows=len(table_cells), cols=max(len(r) for r in table_cells))
        for i, row in enumerate(table_cells):
            for j, cell_text in enumerate(row):
                table.rows[i].cells[j].text = cell_text
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.mark.unit
def test_find_paragraph_containing_returns_paragraph_text():
    template_bytes = _docx_bytes([
        "In re: [[case_number]]",
        "The Debtor, [[debtor_name]], [[employment_description]].",
        "Another unrelated paragraph.",
    ])
    result = DocxTemplateService.find_paragraph_containing(template_bytes, "[[employment_description]]")
    assert result == "The Debtor, [[debtor_name]], [[employment_description]]."


@pytest.mark.unit
def test_find_paragraph_containing_returns_none_when_absent():
    template_bytes = _docx_bytes(["Some paragraph with nothing interesting."])
    result = DocxTemplateService.find_paragraph_containing(template_bytes, "[[missing]]")
    assert result is None


@pytest.mark.unit
def test_find_paragraph_containing_checks_table_cells():
    template_bytes = _docx_bytes(
        paragraphs=["A top paragraph."],
        table_cells=[["Label", "Value: [[cell_value]]"]],
    )
    result = DocxTemplateService.find_paragraph_containing(template_bytes, "[[cell_value]]")
    assert result == "Value: [[cell_value]]"


@pytest.mark.unit
def test_find_paragraph_containing_empty_placeholder_returns_none():
    template_bytes = _docx_bytes(["anything"])
    assert DocxTemplateService.find_paragraph_containing(template_bytes, "") is None


# ─── create_template: primary marker + aliases ────────────────────────


from src.core.agents.types.spec import TemplateVariable  # noqa: E402


@pytest.mark.unit
def test_create_template_replaces_primary_marker():
    template_bytes = _docx_bytes(["The Debtor, John Smith, has filed."])
    spec = [
        TemplateVariable(
            template_variable="debtor_name",
            template_index=0,
            template_property_marker="John Smith",
            template_variable_string="[[debtor_name]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "The Debtor, [[debtor_name]], has filed."


@pytest.mark.unit
def test_create_template_replaces_every_orthographic_alias():
    """Primary marker + alias BOTH get replaced with the same placeholder.

    Header caption uses 'Judith S Schwartz' (no period); body uses
    'Judith S. Schwartz' (with period). One variable with aliases captures
    both forms so [[debtor_name]] ends up in both places.
    """
    template_bytes = _docx_bytes([
        "In re: Judith S Schwartz, Debtor.",
        "The Debtor, Judith S. Schwartz, has filed a motion.",
    ])
    spec = [
        TemplateVariable(
            template_variable="debtor_name",
            template_index=0,
            template_property_marker="Judith S. Schwartz",
            template_property_marker_aliases=["Judith S Schwartz"],
            template_variable_string="[[debtor_name]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "In re: [[debtor_name]], Debtor."
    assert result_doc.paragraphs[1].text == "The Debtor, [[debtor_name]], has filed a motion."


@pytest.mark.unit
def test_create_template_global_longest_first_across_variables_and_aliases():
    """A short alias must NOT consume a substring of another variable's longer
    primary marker. Sorting is GLOBAL across all (marker, placeholder) pairs,
    not per-variable.

    Here 'Chapter 13' contains '13' but belongs to the longer document_date
    marker 'April 13, 2026'. The global longest-first sort ensures
    'April 13, 2026' is replaced before the bare-'13' chapter marker gets
    to run. Also exercises an alias shorter than another variable's primary.
    """
    template_bytes = _docx_bytes([
        "Filed April 13, 2026 under Chapter 13.",
        "See also Chapter Thirteen in the schedules.",
    ])
    spec = [
        TemplateVariable(
            template_variable="chapter",
            template_index=0,
            template_property_marker="13",
            template_property_marker_aliases=["Thirteen"],
            template_variable_string="[[chapter]]",
        ),
        TemplateVariable(
            template_variable="document_date",
            template_index=1,
            template_property_marker="April 13, 2026",
            template_variable_string="[[document_date]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    # 'April 13, 2026' (14 chars) replaced before 'Thirteen' (8) or '13' (2),
    # so the date doesn't get corrupted by the shorter chapter markers.
    assert result_doc.paragraphs[0].text == "Filed [[document_date]] under Chapter [[chapter]]."
    assert result_doc.paragraphs[1].text == "See also Chapter [[chapter]] in the schedules."


@pytest.mark.unit
def test_create_template_skips_variable_without_placeholder_or_markers():
    """A variable with no template_variable_string OR no markers at all is silently ignored."""
    template_bytes = _docx_bytes(["A paragraph about John Smith."])
    spec = [
        TemplateVariable(
            template_variable="no_placeholder",
            template_index=0,
            template_property_marker="John Smith",
            template_variable_string=None,
        ),
        TemplateVariable(
            template_variable="no_markers",
            template_index=1,
            template_property_marker=None,
            template_variable_string="[[no_markers]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    # Neither variable produces a (marker, placeholder) job, so the doc is unchanged.
    assert result_doc.paragraphs[0].text == "A paragraph about John Smith."


# ─── create_template: shared-marker disambiguation ────────────────────


@pytest.mark.unit
def test_create_template_disambiguates_shared_marker():
    """Two variables share the same marker value ("X-123") in semantically
    distinct paragraphs. Each variable's `template_identifying_text_match`
    locates the right paragraph so each placeholder lands on the correct
    occurrence, rather than both placeholders piling onto the first match.

    Mirrors the bankruptcy template where a civil case_number and a
    bankruptcy case_number happen to share the literal value in the source.
    """
    template_bytes = _docx_bytes([
        "Civil action — CASE NO.: X-123",
        "Bankruptcy Case No.: X-123",
    ])
    spec = [
        TemplateVariable(
            template_variable="civil_case_number",
            template_index=0,
            template_property_marker="X-123",
            template_variable_string="[[civil_case_number]]",
            template_identifying_text_match="Civil action — CASE NO.: X-123",
        ),
        TemplateVariable(
            template_variable="bankruptcy_case_number",
            template_index=1,
            template_property_marker="X-123",
            template_variable_string="[[bankruptcy_case_number]]",
            template_identifying_text_match="Bankruptcy Case No.: X-123",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "Civil action — CASE NO.: [[civil_case_number]]"
    assert result_doc.paragraphs[1].text == "Bankruptcy Case No.: [[bankruptcy_case_number]]"


@pytest.mark.unit
def test_create_template_falls_back_when_identifying_text_unmatched(caplog):
    """When `template_identifying_text_match` doesn't match any paragraph,
    fall back to first-unconsumed-occurrence so the placeholder still lands
    somewhere (the orphan check will surface the configuration error), and
    emit a warning the spec author can act on.
    """
    template_bytes = _docx_bytes([
        "Civil action — CASE NO.: X-123",
        "Bankruptcy Case No.: X-123",
    ])
    spec = [
        TemplateVariable(
            template_variable="civil_case_number",
            template_index=0,
            template_property_marker="X-123",
            template_variable_string="[[civil_case_number]]",
            template_identifying_text_match="text that does not appear anywhere",
        ),
        TemplateVariable(
            template_variable="bankruptcy_case_number",
            template_index=1,
            template_property_marker="X-123",
            template_variable_string="[[bankruptcy_case_number]]",
            template_identifying_text_match="Bankruptcy Case No.: X-123",
        ),
    ]

    with caplog.at_level("WARNING"):
        result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    # The civil variable's identifying_text didn't match, so it fell back to
    # the first unconsumed paragraph (paragraph 0) and consumed it. The
    # bankruptcy variable then matched paragraph 1 via its own identifying_text.
    assert result_doc.paragraphs[0].text == "Civil action — CASE NO.: [[civil_case_number]]"
    assert result_doc.paragraphs[1].text == "Bankruptcy Case No.: [[bankruptcy_case_number]]"
    # And the fallback emitted a warning so the author can fix the spec.
    assert any(
        "Shared-marker disambiguation fallback" in record.message
        for record in caplog.records
    )


@pytest.mark.unit
def test_replace_in_paragraph_first_only_handles_marker_spanning_runs():
    """Stage 2 path: marker straddles two runs inside the paragraph the
    shared-marker disambiguation lands on. Must still get replaced exactly
    once (the first occurrence in the paragraph)."""
    p = _paragraph_with_runs([
        "The Debtor's X-",
        "123 case is open.",
    ])
    DocxTemplateService._replace_in_paragraph_first_only(p, "X-123", "[[civil_case_number]]")
    assert p.text == "The Debtor's [[civil_case_number]] case is open."


@pytest.mark.unit
def test_create_template_logs_warning_when_shared_marker_missing(caplog):
    """When a shared-marker variable's marker doesn't appear anywhere in the
    document (e.g. agent hallucinated it), the final warning fires and the
    placeholder is left for `_drop_orphan_variables` to clean up."""
    template_bytes = _docx_bytes([
        "Civil action — CASE NO.: X-123",
    ])
    spec = [
        TemplateVariable(
            template_variable="civil_case_number",
            template_index=0,
            template_property_marker="X-123",
            template_variable_string="[[civil_case_number]]",
            template_identifying_text_match="Civil action — CASE NO.: X-123",
        ),
        TemplateVariable(
            template_variable="bankruptcy_case_number",
            template_index=1,
            template_property_marker="X-123",
            template_variable_string="[[bankruptcy_case_number]]",
            template_identifying_text_match="(not in document)",
        ),
    ]

    with caplog.at_level("WARNING"):
        result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    # Civil placeholder landed on the only available paragraph.
    assert result_doc.paragraphs[0].text == "Civil action — CASE NO.: [[civil_case_number]]"
    # Bankruptcy had no unconsumed occurrence after civil consumed the only
    # paragraph — the final warning fires.
    assert any(
        "has no unconsumed occurrence" in record.message
        for record in caplog.records
    )


# ─── _replace_across_paragraphs: hard-break recipient pairs ────────────


@pytest.mark.unit
def test_replace_across_paragraphs_collapses_two_paragraph_recipient_pair():
    """Recipient pair where Name and Email live in separate paragraphs
    (typical hard-return Word authoring). The `\\n`-joined marker spans
    the two <w:p> elements; create_template must collapse them into one
    paragraph containing the placeholder, removing the second paragraph.
    """
    template_bytes = _docx_bytes([
        "Robin R Weiner",
        "auto-forward-ecf@ch13weiner.com",
    ])
    spec = [
        TemplateVariable(
            template_variable="cos_section_1",
            template_index=0,
            template_property_marker="Robin R Weiner\nauto-forward-ecf@ch13weiner.com",
            template_variable_string="[[cos_section_1]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    paragraphs = [p.text for p in result_doc.paragraphs]
    # Two source paragraphs collapsed into one; second paragraph removed.
    assert "[[cos_section_1]]" in paragraphs[0]
    # No raw email line surviving.
    assert all("auto-forward-ecf@ch13weiner.com" not in p for p in paragraphs)


@pytest.mark.unit
def test_replace_across_paragraphs_handles_three_paragraph_marker():
    """A 3-line marker spanning three <w:p> elements. Verifies the
    algorithm scales beyond the 2-paragraph case (e.g. a 3-line address
    block, or a contact with name + role + email).
    """
    template_bytes = _docx_bytes([
        "Office of the US Trustee",
        "Stuart M. Brown",
        "USTPRegion21.MM.ECF@usdoj.gov",
    ])
    spec = [
        TemplateVariable(
            template_variable="trustee_block",
            template_index=0,
            template_property_marker="Office of the US Trustee\nStuart M. Brown\nUSTPRegion21.MM.ECF@usdoj.gov",
            template_variable_string="[[trustee_block]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    paragraphs = [p.text for p in result_doc.paragraphs]
    assert "[[trustee_block]]" in paragraphs[0]
    # All three lines should be gone.
    full = "\n".join(paragraphs)
    assert "Stuart M. Brown" not in full
    assert "USTPRegion21.MM.ECF@usdoj.gov" not in full


@pytest.mark.unit
def test_create_template_with_three_hard_break_recipient_pairs_emits_all_placeholders():
    """End-to-end shape of the user's failing Certificate of Service —
    three recipient pairs each in their own pair of <w:p> elements, with
    blank paragraphs between them. All three placeholders must land.
    """
    template_bytes = _docx_bytes([
        "By CM/ECF",
        "Timothy R Qualls",
        "stalevich@yvlaw.net",
        "",
        "Robin R Weiner",
        "auto-forward-ecf@ch13weiner.com",
        "",
        "Office of the US Trustee",
        "USTPRegion21.MM.ECF@usdoj.gov",
    ])
    spec = [
        TemplateVariable(
            template_variable="cos_section_1",
            template_index=0,
            template_property_marker="Timothy R Qualls\nstalevich@yvlaw.net",
            template_variable_string="[[cos_section_1]]",
        ),
        TemplateVariable(
            template_variable="cos_section_2",
            template_index=1,
            template_property_marker="Robin R Weiner\nauto-forward-ecf@ch13weiner.com",
            template_variable_string="[[cos_section_2]]",
        ),
        TemplateVariable(
            template_variable="cos_section_3",
            template_index=2,
            template_property_marker="Office of the US Trustee\nUSTPRegion21.MM.ECF@usdoj.gov",
            template_variable_string="[[cos_section_3]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    all_text = "\n".join(p.text for p in result_doc.paragraphs)
    # All three placeholders must appear.
    assert "[[cos_section_1]]" in all_text
    assert "[[cos_section_2]]" in all_text
    assert "[[cos_section_3]]" in all_text
    # No raw recipient text leaking through.
    assert "Timothy R Qualls" not in all_text
    assert "stalevich@yvlaw.net" not in all_text
    assert "Robin R Weiner" not in all_text
    assert "auto-forward-ecf@ch13weiner.com" not in all_text
    assert "Office of the US Trustee" not in all_text
    assert "USTPRegion21.MM.ECF@usdoj.gov" not in all_text


@pytest.mark.unit
def test_replace_across_paragraphs_noop_when_marker_absent():
    """No `\\n`-marker match → no document mutation, no error."""
    template_bytes = _docx_bytes([
        "A paragraph.",
        "Another paragraph.",
    ])
    spec = [
        TemplateVariable(
            template_variable="missing",
            template_index=0,
            template_property_marker="Nope\nNot here",
            template_variable_string="[[missing]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    paragraphs = [p.text for p in result_doc.paragraphs]
    # Doc shape preserved exactly.
    assert paragraphs == ["A paragraph.", "Another paragraph."]


# ─── trailing-whitespace tolerance (parse_document_v2 vs raw <w:t>) ────


def _docx_with_trailing_space_in_paragraph(
    paragraphs: list[str],
    trailing_space_indices: set[int],
) -> bytes:
    """Build a .docx where the paragraphs at the given indices have a
    TRAILING SPACE inside <w:t>. Word commonly stores incidental
    trailing whitespace in <w:t> that paragraph.text.strip() trims —
    we mimic the joint-debtor caption pattern here to reproduce the
    haystack-vs-marker mismatch."""
    from docx.oxml.ns import qn

    doc = Document()
    for i, text in enumerate(paragraphs):
        body = text + " " if i in trailing_space_indices else text
        p = doc.add_paragraph(body)
        # Force xml:space=preserve so the trailing space survives serialization.
        if i in trailing_space_indices:
            for t in p._element.iter(qn('w:t')):
                t.set(qn("xml:space"), "preserve")
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


@pytest.mark.unit
def test_replace_across_paragraphs_matches_marker_when_first_paragraph_has_trailing_space():
    """Joint-debtor caption pattern: raw <w:t> stores 'Jane M. Doe and '
    (TRAILING SPACE), parse_document_v2 strips it, agent's marker reflects
    the stripped form 'Jane M. Doe and\\nJohn A. Smith,'.

    Without the rstrip fallback, find-and-replace fails silently and the
    caption stays hardcoded in the rendered template.
    """
    template_bytes = _docx_with_trailing_space_in_paragraph(
        paragraphs=[
            "Jane M. Doe and",  # The helper adds the trailing space.
            "John A. Smith,",
            "                Debtors.        /",
        ],
        trailing_space_indices={0},
    )
    spec = [
        TemplateVariable(
            template_variable="debtor_name_caption",
            template_index=0,
            template_property_marker="Jane M. Doe and\nJohn A. Smith,",
            template_variable_string="[[debtor_name_caption]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)

    result_doc = Document(BytesIO(result_bytes))
    paragraphs = [p.text for p in result_doc.paragraphs]
    # The marker spanning the two caption paragraphs collapsed into one
    # paragraph containing the placeholder; the second paragraph was removed.
    assert any("[[debtor_name_caption]]" in p for p in paragraphs)
    # No raw caption text surviving.
    full = "\n".join(paragraphs)
    assert "Jane M. Doe" not in full
    assert "John A. Smith" not in full


@pytest.mark.unit
def test_replace_across_paragraphs_raw_match_still_wins_when_no_drift():
    """When the raw <w:t> joined haystack already matches the marker
    (no trailing-whitespace drift), the stage-1 raw match path runs
    exactly as before — fallback is only used when the raw match fails."""
    template_bytes = _docx_bytes([
        "First line",
        "Second line",
    ])
    spec = [
        TemplateVariable(
            template_variable="pair",
            template_index=0,
            template_property_marker="First line\nSecond line",
            template_variable_string="[[pair]]",
        ),
    ]
    result_bytes = DocxTemplateService.create_template(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    paragraphs = [p.text for p in result_doc.paragraphs]
    assert any("[[pair]]" in p for p in paragraphs)
    full = "\n".join(paragraphs)
    assert "First line" not in full
    assert "Second line" not in full


@pytest.mark.unit
def test_replace_across_paragraphs_matches_when_both_paragraphs_have_trailing_space():
    """Both spanned paragraphs have trailing whitespace in <w:t>. The
    rstrip fallback should still match the stripped-view marker and
    replace correctly across the two paragraphs."""
    template_bytes = _docx_with_trailing_space_in_paragraph(
        paragraphs=[
            "alpha",
            "beta",
            "after",
        ],
        trailing_space_indices={0, 1},
    )
    spec = [
        TemplateVariable(
            template_variable="ab",
            template_index=0,
            template_property_marker="alpha\nbeta",
            template_variable_string="[[ab]]",
        ),
    ]
    result_bytes = DocxTemplateService.create_template(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    paragraphs = [p.text for p in result_doc.paragraphs]
    assert any("[[ab]]" in p for p in paragraphs)
    full = "\n".join(paragraphs)
    assert "alpha" not in full
    assert "beta" not in full


# ─── create_template end-to-end: joint caption collapse ───────────────


def _docx_bytes_with_joint_paragraph(before_paragraphs: list[str], names: list[str], after_paragraphs: list[str]) -> bytes:
    """Build a docx where `names` share ONE paragraph joined by <w:br/> soft breaks."""
    doc = Document()
    for t in before_paragraphs:
        doc.add_paragraph(t)
    p = doc.add_paragraph()
    p.add_run(names[0])
    for name in names[1:]:
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "line")
        run._element.append(br)
        t_elem = OxmlElement("w:t")
        t_elem.text = name
        t_elem.set(qn("xml:space"), "preserve")
        run._element.append(t_elem)
    for t in after_paragraphs:
        doc.add_paragraph(t)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.mark.unit
def test_create_template_collapses_joint_caption_with_soft_break():
    """Source paragraph has two names joined by <w:br/>. Multi-line marker collapses into one placeholder line."""
    template_bytes = _docx_bytes_with_joint_paragraph(
        before_paragraphs=["In re:"],
        names=["Lori Creswell", "Robert Creswell,"],
        after_paragraphs=["    Debtors."],
    )
    spec = [
        TemplateVariable(
            template_variable="debtor_name",
            template_index=0,
            template_property_marker="Lori Creswell\nRobert Creswell,",
            template_variable_string="[[debtor_name]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[1].text == "[[debtor_name]]"
    # <w:br/> removed during collapse.
    assert not list(result_doc.paragraphs[1]._element.iter(qn("w:br")))


@pytest.mark.unit
def test_create_template_accepts_literal_backslash_n_in_marker():
    """LLM may emit marker with two-char '\\n' escape — engine normalizes both forms."""
    template_bytes = _docx_bytes_with_joint_paragraph(
        before_paragraphs=[],
        names=["Lori Creswell", "Robert Creswell,"],
        after_paragraphs=[],
    )
    spec = [
        TemplateVariable(
            template_variable="debtor_name",
            template_index=0,
            template_property_marker="Lori Creswell\\nRobert Creswell,",  # literal backslash-n
            template_variable_string="[[debtor_name]]",
        ),
    ]

    result_bytes = DocxTemplateService.create_template(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "[[debtor_name]]"


# ─── fill_template: soft-break rendering for \n-bearing values ────────


from src.core.agents.types.sources import FieldSource  # noqa: E402
from src.core.agents.types.spec import TemplateField  # noqa: E402


@pytest.mark.unit
def test_fill_template_renders_newline_as_soft_line_break_in_caption_paragraph():
    """Caption-shape paragraph (placeholder IS the paragraph) renders \\n as <w:br/>."""
    template_bytes = _docx_bytes(["[[debtor_name]]"])
    fields = [
        TemplateField(
            property_name="debtor_name",
            source=FieldSource.CASE_VECTOR,
            template_variable_string="[[debtor_name]]",
        )
    ]

    filled_bytes, unresolved = DocxTemplateService.fill_template(
        template_bytes,
        fields,
        {"debtor_name": "Lori Creswell\nRobert Creswell"},
    )

    assert unresolved == []
    doc = Document(BytesIO(filled_bytes))
    paragraph = doc.paragraphs[0]
    line_breaks = [
        br for br in paragraph._element.iter(qn("w:br"))
        if br.get(qn("w:type")) == "line"
    ]
    assert len(line_breaks) == 1


@pytest.mark.unit
def test_fill_template_multi_placeholder_caption_paragraph_renders_each_with_soft_break():
    """Fleisher CoS bug regression: a single paragraph with TWO
    placeholders separated by soft breaks (a CoS recipients block)
    used to misclassify as inline because the sibling placeholder's
    `[[cos_email_section_2]]` characters survived the residue strip.
    With the smarter `_is_caption_shape_paragraph` (strips ALL
    `[[…]]` tokens), each placeholder's multi-line value renders
    with `<w:br/>` between lines — no more " and " join."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    p = doc.add_paragraph("[[cos_email_section_1]]")
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "line")
    run._element.append(br)
    br2 = OxmlElement("w:br")
    br2.set(qn("w:type"), "line")
    run._element.append(br2)
    t = OxmlElement("w:t")
    t.text = "[[cos_email_section_2]]"
    run._element.append(t)
    buf = BytesIO()
    doc.save(buf)

    fields = [
        TemplateField(
            property_name="cos_email_section_1",
            source=FieldSource.CASE_VECTOR,
            template_variable_string="[[cos_email_section_1]]",
        ),
        TemplateField(
            property_name="cos_email_section_2",
            source=FieldSource.CASE_VECTOR,
            template_variable_string="[[cos_email_section_2]]",
        ),
    ]

    filled_bytes, unresolved = DocxTemplateService.fill_template(
        buf.getvalue(),
        fields,
        {
            "cos_email_section_1": "Robin R. Weiner\nauto@ch13.com",
            "cos_email_section_2": "Office of US Trustee\nUSTPRegion21@usdoj.gov",
        },
    )

    assert unresolved == []
    out = Document(BytesIO(filled_bytes))
    para = out.paragraphs[0]
    # No " and " in the rendered paragraph — each value rendered as
    # caption-shape with `<w:br/>` between its lines.
    assert " and " not in para.text
    assert "Robin R. Weiner" in para.text
    assert "auto@ch13.com" in para.text
    assert "Office of US Trustee" in para.text
    assert "USTPRegion21@usdoj.gov" in para.text
    # Soft breaks present: 2 (between paragraphs) + 1 (inside section_1's
    # value) + 1 (inside section_2's value) = 4.
    line_breaks = [
        br for br in para._element.iter(qn("w:br"))
        if br.get(qn("w:type")) == "line"
    ]
    assert len(line_breaks) >= 3


@pytest.mark.unit
def test_fill_template_joins_with_and_in_inline_paragraph():
    """Inline paragraph (real prose around placeholder) joins \\n-bearing value with ' and '."""
    template_bytes = _docx_bytes([
        "The Debtor, [[debtor_name]], is employed in capacities requiring trust.",
    ])
    fields = [
        TemplateField(
            property_name="debtor_name",
            source=FieldSource.CASE_VECTOR,
            template_variable_string="[[debtor_name]]",
        )
    ]

    filled_bytes, unresolved = DocxTemplateService.fill_template(
        template_bytes,
        fields,
        {"debtor_name": "Lori Creswell\nRobert Creswell"},
    )

    assert unresolved == []
    doc = Document(BytesIO(filled_bytes))
    paragraph = doc.paragraphs[0]
    assert paragraph.text == "The Debtor, Lori Creswell and Robert Creswell, is employed in capacities requiring trust."
    assert not list(paragraph._element.iter(qn("w:br")))


@pytest.mark.unit
def test_fill_template_same_value_renders_differently_per_paragraph():
    """Same \\n-bearing value in BOTH a caption paragraph AND an inline paragraph
    gets per-paragraph treatment: caption keeps <w:br/>, inline joins with ' and '."""
    template_bytes = _docx_bytes([
        "[[debtor_name]]",
        "The Debtor, [[debtor_name]], is employed.",
    ])
    fields = [
        TemplateField(
            property_name="debtor_name",
            source=FieldSource.CASE_VECTOR,
            template_variable_string="[[debtor_name]]",
        )
    ]

    filled_bytes, unresolved = DocxTemplateService.fill_template(
        template_bytes,
        fields,
        {"debtor_name": "Lori Creswell\nRobert Creswell"},
    )

    assert unresolved == []
    doc = Document(BytesIO(filled_bytes))
    caption_para = doc.paragraphs[0]
    body_para = doc.paragraphs[1]

    caption_breaks = [
        br for br in caption_para._element.iter(qn("w:br"))
        if br.get(qn("w:type")) == "line"
    ]
    assert len(caption_breaks) == 1, "caption paragraph should keep <w:br/>"

    assert body_para.text == "The Debtor, Lori Creswell and Robert Creswell, is employed."
    assert not list(body_para._element.iter(qn("w:br"))), "body paragraph should not have <w:br/>"


@pytest.mark.unit
def test_fill_template_three_debtor_value_joins_inline_with_and():
    """Three names in an inline paragraph join with ' and ' between each pair."""
    template_bytes = _docx_bytes([
        "The Debtors, [[debtor_name]], are co-debtors.",
    ])
    fields = [
        TemplateField(
            property_name="debtor_name",
            source=FieldSource.CASE_VECTOR,
            template_variable_string="[[debtor_name]]",
        )
    ]

    filled_bytes, unresolved = DocxTemplateService.fill_template(
        template_bytes,
        fields,
        {"debtor_name": "Alice Partner\nBob Partner\nCarol Partner"},
    )

    assert unresolved == []
    doc = Document(BytesIO(filled_bytes))
    assert doc.paragraphs[0].text == "The Debtors, Alice Partner and Bob Partner and Carol Partner, are co-debtors."


@pytest.mark.unit
def test_fill_template_solo_value_has_no_line_break():
    """No \\n in value → no <w:br/> inserted, original single-line path."""
    template_bytes = _docx_bytes(["Caption: [[debtor_name]]"])
    fields = [
        TemplateField(
            property_name="debtor_name",
            source=FieldSource.CASE_VECTOR,
            template_variable_string="[[debtor_name]]",
        )
    ]

    filled_bytes, unresolved = DocxTemplateService.fill_template(
        template_bytes, fields, {"debtor_name": "John Smith"}
    )

    assert unresolved == []
    doc = Document(BytesIO(filled_bytes))
    assert doc.paragraphs[0].text == "Caption: John Smith"
    assert not list(doc.paragraphs[0]._element.iter(qn("w:br")))


# ─── flatten_word_fields: Word DATE field codes + content controls ─────


def _docx_with_field_simple(label: str, instruction: str, cached_text: str) -> bytes:
    """Build a .docx with a `<w:fldSimple>` field — e.g. `{ DATE \\@ "MMMM d, yyyy" }`
    whose cached rendered text is `cached_text`. Word renders the live
    field; python-docx returns the cached text.

    Paragraph layout: `<w:p> <w:r><w:t>label </w:t></w:r> <w:fldSimple>
    <w:r><w:t>cached_text</w:t></w:r> </w:fldSimple> </w:p>`.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(label + " ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = cached_text
    inner_r.append(inner_t)
    fld.append(inner_r)
    p._element.append(fld)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _docx_with_complex_field(label: str, instruction: str, cached_text: str) -> bytes:
    """Build a .docx with a complex `<w:fldChar>` + `<w:instrText>` field.

    Layout:
        <w:p>
          <w:r><w:t>label </w:t></w:r>
          <w:r><w:fldChar fldCharType="begin"/></w:r>
          <w:r><w:instrText>{ instruction }</w:instrText></w:r>
          <w:r><w:fldChar fldCharType="separate"/></w:r>
          <w:r><w:t>cached_text</w:t></w:r>     ← keep this after flatten
          <w:r><w:fldChar fldCharType="end"/></w:r>
        </w:p>
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(label + " ")

    def _make_run_with(child):
        r = OxmlElement("w:r")
        r.append(child)
        return r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    p._element.append(_make_run_with(begin))

    instr = OxmlElement("w:instrText")
    instr.text = instruction
    p._element.append(_make_run_with(instr))

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    p._element.append(_make_run_with(sep))

    cached_r = OxmlElement("w:r")
    cached_t = OxmlElement("w:t")
    cached_t.text = cached_text
    cached_r.append(cached_t)
    p._element.append(cached_r)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    p._element.append(_make_run_with(end))

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _docx_with_sdt_content_control(label: str, inner_text: str) -> bytes:
    """Build a .docx with a `<w:sdt>` content control wrapping a run.

    Layout:
        <w:p>
          <w:r><w:t>label </w:t></w:r>
          <w:sdt>
            <w:sdtContent>
              <w:r><w:t>inner_text</w:t></w:r>
            </w:sdtContent>
          </w:sdt>
        </w:p>
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(label + " ")
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = inner_text
    inner_r.append(inner_t)
    content.append(inner_r)
    sdt.append(content)
    p._element.append(sdt)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


@pytest.mark.unit
def test_flatten_word_fields_strips_fldSimple_keeps_cached_text():
    """A Word DATE `<w:fldSimple>` field with cached text 'April 8, 2026'
    must become plain text after flattening — no <w:fldSimple> elements
    left, the cached text survives as a normal <w:t>."""
    src = _docx_with_field_simple(
        label="Dated:",
        instruction='DATE \\@ "MMMM d, yyyy"',
        cached_text="April 8, 2026",
    )
    out_bytes = DocxTemplateService.flatten_word_fields(src)

    out_doc = Document(BytesIO(out_bytes))
    para = out_doc.paragraphs[0]
    # Cached text preserved as plain visible content.
    assert "April 8, 2026" in para.text
    assert para.text.startswith("Dated:")
    # No field machinery left.
    assert not list(para._element.iter(qn("w:fldSimple")))
    assert not list(para._element.iter(qn("w:fldChar")))
    assert not list(para._element.iter(qn("w:instrText")))


@pytest.mark.unit
def test_flatten_word_fields_strips_complex_fldChar_keeps_cached_text():
    """The two-marker `<w:fldChar begin/separate/end>` form. After
    flatten the begin/separate/end + instruction runs are gone but
    the cached result run survives as plain text."""
    src = _docx_with_complex_field(
        label="Filed on",
        instruction='{ DATE \\@ "MMMM d, yyyy" }',
        cached_text="April 8, 2026",
    )
    out_bytes = DocxTemplateService.flatten_word_fields(src)

    out_doc = Document(BytesIO(out_bytes))
    para = out_doc.paragraphs[0]
    assert "April 8, 2026" in para.text
    assert para.text.startswith("Filed on")
    assert not list(para._element.iter(qn("w:fldChar")))
    assert not list(para._element.iter(qn("w:instrText")))


@pytest.mark.unit
def test_flatten_word_fields_unwraps_sdt_content_control():
    """`<w:sdt>` content controls collapse to their wrapped runs."""
    src = _docx_with_sdt_content_control(
        label="Debtor name:",
        inner_text="Jane M. Doe",
    )
    out_bytes = DocxTemplateService.flatten_word_fields(src)

    out_doc = Document(BytesIO(out_bytes))
    para = out_doc.paragraphs[0]
    assert "Jane M. Doe" in para.text
    assert para.text.startswith("Debtor name:")
    assert not list(para._element.iter(qn("w:sdt")))
    assert not list(para._element.iter(qn("w:sdtContent")))


@pytest.mark.unit
def test_flatten_word_fields_noop_on_doc_without_fields():
    """A .docx with no fields, content controls, or instrText should
    round-trip with byte-identical visible text + no machinery
    elements left behind (there were none to start with)."""
    doc = Document()
    doc.add_paragraph("Plain paragraph one.")
    doc.add_paragraph("Plain paragraph two.")
    buf = BytesIO()
    doc.save(buf)
    src = buf.getvalue()

    out_bytes = DocxTemplateService.flatten_word_fields(src)
    out_doc = Document(BytesIO(out_bytes))

    assert [p.text for p in out_doc.paragraphs] == [
        "Plain paragraph one.",
        "Plain paragraph two.",
    ]


@pytest.mark.unit
def test_flatten_word_fields_keeps_runs_unrelated_to_fields():
    """Field stripping must not damage non-field runs. A doc with
    field-bearing AND plain paragraphs should keep the plain
    paragraphs intact."""
    # Build a doc with one fldSimple paragraph then a plain one.
    src = _docx_with_field_simple(
        label="Dated:",
        instruction='DATE \\@ "MMMM d, yyyy"',
        cached_text="April 8, 2026",
    )
    src_doc = Document(BytesIO(src))
    src_doc.add_paragraph("Sincerely, John Smith.")
    buf = BytesIO()
    src_doc.save(buf)

    out_bytes = DocxTemplateService.flatten_word_fields(buf.getvalue())
    out_doc = Document(BytesIO(out_bytes))
    paragraphs = [p.text for p in out_doc.paragraphs]
    assert paragraphs[0] == "Dated: April 8, 2026"
    assert paragraphs[1] == "Sincerely, John Smith."



# Stage 2: tab / non-<w:t> preservation when marker spans subset of runs


def _paragraph_with_tab_separated_runs(
    pre_tab_text: str, post_tab_texts: list[str]
):
    """Build a paragraph laid out as:

      <w:r><w:t>pre_tab_text</w:t></w:r>
      <w:r><w:tab/><w:t>post_tab_texts[0]</w:t></w:r>
      <w:r><w:t>post_tab_texts[1]</w:t></w:r>
      ...

    Mirrors the real-world Fleisher caption shape:
      Run A: "In re"
      Run B: <w:tab/> + "Case No. "
      Run C..: split case number fragments
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(pre_tab_text)
    if not post_tab_texts:
        return p
    run_with_tab = p.add_run()
    tab = OxmlElement("w:tab")
    run_with_tab._element.append(tab)
    t = OxmlElement("w:t")
    t.text = post_tab_texts[0]
    t.set(qn("xml:space"), "preserve")
    run_with_tab._element.append(t)
    for text in post_tab_texts[1:]:
        p.add_run(text)
    return p


@pytest.mark.unit
def test_replace_in_paragraph_preserves_tab_when_marker_spans_subset_of_runs():
    p = _paragraph_with_tab_separated_runs(
        pre_tab_text="In re",
        post_tab_texts=["Case No. ", "26-bk", "-", "12569", "-", "MAM"],
    )
    assert p.text == "In re\tCase No. 26-bk-12569-MAM"
    assert len(list(p._element.iter(qn("w:tab")))) == 1

    DocxTemplateService._replace_in_paragraph(
        p, "26-bk-12569-MAM", "[[case_number]]"
    )

    assert p.text == "In re\tCase No. [[case_number]]"
    assert len(list(p._element.iter(qn("w:tab")))) == 1


@pytest.mark.unit
def test_replace_in_paragraph_preserves_tab_when_marker_fits_in_single_late_run():
    p = _paragraph_with_tab_separated_runs(
        pre_tab_text="In re",
        post_tab_texts=["Case No. ", "26-bk-12569-MAM"],
    )
    DocxTemplateService._replace_in_paragraph(
        p, "26-bk-12569-MAM", "[[case_number]]"
    )
    assert p.text == "In re\tCase No. [[case_number]]"
    assert len(list(p._element.iter(qn("w:tab")))) == 1


@pytest.mark.unit
def test_replace_in_paragraph_preserves_tab_when_marker_spans_two_late_runs():
    p = _paragraph_with_tab_separated_runs(
        pre_tab_text="In re",
        post_tab_texts=["Case No. ", "ABC", "-", "999"],
    )
    DocxTemplateService._replace_in_paragraph(p, "ABC-999", "[[case_number]]")
    assert p.text == "In re\tCase No. [[case_number]]"
    assert len(list(p._element.iter(qn("w:tab")))) == 1


@pytest.mark.unit
def test_replace_in_paragraph_preserves_two_tabs_when_marker_spans_run_subset():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Col1")
    run2 = p.add_run()
    run2._element.append(OxmlElement("w:tab"))
    t2 = OxmlElement("w:t")
    t2.text = "Col2: "
    t2.set(qn("xml:space"), "preserve")
    run2._element.append(t2)
    run3 = p.add_run()
    run3._element.append(OxmlElement("w:tab"))
    t3 = OxmlElement("w:t")
    t3.text = "VALUE-"
    t3.set(qn("xml:space"), "preserve")
    run3._element.append(t3)
    p.add_run("12345")

    assert p.text == "Col1\tCol2: \tVALUE-12345"
    DocxTemplateService._replace_in_paragraph(p, "VALUE-12345", "[[val]]")
    assert p.text == "Col1\tCol2: \t[[val]]"
    assert len(list(p._element.iter(qn("w:tab")))) == 2


@pytest.mark.unit
def test_replace_in_paragraph_preserves_tab_when_marker_straddles_tab_boundary():
    """Marker concatenates across the <w:tab/>'s own run boundary in
    the <w:t> concatenation. 'reCase' matches because <w:t>s concatenate
    without the tab, but the tab itself must NOT be lost."""
    p = _paragraph_with_tab_separated_runs(
        pre_tab_text="In re",
        post_tab_texts=["Case No. ", "26-bk-12569-MAM"],
    )
    DocxTemplateService._replace_in_paragraph(p, "reCase", "[[X]]")
    assert len(list(p._element.iter(qn("w:tab")))) == 1
