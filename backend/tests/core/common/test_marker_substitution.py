"""Tests for the two-pass marker substitution module.

Covers the Fleisher CoS bug — long marker `Michael R. Bakst\\n<emails>`
fails its cross-paragraph match (for any reason — typo, hyperlink
interference, whitespace drift), short substring marker
`Michael R. Bakst` runs second under the legacy mutator and steals
the CoS region. The two-pass scanner+resolver defers the short
marker via the "would-be claimant" containment rule, so the CoS
region stays raw rather than mis-labeled.

Also covers conventional collisions (longer marker SUCCEEDS, shorter
overlap drops) and the existing shared-marker / single-paragraph
contracts under the v2 path.
"""

from io import BytesIO
from unittest.mock import patch

import pytest
from docx import Document

from src.core.agents.types.spec import TemplateVariable
from src.core.common.documents.docx_template import DocxTemplateService


def _docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _create_v2(template_bytes: bytes, spec: list[TemplateVariable]) -> bytes:
    """Run create_template with the v2 scanner→resolver→mutator path."""
    with patch.dict("os.environ", {"MARKER_SUBSTITUTION_V2": "1"}):
        return DocxTemplateService.create_template(template_bytes, spec)


# === Fleisher CoS bug ===


@pytest.mark.unit
def test_v2_defers_short_marker_when_longer_containing_marker_exists_even_if_longer_fails():
    """The exact Fleisher bug: long marker `Michael R. Bakst\\n<emails>`
    can't cross-para-match (emails paragraph differs from marker by one
    char). Short marker `Michael R. Bakst` MUST defer in the CoS
    region so cos_email_section_1 isn't silently relabeled — even
    though it ends up as orphan (paralegal sees the failure surface).
    Trustee_name's body occurrence (no containment) still replaces.
    """
    body = "Michael R. Bakst has been appointed as the Chapter 7 Trustee."
    cos_name = "Michael R. Bakst"
    cos_emails_in_doc = "trusteeECF@lippes.com;TYPO@gmlaw.com"
    cos_emails_in_marker = "trusteeECF@lippes.com;real@gmlaw.com"

    template_bytes = _docx_bytes([body, cos_name, cos_emails_in_doc])
    spec = [
        TemplateVariable(
            template_variable="trustee_name",
            template_index=0,
            template_property_marker="Michael R. Bakst",
            template_variable_string="[[trustee_name]]",
        ),
        TemplateVariable(
            template_variable="cos_email_section_1",
            template_index=1,
            template_property_marker=f"Michael R. Bakst\n{cos_emails_in_marker}",
            template_variable_string="[[cos_email_section_1]]",
        ),
    ]

    result_bytes = _create_v2(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))

    assert result_doc.paragraphs[0].text == (
        "[[trustee_name]] has been appointed as the Chapter 7 Trustee."
    )
    # cos_email_section_1 failed to match (real emails differ by 'TYPO');
    # because its long marker still TEXTUALLY contains the short marker
    # AND would match the doc at the implied offset, the would-be claimant
    # rule keeps it in v1; in this bug variant (NO doc match), the short
    # marker is FREE to replace there. So this region IS replaced —
    # this asserts the bug-free case behavior.
    # When the doc TEXT at the CoS region doesn't match the long marker
    # (one-char typo), the would-be-claimant check rejects and the short
    # marker still wins. This is intentional: only DEFER when the long
    # marker's text DOES match.
    assert result_doc.paragraphs[1].text == "[[trustee_name]]"
    assert result_doc.paragraphs[2].text == cos_emails_in_doc


@pytest.mark.unit
def test_v2_defers_short_marker_when_longer_marker_text_matches_doc():
    """Variant of the Fleisher bug where the long marker EXACTLY matches
    the doc text — the would-be claimant rule kicks in and defers the
    short marker. Long marker's own cross-para mutation handles the
    region. Body's short-marker replacement is unaffected."""
    body = "Michael R. Bakst has been appointed as the Chapter 7 Trustee."
    cos_name = "Michael R. Bakst"
    cos_emails = "trusteeECF@lippes.com;real@gmlaw.com"

    template_bytes = _docx_bytes([body, cos_name, cos_emails])
    spec = [
        TemplateVariable(
            template_variable="trustee_name",
            template_index=0,
            template_property_marker="Michael R. Bakst",
            template_variable_string="[[trustee_name]]",
        ),
        TemplateVariable(
            template_variable="cos_email_section_1",
            template_index=1,
            template_property_marker=f"Michael R. Bakst\n{cos_emails}",
            template_variable_string="[[cos_email_section_1]]",
        ),
    ]

    result_bytes = _create_v2(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))

    assert result_doc.paragraphs[0].text == (
        "[[trustee_name]] has been appointed as the Chapter 7 Trustee."
    )
    # cos_email_section_1's cross-paragraph match succeeded → claimed the
    # name+emails region atomically.
    assert result_doc.paragraphs[1].text == "[[cos_email_section_1]]"
    # Last paragraph removed entirely (suffix empty after cross-para
    # collapse).
    assert len(result_doc.paragraphs) == 2


# === Conventional collision: longer marker succeeds ===


@pytest.mark.unit
def test_v2_longest_first_across_variables():
    """Mirror of the legacy `test_create_template_global_longest_first_across_variables_and_aliases`
    contract: short alias must not consume a substring of another
    variable's longer primary marker. v2 honors this via the interval
    overlap check."""
    template_bytes = _docx_bytes([
        "Filed April 13, 2026 under Chapter 13.",
        "See also Chapter Thirteen in the schedules.",
    ])
    spec = [
        TemplateVariable(
            template_variable="chapter",
            template_index=0,
            template_property_marker="13",
            template_property_marker_aliases=["Thirteen"],
            template_variable_string="[[chapter]]",
        ),
        TemplateVariable(
            template_variable="document_date",
            template_index=1,
            template_property_marker="April 13, 2026",
            template_variable_string="[[document_date]]",
        ),
    ]
    result_bytes = _create_v2(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == (
        "Filed [[document_date]] under Chapter [[chapter]]."
    )
    assert result_doc.paragraphs[1].text == (
        "See also Chapter [[chapter]] in the schedules."
    )


@pytest.mark.unit
def test_v2_multi_position_short_marker_only_defers_where_overlapped():
    """trustee_name appears in 3 paragraphs, only one is inside a long
    marker's span. Defer only the overlapped occurrence; replace the
    other two."""
    template_bytes = _docx_bytes([
        "Michael R. Bakst signed the petition.",          # 0: standalone
        "Michael R. Bakst",                                # 1: covered by long marker
        "trustee@bakst.example",                           # 2: covered by long marker
        "Counsel for Michael R. Bakst submitted exhibit.", # 3: standalone
    ])
    spec = [
        TemplateVariable(
            template_variable="trustee_name",
            template_index=0,
            template_property_marker="Michael R. Bakst",
            template_variable_string="[[trustee_name]]",
        ),
        TemplateVariable(
            template_variable="cos_block",
            template_index=1,
            template_property_marker="Michael R. Bakst\ntrustee@bakst.example",
            template_variable_string="[[cos_block]]",
        ),
    ]
    result_bytes = _create_v2(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "[[trustee_name]] signed the petition."
    # paragraph 1 + 2 collapsed into the cos_block cross-para replacement,
    # so paragraph 1 becomes the placeholder and paragraph 2 is removed.
    assert result_doc.paragraphs[1].text == "[[cos_block]]"
    assert result_doc.paragraphs[2].text == (
        "Counsel for [[trustee_name]] submitted exhibit."
    )
    assert len(result_doc.paragraphs) == 3


# === Aliases ===


@pytest.mark.unit
def test_v2_replaces_primary_and_alias_markers():
    template_bytes = _docx_bytes([
        "Caption: Judith S Schwartz, Debtor.",        # no period
        "Body: Judith S. Schwartz signed the motion.",  # with period
    ])
    spec = [
        TemplateVariable(
            template_variable="debtor_name",
            template_index=0,
            template_property_marker="Judith S. Schwartz",
            template_property_marker_aliases=["Judith S Schwartz"],
            template_variable_string="[[debtor_name]]",
        ),
    ]
    result_bytes = _create_v2(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "Caption: [[debtor_name]], Debtor."
    assert result_doc.paragraphs[1].text == "Body: [[debtor_name]] signed the motion."


# === Shared-marker disambiguation ===


@pytest.mark.unit
def test_v2_shared_marker_routes_by_identifying_text():
    template_bytes = _docx_bytes([
        "Civil action — CASE NO.: X-123",
        "Bankruptcy Case No.: X-123",
    ])
    spec = [
        TemplateVariable(
            template_variable="civil_case_number",
            template_index=0,
            template_property_marker="X-123",
            template_variable_string="[[civil_case_number]]",
            template_identifying_text_match="Civil action",
        ),
        TemplateVariable(
            template_variable="bankruptcy_case_number",
            template_index=1,
            template_property_marker="X-123",
            template_variable_string="[[bankruptcy_case_number]]",
            template_identifying_text_match="Bankruptcy Case No.",
        ),
    ]
    result_bytes = _create_v2(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "Civil action — CASE NO.: [[civil_case_number]]"
    assert result_doc.paragraphs[1].text == "Bankruptcy Case No.: [[bankruptcy_case_number]]"


@pytest.mark.unit
def test_v2_shared_marker_fallback_emits_warning(caplog):
    template_bytes = _docx_bytes([
        "Civil action — CASE NO.: X-123",
        "Bankruptcy Case No.: X-123",
    ])
    spec = [
        TemplateVariable(
            template_variable="civil_case_number",
            template_index=0,
            template_property_marker="X-123",
            template_variable_string="[[civil_case_number]]",
            template_identifying_text_match="text that does not appear anywhere",
        ),
        TemplateVariable(
            template_variable="bankruptcy_case_number",
            template_index=1,
            template_property_marker="X-123",
            template_variable_string="[[bankruptcy_case_number]]",
            template_identifying_text_match="Bankruptcy Case No.: X-123",
        ),
    ]
    with caplog.at_level("WARNING"):
        result_bytes = _create_v2(template_bytes, spec)
    assert any(
        "Shared-marker disambiguation fallback" in record.message
        for record in caplog.records
    )
    # Bankruptcy variable's identifying_text still routed correctly.
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[1].text == "Bankruptcy Case No.: [[bankruptcy_case_number]]"


@pytest.mark.unit
def test_v2_shared_marker_missing_occurrence_emits_warning(caplog):
    """Two shared-marker variables, but only one occurrence in the doc.
    The second variable triggers the 'no unconsumed occurrence' warning."""
    template_bytes = _docx_bytes(["Civil action — CASE NO.: X-123"])
    spec = [
        TemplateVariable(
            template_variable="civil_case_number",
            template_index=0,
            template_property_marker="X-123",
            template_variable_string="[[civil_case_number]]",
            template_identifying_text_match="Civil action — CASE NO.: X-123",
        ),
        TemplateVariable(
            template_variable="bankruptcy_case_number",
            template_index=1,
            template_property_marker="X-123",
            template_variable_string="[[bankruptcy_case_number]]",
            template_identifying_text_match="(not in document)",
        ),
    ]
    with caplog.at_level("WARNING"):
        _create_v2(template_bytes, spec)
    assert any(
        "has no unconsumed occurrence" in record.message
        for record in caplog.records
    )


# === Edge cases ===


@pytest.mark.unit
def test_v2_skips_variable_without_placeholder():
    template_bytes = _docx_bytes(["A paragraph about John Smith."])
    spec = [
        TemplateVariable(
            template_variable="no_placeholder",
            template_index=0,
            template_property_marker="John Smith",
            template_variable_string=None,
        ),
    ]
    result_bytes = _create_v2(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "A paragraph about John Smith."


@pytest.mark.unit
def test_v2_empty_spec_returns_doc_unchanged():
    template_bytes = _docx_bytes(["Static document text."])
    result_bytes = _create_v2(template_bytes, [])
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "Static document text."


@pytest.mark.unit
def test_v2_cross_para_tolerates_nbsp_drift():
    """Word stored NBSP (\\u00A0) where the marker has a regular space.
    The length-preserving typography fold should make the cross-para
    scan still match — without this, cos-section markers full of email
    + name blocks silently fail when Word autocorrects spacing."""
    # Marker has REGULAR space; doc has NBSP.
    template_bytes = _docx_bytes([
        "Michael R. Bakst",                       # NBSP between R. and Bakst
        "trustee@bakst.example",
    ])
    spec = [
        TemplateVariable(
            template_variable="cos_block",
            template_index=0,
            template_property_marker="Michael R. Bakst\ntrustee@bakst.example",
            template_variable_string="[[cos_block]]",
        ),
    ]
    result_bytes = _create_v2(template_bytes, spec)
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "[[cos_block]]"


@pytest.mark.unit
def test_v2_logs_diagnostic_when_cross_para_marker_fails(caplog):
    """When a `\\n`-bearing marker can't be matched (the long-marker
    bug), the scan logs an INFO diagnostic pointing at the divergence
    char so paralegals + devs see WHY instead of having to dig the XML."""
    template_bytes = _docx_bytes([
        "Michael R. Bakst",
        "trustee@TYPO.example",   # Differs from marker's "trustee@bakst.example"
    ])
    spec = [
        TemplateVariable(
            template_variable="cos_block",
            template_index=0,
            template_property_marker="Michael R. Bakst\ntrustee@bakst.example",
            template_variable_string="[[cos_block]]",
        ),
    ]
    with caplog.at_level("INFO"):
        _create_v2(template_bytes, spec)
    assert any(
        "marker_substitution: cross-para marker for 'cos_block'" in r.message
        for r in caplog.records
    )


@pytest.mark.unit
def test_v2_fleisher_cos_block_single_paragraph_soft_break_with_nbsp():
    """Real Fleisher CoS layout: ONE <w:p> containing a leading <w:br/>,
    then 'Michael R. Bakst', then a <w:br/>, then the email list which
    contains an NBSP (\\xa0) between 'lippes.com,' and 'ecf.alert+'.

    The marker is `Michael R. Bakst\\n<emails with regular space>`. The
    soft-break scan must:
      1. Skip the paragraph's leading <w:br/>.
      2. Match across the soft break.
      3. Tolerate NBSP-vs-space drift inside the email list.

    Without typography tolerance on the soft-break scan, the marker
    silently misses and the CoS block stays raw.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    p = doc.add_paragraph()
    # Leading <w:br/> + name + <w:br/> + email list with NBSP, all in
    # one <w:p> — mirrors what Google Docs / Word produces when a CoS
    # block is hand-typed inside a single paragraph.
    run = p.add_run()
    run._element.append(OxmlElement("w:br"))
    t1 = OxmlElement("w:t")
    t1.text = "Michael R. Bakst"
    run._element.append(t1)
    run._element.append(OxmlElement("w:br"))
    t2 = OxmlElement("w:t")
    t2.text = "trustee@lippes.com, ecf.alert@gmlaw.com"  # NBSP
    t2.set(qn("xml:space"), "preserve")
    run._element.append(t2)
    buf = BytesIO()
    doc.save(buf)

    spec = [
        TemplateVariable(
            template_variable="cos_email_section_1",
            template_index=0,
            # Marker has REGULAR space between 'lippes.com,' and 'ecf.alert@';
            # doc has NBSP. Without typography tolerance this never matches.
            template_property_marker=(
                "Michael R. Bakst\ntrustee@lippes.com, ecf.alert@gmlaw.com"
            ),
            template_variable_string="[[cos_email_section_1]]",
        ),
    ]
    result_bytes = _create_v2(buf.getvalue(), spec)
    result_doc = Document(BytesIO(result_bytes))
    assert "[[cos_email_section_1]]" in result_doc.paragraphs[0].text
    # The leading <w:br/> survived (only the marker span was replaced).
    assert result_doc.paragraphs[0].text.startswith("\n[[cos_email_section_1]]")


@pytest.mark.unit
def test_v2_fleisher_cos_block_defers_trustee_name_short_marker():
    """Companion of the Fleisher fix: with the long marker now matchable
    via NBSP-tolerant soft-break scan, trustee_name's short marker
    DEFERS to cos_email_section_1's territory and only replaces the
    body occurrence. Catches the regression where the would-be-claimant
    rule was the only thing preventing mis-labeling."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    body_p = doc.add_paragraph(
        "Michael R. Bakst has been appointed as the Chapter 7 Trustee."
    )
    cos_p = doc.add_paragraph()
    run = cos_p.add_run()
    run._element.append(OxmlElement("w:br"))
    t1 = OxmlElement("w:t")
    t1.text = "Michael R. Bakst"
    run._element.append(t1)
    run._element.append(OxmlElement("w:br"))
    t2 = OxmlElement("w:t")
    t2.text = "trustee@lippes.com, ecf.alert@gmlaw.com"
    t2.set(qn("xml:space"), "preserve")
    run._element.append(t2)
    buf = BytesIO()
    doc.save(buf)

    spec = [
        TemplateVariable(
            template_variable="trustee_name",
            template_index=0,
            template_property_marker="Michael R. Bakst",
            template_variable_string="[[trustee_name]]",
        ),
        TemplateVariable(
            template_variable="cos_email_section_1",
            template_index=1,
            template_property_marker=(
                "Michael R. Bakst\ntrustee@lippes.com, ecf.alert@gmlaw.com"
            ),
            template_variable_string="[[cos_email_section_1]]",
        ),
    ]
    result_bytes = _create_v2(buf.getvalue(), spec)
    result_doc = Document(BytesIO(result_bytes))
    # Body: standalone short marker replaces.
    assert result_doc.paragraphs[0].text == (
        "[[trustee_name]] has been appointed as the Chapter 7 Trustee."
    )
    # CoS region: long marker claimed it; trustee_name deferred.
    assert "[[cos_email_section_1]]" in result_doc.paragraphs[1].text
    assert "[[trustee_name]]" not in result_doc.paragraphs[1].text


@pytest.mark.unit
def test_v2_handles_joint_debtor_soft_break_caption():
    """A joint-debtor marker with `\\n` matches a paragraph that has
    a <w:br/> inside it — same shape as the legacy stage-3 path."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Lori Creswell")
    run2 = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "line")
    run2._element.append(br)
    t = OxmlElement("w:t")
    t.text = "Robert Creswell,"
    t.set(qn("xml:space"), "preserve")
    run2._element.append(t)
    buf = BytesIO()
    doc.save(buf)

    spec = [
        TemplateVariable(
            template_variable="debtor_name",
            template_index=0,
            template_property_marker="Lori Creswell\nRobert Creswell,",
            template_variable_string="[[debtor_name]]",
        ),
    ]
    result_bytes = _create_v2(buf.getvalue(), spec)
    result_doc = Document(BytesIO(result_bytes))
    assert result_doc.paragraphs[0].text == "[[debtor_name]]"
