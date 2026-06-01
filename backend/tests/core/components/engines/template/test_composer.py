"""Tests for composer-level helpers.

Focused on pure-function helpers like _enforce_merges (post-agent merge
validation + stale-source cleanup). Full regenerate_template coverage
requires R2 + repo + agent mocks and lives in a future integration pass.
"""

import pytest
from fastapi import HTTPException

from src.core.agents.types.spec import TemplateVariable
from src.core.components.engines.template.composer import _enforce_merges
from src.core.components.engines.template.schemas import MergeOperation


def _var(name: str, marker: str = "") -> TemplateVariable:
    return TemplateVariable(
        template_variable=name,
        template_index=0,
        template_property_marker=marker,
        template_variable_string=f"[[{name}]]",
    )


@pytest.mark.unit
def test_enforce_merges_noop_when_merges_empty():
    spec = [_var("a"), _var("b")]
    result = _enforce_merges(spec, merges=[])
    assert result == spec


@pytest.mark.unit
def test_enforce_merges_drops_stale_sources_when_merged_present():
    """Agent followed the merge instruction but forgot to remove the source
    variables. _enforce_merges cleans them up."""
    spec = [
        _var("ecf_number", marker="3"),
        _var("document_title", marker="Notice of Appearance"),
        _var("ecf_number_document_title", marker="3 being a document Notice of Appearance"),
        _var("debtor_name", marker="Jane Doe"),
    ]

    result = _enforce_merges(
        spec,
        merges=[
            MergeOperation(source_variables=["ecf_number", "document_title"]),
        ],
    )

    names = [v.template_variable for v in result]
    assert "ecf_number" not in names
    assert "document_title" not in names
    assert "ecf_number_document_title" in names
    assert "debtor_name" in names  # untouched


@pytest.mark.unit
def test_enforce_merges_raises_when_merged_variable_missing():
    """Agent failed to produce the merged variable — error so caller can retry."""
    spec = [
        _var("ecf_number", marker="3"),
        _var("document_title", marker="Notice of Appearance"),
    ]

    with pytest.raises(HTTPException) as exc:
        _enforce_merges(
            spec,
            merges=[
                MergeOperation(source_variables=["ecf_number", "document_title"]),
            ],
        )

    assert exc.value.status_code == 422
    errors = exc.value.detail["merge_errors"]
    assert any("ecf_number_document_title" in e and "missing" in e for e in errors)


@pytest.mark.unit
def test_enforce_merges_handles_multiple_merges():
    spec = [
        _var("a1"), _var("a2"), _var("a1_a2"),
        _var("b1"), _var("b2"), _var("b1_b2"),
        _var("unrelated"),
    ]

    result = _enforce_merges(
        spec,
        merges=[
            MergeOperation(source_variables=["a1", "a2"]),
            MergeOperation(source_variables=["b1", "b2"]),
        ],
    )

    names = [v.template_variable for v in result]
    assert set(names) == {"a1_a2", "b1_b2", "unrelated"}


@pytest.mark.unit
def test_enforce_merges_preserves_untouched_variables_unchanged():
    a = _var("ecf_number", marker="3")
    b = _var("document_title", marker="Notice of Appearance")
    merged = _var("ecf_number_document_title", marker="3 being a document Notice of Appearance")
    untouched = _var("case_number", marker="26-10700-SMG")

    result = _enforce_merges(
        [a, b, merged, untouched],
        merges=[
            MergeOperation(source_variables=["ecf_number", "document_title"]),
        ],
    )

    # untouched variable is returned intact (same instance)
    by_name = {v.template_variable: v for v in result}
    assert by_name["case_number"] is untouched
    assert by_name["ecf_number_document_title"] is merged


@pytest.mark.unit
def test_merge_operation_derives_name_from_source_variables():
    merge = MergeOperation(source_variables=["ecf_number", "document_title"])
    assert merge.resolve_variable_name() == "ecf_number_document_title"


@pytest.mark.unit
def test_merge_operation_rejects_extras_including_legacy_new_variable_name():
    """No manual override: extras (including the old new_variable_name field)
    must fail validation — keeps the contract narrow."""
    from pydantic import ValidationError

    with pytest.raises(ValidationError):
        MergeOperation.model_validate({
            "source_variables": ["a", "b"],
            "new_variable_name": "custom_name",
        })


@pytest.mark.unit
def test_enforce_merges_works_with_auto_generated_name():
    """Source variables dropped + auto-generated merged variable recognized in output."""
    spec = [
        _var("ecf_number", marker="3"),
        _var("document_title", marker="Notice of Appearance"),
        _var("ecf_number_document_title", marker="3 being a document Notice of Appearance"),
    ]

    result = _enforce_merges(
        spec,
        merges=[
            MergeOperation(source_variables=["ecf_number", "document_title"]),
        ],
    )

    names = [v.template_variable for v in result]
    assert names == ["ecf_number_document_title"]


@pytest.mark.unit
def test_enforce_merges_tolerates_agent_already_removing_sources():
    """Agent may have correctly removed the sources on its own — _enforce_merges
    should still succeed (merged variable is present)."""
    spec = [
        _var("ecf_number_document_title", marker="3 being a document Notice of Appearance"),
        _var("debtor_name", marker="Jane Doe"),
    ]

    result = _enforce_merges(
        spec,
        merges=[
            MergeOperation(source_variables=["ecf_number", "document_title"]),
        ],
    )

    names = [v.template_variable for v in result]
    assert set(names) == {"ecf_number_document_title", "debtor_name"}


# ─── _ensure_joint_debtor_variable — deterministic joint-caption detector ──


from io import BytesIO  # noqa: E402

from docx import Document as _Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from src.core.components.engines.template.composer import _ensure_joint_debtor_variable  # noqa: E402


def _caption_docx(body_paragraphs: list[str]) -> bytes:
    doc = _Document()
    for t in body_paragraphs:
        doc.add_paragraph(t)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _caption_docx_with_soft_break(pre: list[str], names: list[str], post: list[str]) -> bytes:
    """Build a docx where `names` share one paragraph joined by <w:br/>."""
    doc = _Document()
    for t in pre:
        doc.add_paragraph(t)
    p = doc.add_paragraph()
    p.add_run(names[0])
    for name in names[1:]:
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "line")
        run._element.append(br)
        t_elem = OxmlElement("w:t")
        t_elem.text = name
        t_elem.set(qn("xml:space"), "preserve")
        run._element.append(t_elem)
    for t in post:
        doc.add_paragraph(t)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.unit
def test_ensure_joint_debtor_detects_multi_paragraph_caption():
    """Classic shape: 'In re:' + two name paragraphs + 'Debtors.' line."""
    docx = _caption_docx([
        "In re:",
        "Lori Creswell",
        "Robert Creswell,",
        "    Debtors.",
    ])
    result = _ensure_joint_debtor_variable(docx, template_spec=[])
    assert len(result) == 1
    assert result[0].template_variable == "debtor_name"
    assert result[0].template_property_marker == "Lori Creswell\nRobert Creswell,"


@pytest.mark.unit
def test_ensure_joint_debtor_detects_soft_break_joined_names():
    """Real Creswell shape: names share ONE paragraph with <w:br/> between them."""
    docx = _caption_docx_with_soft_break(
        pre=["In re:"],
        names=["Lori Creswell", "Robert Creswell,"],
        post=["Debtors."],
    )
    result = _ensure_joint_debtor_variable(docx, template_spec=[])
    assert len(result) == 1
    assert result[0].template_property_marker == "Lori Creswell\nRobert Creswell,"


@pytest.mark.unit
def test_ensure_joint_debtor_detects_in_re_glued_to_case_no():
    """'In re:\\tCase No: ...' glued header + 'Chapter N' metadata line + names."""
    docx = _caption_docx([
        "In re:\tCase No: 25-14980-PDR",
        "Chapter 13",
        "Lori Creswell",
        "Robert Creswell,",
        "Debtors.",
    ])
    result = _ensure_joint_debtor_variable(docx, template_spec=[])
    assert len(result) == 1
    assert result[0].template_property_marker == "Lori Creswell\nRobert Creswell,"


@pytest.mark.unit
def test_ensure_joint_debtor_replaces_broken_existing_variable():
    """Agent emitted a single-name marker — detector overrides with correct joint marker."""
    docx = _caption_docx([
        "In re:",
        "Lori Creswell",
        "Robert Creswell,",
        "Debtors.",
    ])
    existing = TemplateVariable(
        template_variable="debtor_name",
        template_index=0,
        template_property_marker="Lori Creswell",
        template_variable_string="[[debtor_name]]",
    )

    result = _ensure_joint_debtor_variable(docx, [existing])
    assert len(result) == 1
    assert result[0].template_property_marker == "Lori Creswell\nRobert Creswell,"


@pytest.mark.unit
def test_ensure_joint_debtor_idempotent_when_marker_already_correct():
    """Agent already emitted the correct joint marker — no churn."""
    docx = _caption_docx([
        "In re:",
        "Lori Creswell",
        "Robert Creswell,",
        "Debtors.",
    ])
    correct = TemplateVariable(
        template_variable="debtor_name",
        template_index=0,
        template_property_marker="Lori Creswell\nRobert Creswell,",
        template_variable_string="[[debtor_name]]",
    )
    result = _ensure_joint_debtor_variable(docx, [correct])
    assert result[0] is correct


@pytest.mark.unit
def test_ensure_joint_debtor_noop_on_solo_caption():
    """One debtor before 'Debtor.' → no synthesis, spec unchanged."""
    docx = _caption_docx([
        "In re:",
        "John Smith,",
        "Debtor.",
    ])
    result = _ensure_joint_debtor_variable(docx, template_spec=[])
    assert result == []


@pytest.mark.unit
def test_ensure_joint_debtor_noop_when_no_in_re_block():
    """No 'In re:' paragraph → spec unchanged."""
    docx = _caption_docx([
        "NOTICE OF WITHDRAWAL",
        "Motion body content with no caption.",
    ])
    result = _ensure_joint_debtor_variable(docx, template_spec=[])
    assert result == []


@pytest.mark.unit
def test_ensure_joint_debtor_handles_three_debtor_caption():
    """Multi-party filing: 3 names produce a 3-line marker."""
    docx = _caption_docx([
        "In re:",
        "Alice Partner",
        "Bob Partner",
        "Carol Partner,",
        "Debtors.",
    ])
    result = _ensure_joint_debtor_variable(docx, template_spec=[])
    assert len(result) == 1
    assert result[0].template_property_marker == "Alice Partner\nBob Partner\nCarol Partner,"


# ─── _drop_orphan_variables — virtual + auto_derive carve-outs ────────


from io import BytesIO  # noqa: E402

from docx import Document as _DocxDocument  # noqa: E402

from src.core.agents.types.sources import (  # noqa: E402
    AutoDerivedSourceParams,
    CaseVectorSourceParams,
    DropdownEmailSourceParams,
    FieldSource,
)
from src.core.components.engines.template.composer import (  # noqa: E402
    _drop_orphan_variables,
    build_agent_config,
)


def _docx_with_paragraphs(paragraphs: list[str]) -> bytes:
    doc = _DocxDocument()
    for t in paragraphs:
        doc.add_paragraph(t)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.unit
def test_drop_orphan_preserves_virtual_parent_referenced_by_auto_derive_child():
    """A virtual parent has no [[placeholder]] in the docx by design.
    With at least one auto_derive child depending on it, the orphan-drop
    must keep it; without the carve-out it'd get nuked."""
    template_bytes = _docx_with_paragraphs(["Row claim: [[claim_number]]"])

    spec = [
        TemplateVariable(
            template_variable="proof_of_claim_row",
            template_index=0,
            template_property_marker="4 - Bank of America - $3,000",
            template_variable_string=None,  # virtual
            source=FieldSource.DROPDOWN_FROM_GMAIL,
            source_params=DropdownEmailSourceParams(
                label="Proof of Claim",
                example_format="Claim 4 - Bank of America - $3,000",
            ),
        ),
        TemplateVariable(
            template_variable="claim_number",
            template_index=1,
            template_property_marker="4",
            template_variable_string="[[claim_number]]",
            source=FieldSource.AUTO_DERIVED_FROM_VARIABLE,
            source_params=AutoDerivedSourceParams(dependent_variable="proof_of_claim_row"),
        ),
    ]
    result = _drop_orphan_variables(template_bytes, spec)
    names = [v.template_variable for v in result]
    assert "proof_of_claim_row" in names
    assert "claim_number" in names


@pytest.mark.unit
def test_drop_orphan_drops_truly_orphaned_physical_variable():
    """Physical variable whose [[placeholder]] is missing AND no auto_derive
    child depends on it → still dropped (regression)."""
    template_bytes = _docx_with_paragraphs(["The Debtor, [[real_var]], filed."])
    spec = [
        TemplateVariable(
            template_variable="real_var",
            template_index=0,
            template_property_marker="x",
            template_variable_string="[[real_var]]",
        ),
        TemplateVariable(
            template_variable="ghost_var",
            template_index=1,
            template_property_marker="never appears",
            template_variable_string="[[ghost_var]]",
        ),
    ]
    result = _drop_orphan_variables(template_bytes, spec)
    names = [v.template_variable for v in result]
    assert names == ["real_var"]


@pytest.mark.unit
def test_drop_orphan_preserves_physical_parent_when_child_depends_even_if_marker_missing():
    """A physical parent whose marker got swallowed should NOT be dropped if a
    child depends on it — defense-in-depth."""
    template_bytes = _docx_with_paragraphs(["The Debtor, [[child_var]], filed."])
    spec = [
        TemplateVariable(
            template_variable="parent_var",
            template_index=0,
            template_property_marker="parent_marker",
            template_variable_string="[[parent_var]]",  # missing in docx
        ),
        TemplateVariable(
            template_variable="child_var",
            template_index=1,
            template_property_marker="child_marker",
            template_variable_string="[[child_var]]",
            source=FieldSource.AUTO_DERIVED_FROM_VARIABLE,
            source_params=AutoDerivedSourceParams(dependent_variable="parent_var"),
        ),
    ]
    result = _drop_orphan_variables(template_bytes, spec)
    names = [v.template_variable for v in result]
    assert "parent_var" in names  # preserved despite missing placeholder
    assert "child_var" in names


# ─── case_vector source_params preservation ────────────────────────────


from unittest.mock import AsyncMock  # noqa: E402

from src.core.common.storage.database import ReferenceDataRepository  # noqa: E402


@pytest.mark.unit
async def test_build_agent_config_preserves_case_vector_text_query(monkeypatch):
    """Regression: composer used to nuke source_params for every case_vector
    field unconditionally, which silently broke both the pgvector retrieval
    handler (fell back to property-name-derived query) AND the
    CaseVectorVisionAgent prompt (no `topical query` line). Now the author's
    `CaseVectorSourceParams.text_query` must survive into the runtime
    AgentConfig so both paths see it."""
    monkeypatch.setattr(ReferenceDataRepository, "list", AsyncMock(return_value=[]))

    template_spec = [
        TemplateVariable(
            template_variable="plaintiff_name",
            template_index=0,
            template_property_marker="JP Morgan Chase Bank",
            template_variable_string="[[plaintiff_name]]",
            source=FieldSource.CASE_VECTOR,
            source_params=CaseVectorSourceParams(
                text_query="Case filed against the debtor under SOFA question number 9",
            ),
        ),
    ]

    config = await build_agent_config("tpl_test", template_spec)

    assert len(config.template_fields) == 1
    field = config.template_fields[0]
    assert field.property_name == "plaintiff_name"
    assert field.source == FieldSource.CASE_VECTOR
    # The author's topical query must reach the runtime field; the
    # pgvector handler + vision agent both read it from here.
    assert isinstance(field.source_params, CaseVectorSourceParams)
    assert field.source_params.text_query == (
        "Case filed against the debtor under SOFA question number 9"
    )


@pytest.mark.unit
async def test_build_agent_config_leaves_case_vector_source_params_none_when_unset(monkeypatch):
    """When the author didn't supply explicit source_params on a case_vector
    field (auto-derived-from-variable-name behavior), `source_params` stays
    None — the pgvector handler falls back to the property name."""
    monkeypatch.setattr(ReferenceDataRepository, "list", AsyncMock(return_value=[]))

    template_spec = [
        TemplateVariable(
            template_variable="debtor_name",
            template_index=0,
            template_property_marker="John Smith",
            template_variable_string="[[debtor_name]]",
            source=FieldSource.CASE_VECTOR,
            source_params=None,
        ),
    ]

    config = await build_agent_config("tpl_test", template_spec)

    field = config.template_fields[0]
    assert field.source == FieldSource.CASE_VECTOR
    assert field.source_params is None
